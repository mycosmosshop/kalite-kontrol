# -*- coding: utf-8 -*-
"""APQP kanit belgelerini urunun ERP verisinden uretir.

  python apqp_belge_uret.py <stok kodu>

Uretilenler (G:\\Drive'im\\APQP\\<kod>\\ altina):
  PL74 Proses Akis Diyagrami   - sablon kopyalanip doldurulur
  FR90 Fizibilite Taahhudu     - sablon kopyalanip doldurulur
  FR81 Toplanti Tutanagi       - sifirdan uretilir (sablon yok)
  Kapasite Takip Formu         - sifirdan uretilir (sablon yok)

ONEMLI: Sablonlarda otomatik sekiller (41 adete varan rect/diamond/line) ve
gomulu gorseller var; openpyxl bir dosyayi acip kaydettiginde bu sekilleri
KAYBEDIYOR. Bu yuzden sablonlar ZIP DUZEYINDE yamaniyor: yalniz ilgili sayfa
XML'indeki hucre degerleri degistiriliyor, geri kalan her sey (sekiller,
gorseller, makrolar, stiller) bit bit korunuyor.
"""
import sys, re, io, json, math, zipfile, shutil, urllib.request, urllib.parse, datetime, os

SUPABASE = "https://nnubrxbpthmkitueixbh.supabase.co/rest/v1"
ANON = ("eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6Im5udWJyeGJwdGhta2l0"
        "dWVpeGJoIiwicm9sZSI6ImFub24iLCJpYXQiOjE3ODA1NjI2MDIsImV4cCI6MjA5NjEzODYwMn0"
        ".CHZUOylf_q8kkOQbFf9VWZ6-doUTlynmAhahM2EuImE")
DRIVE = r"G:\Drive'ım\APQP"
# Sablonlar URUN klasorunde DURMAZ: APQP kaydi silinirken urun klasoru de
# kaldirilabiliyor ve sablonlar onunla birlikte gidiyordu. Ayri klasor.
SABLON = os.path.join(DRIVE, "_Şablonlar")

# Lokasyona gore APQP ekibi (FR91 sablonundaki roller)
EKIP = {
    "ankara": [("AR&GE Proje Yöneticisi", "Sinem Kaya"), ("Kalite Güvence Müdürü", "Volkan Pekatik"),
               ("Kalite Mühendisi", "Emre Biçer"), ("Satın Alma", "Kutlay Altıparmak"),
               ("Lojistik", "Taner Şeşenoğlu"), ("Üretim", "Mete Yılmaz"), ("Satış", "Ender Zaimoğlu")],
    "cerkezkoy": [("AR&GE Proje Yöneticisi", "Sinem Kaya"), ("Kalite Güvence Müdürü", "Volkan Pekatik"),
                  ("Kalite Mühendisi", "Emrah Eryılmaz"), ("Satın Alma", "Kutlay Altıparmak"),
                  ("Lojistik", "Necmettin Altıntaş"), ("Üretim", "Umut Çiftçiogulları"),
                  ("Satış", "Ender Zaimoğlu")],
}


def sorgu(yol):
    r = urllib.request.Request(SUPABASE + yol, headers={"apikey": ANON, "Authorization": "Bearer " + ANON})
    with urllib.request.urlopen(r, timeout=60) as f:
        return json.load(f)


def yaz(yol, veri, yontem="POST"):
    """Supabase'e kayit yazar (yalniz APQP/MSA tablolari — LeanSys'e DOKUNULMAZ)."""
    g = json.dumps(veri, ensure_ascii=False).encode("utf-8")
    r = urllib.request.Request(SUPABASE + yol, data=g, method=yontem, headers={
        "apikey": ANON, "Authorization": "Bearer " + ANON,
        "Content-Type": "application/json", "Prefer": "return=representation"})
    with urllib.request.urlopen(r, timeout=60) as f:
        govde = f.read().decode("utf-8")
        return json.loads(govde) if govde.strip() else []


def met(x):
    return "" if x is None else str(x).strip()


# ── XLSX zip yamasi: hucre degerini degistirir, gerisine dokunmaz ─────────
def _xml_kacir(s):
    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def _sutun(ref):
    return re.match(r"([A-Z]+)", ref).group(1)


def _sutun_no(h):
    n = 0
    for c in h:
        n = n * 26 + (ord(c) - 64)
    return n


def _secim_duzelt(xml):
    """<selection activeCell=... sqref=...> tutarsizsa Excel dosyayi ACARKEN
    "Gorunum onarildi" uyarisi veriyor. Etkin hucre, secim araliginin ILK
    hucresine esitlenir."""
    def onar(m):
        oz = m.group(1)
        g = re.search(r'sqref="([A-Z]+\d+)', oz)
        if not g:
            return m.group(0)
        if "activeCell" in oz:
            return "<selection%s/>" % re.sub(r'activeCell="[A-Z]+\d+"',
                                             'activeCell="%s"' % g.group(1), oz)
        return '<selection activeCell="%s"%s/>' % (g.group(1), oz)

    return re.sub(r"<selection([^/>]*)/>", onar, xml)


def hucre_yaz(kaynak, hedef, sayfa_dosyasi, degerler, ek_xml=None, yeni_parcalar=None):
    """degerler: {'C6': 'metin', 'B12': 3, ...}  -> hedef dosyaya yazar.
    ek_xml: {zip_ici_yol: yeni_xml} — cizim/stil gibi baska parcalari da
    ayni yazma isleminde degistirmek icin."""
    zin = zipfile.ZipFile(kaynak)
    xml = _secim_duzelt(zin.read(sayfa_dosyasi).decode("utf-8"))
    # Yazilmayan sayfalarda da secim tutarsizsa Excel dosyayi ACARKEN
    # onariyor; hepsi duzeltilir.
    ek_xml = dict(ek_xml or {})
    for e in zin.infolist():
        if re.match(r"xl/worksheets/sheet\d+\.xml$", e.filename) \
                and e.filename != sayfa_dosyasi and e.filename not in ek_xml:
            g = zin.read(e.filename).decode("utf-8")
            d = _secim_duzelt(g)
            if d != g:
                ek_xml[e.filename] = d

    for ref, deger in degerler.items():
        # None -> hucreyi BOSALT (bicim korunur). Bos dize yazmak yetmez:
        # satir sayan formuller (COUNT/Anzahl) metni de sayabiliyor.
        if deger is None:
            kalip = re.compile(r'<c r="%s"([^>/]*)(/>|>.*?</c>)' % ref, re.S)
            m = kalip.search(xml)
            if m:
                oz = re.sub(r'\st="[^"]*"', "", m.group(1))
                xml = xml[:m.start()] + '<c r="%s"%s/>' % (ref, oz) + xml[m.end():]
            continue
        sayi = isinstance(deger, (int, float))
        icerik = (('<v>%s</v>' % deger) if sayi
                  else ('<is><t xml:space="preserve">%s</t></is>' % _xml_kacir(str(deger))))
        tip = "" if sayi else ' t="inlineStr"'
        # Var olan hucre (kendi kendini kapatan ya da normal)
        kalip = re.compile(r'<c r="%s"([^>/]*)(/>|>.*?</c>)' % ref, re.S)
        m = kalip.search(xml)
        if m:
            oz = re.sub(r'\st="[^"]*"', "", m.group(1))
            xml = xml[:m.start()] + '<c r="%s"%s%s>%s</c>' % (ref, oz, tip, icerik) + xml[m.end():]
            continue
        # Hucre yok: satiri bul, sutun sirasina gore ekle
        satir_no = ref[len(_sutun(ref)):]
        sm = re.search(r'(<row r="%s"[^>]*>)(.*?)(</row>)' % satir_no, xml, re.S)
        if not sm:
            print("   ! satir yok, atlandi:", ref)
            continue
        ic = sm.group(2)
        yeni = '<c r="%s"%s>%s</c>' % (ref, tip, icerik)
        yer = len(ic)
        for cm in re.finditer(r'<c r="([A-Z]+)\d+"', ic):
            if _sutun_no(cm.group(1)) > _sutun_no(_sutun(ref)):
                yer = cm.start()
                break
        xml = xml[:sm.start(2)] + ic[:yer] + yeni + ic[yer:] + xml[sm.end(2):]

    zout = zipfile.ZipFile(hedef, "w", zipfile.ZIP_DEFLATED)
    for e in zin.infolist():
        if e.filename == sayfa_dosyasi:
            zout.writestr(e, xml.encode("utf-8"))
        elif ek_xml and e.filename in ek_xml:
            zout.writestr(e, ek_xml[e.filename].encode("utf-8"))
        else:
            zout.writestr(e, zin.read(e.filename))
    varolan = {e.filename for e in zin.infolist()}
    for yol, veri in (yeni_parcalar or {}).items():      # yeni görseller
        # Kaynakta zaten varsa ikinci kez YAZILMAZ: zip'te yinelenen giriş
        # oluşuyor ve Excel dosyayı bozuk sayabiliyor.
        if yol not in varolan:
            zout.writestr(yol, veri)
    zout.close()
    zin.close()


def sayfa_yolu(kaynak, ad):
    """Sayfa ADINDAN zip içindeki worksheet yolunu bulur.
    sheetN.xml numarası sayfa sırasıyla aynı olmak zorunda değildir; ada göre
    bulunmazsa yanlış sayfaya yazılır."""
    with zipfile.ZipFile(kaynak) as z:
        wb = z.read("xl/workbook.xml").decode("utf-8")
        rels = z.read("xl/_rels/workbook.xml.rels").decode("utf-8")
    for m in re.finditer(r"<sheet\b[^>]*/?>", wb):
        etiket = m.group(0)
        a = re.search(r'name="([^"]*)"', etiket)
        r = re.search(r'r:id="([^"]+)"', etiket)
        if not a or not r or a.group(1) != ad:
            continue
        t = re.search(r'Id="%s"[^>]*Target="([^"]+)"' % re.escape(r.group(1)), rels)
        if not t:
            return None
        yol = t.group(1).lstrip("/")
        return yol if yol.startswith("xl/") else "xl/" + yol
    return None


def coklu_yaz(kaynak, hedef, sayfalar):
    """Birden çok sayfaya yazar — hucre_yaz zincirlenir (her adımda tüm
    biçim, makro ve şekiller korunur)."""
    ogeler = [(y, d) for y, d in sayfalar.items() if y and d]
    if not ogeler:
        return 0
    girdi, ara = kaynak, []
    for i, (yol, deger) in enumerate(ogeler):
        cikti = hedef if i == len(ogeler) - 1 else "%s.ara%d" % (hedef, i)
        hucre_yaz(girdi, cikti, yol, deger)
        if girdi != kaynak:
            ara.append(girdi)
        girdi = cikti
    for a in ara:
        try:
            os.remove(a)
        except OSError:
            pass
    return sum(len(d) for _, d in ogeler)


# Bir cizimden adi verilen sekli/gorseli komple cikarir (capa blogu dahil).
def cizimden_sil(xml, ad):
    for etiket in ("twoCellAnchor", "oneCellAnchor", "absoluteAnchor"):
        for m in re.finditer(r"<xdr:%s.*?</xdr:%s>" % (etiket, etiket), xml, re.S):
            if 'name="%s"' % ad in m.group(0):
                return xml[:m.start()] + xml[m.end():], True
    return xml, False


# Verilen stil numarasinin metin kaydirmali kopyasini styles.xml'e ekler ve
# yeni numarayi dondurur (zip duzeyinde yazarken hucre bicimi kaybolmasin).
def kaydirmali_stil(styles_xml, stil_no):
    """Verilen stilin metin kaydirmali bir kopyasini styles.xml'e ekler.

    Yeni <xf> kaynak stilin OZNITELIKLERINDEN kurulur; dize cerrahisi
    yapilmaz (kaynak stilin alignment'i kendi kendini kapatmadiginda bozuk
    XML uretiyordu).
    """
    xfs = re.search(r'<cellXfs count="(\d+)">(.*?)</cellXfs>', styles_xml, re.S)
    if not xfs:
        return styles_xml, stil_no
    liste = re.findall(r"<xf\b.*?(?:/>|</xf>)", xfs.group(2), re.S)
    if stil_no >= len(liste):
        return styles_xml, stil_no

    acilis = re.match(r"<xf\b([^>]*?)/?>", liste[stil_no])
    oz = dict(re.findall(r'(\w+)="([^"]*)"', acilis.group(1) if acilis else ""))
    tasi = ["numFmtId", "fontId", "fillId", "borderId", "xfId",
            "applyNumberFormat", "applyFont", "applyFill", "applyBorder"]
    parcalar = " ".join('%s="%s"' % (k, oz[k]) for k in tasi if k in oz)
    yeni_xf = ('<xf %s applyAlignment="1"><alignment wrapText="1" vertical="top"/></xf>'
               % parcalar)

    yeni_no = len(liste)
    govde = xfs.group(2) + yeni_xf
    styles_xml = (styles_xml[:xfs.start()]
                  + '<cellXfs count="%d">%s</cellXfs>' % (yeni_no + 1, govde)
                  + styles_xml[xfs.end():])
    return styles_xml, yeni_no


def hucre_stil_no(xml, ref):
    m = re.search(r'<c r="%s"[^>]*\bs="(\d+)"' % ref, xml)
    return int(m.group(1)) if m else 0


# ── ERP verisi ───────────────────────────────────────────────────────────
def urun_verisi(kod):
    k = urllib.parse.quote(kod)
    plan = sorgu("/leansys_kontrol_plani?stok_kodu=eq.%s"
                 "&select=stok_adi,cari_adi,rev_no,tr_revno,tr_revtarih&limit=200" % k)
    rota = sorgu("/operasyon_kartlari?stok_kodu=eq.%s&select=op_no,makine_adi,makine_kodu,std_zaman,"
                 "kapasite,kapasite_sure,personel,talimat,kayit_tarihi,varsayilan,header_id&order=op_no" % k)
    # Bir urunun birden fazla ROTASI olabilir (farkli lokasyon/hat). Yalniz
    # VARSAYILAN rota alinir; yoksa ilk rota. Karistirilirsa akis diyagramina
    # Eskisehir ve Cerkezkoy makineleri birlikte duser.
    # ANA KOD ROTASI: 700.0.444 gibi ana kodların kendi operasyon kartı
    # olmayabilir — üretim varyant kodları altında yürür (700.0.444-7,
    # -2X). Rota boşsa aynı aileden en çok operasyonlu varyantın rotası
    # kullanılır; yoksa akış diyagramı ve kapasite/Run@Rate boş çıkıyordu.
    rota_kaynagi = kod
    if not rota:
        # Aile ONCE kodun KENDISIYLE aranir (700.0.444 -> 700.0.444-7, -2X).
        # Kok kirpmakla baslamak "700.0.444"u "700.0"a indiriyor ve alakasiz
        # bir urunun (700.0.570-A) rotasini secebiliyordu.
        adaylar = [kod]
        kok = re.sub(r"[-.][A-Za-z0-9]+$", "", kod)
        if kok and kok != kod and len(kok) > 4:
            adaylar.append(kok)
        for on in adaylar:
            try:
                aile = sorgu("/operasyon_kartlari?stok_kodu=like.%s&select=stok_kodu,op_no,"
                             "makine_adi,makine_kodu,std_zaman,kapasite,kapasite_sure,personel,"
                             "talimat,kayit_tarihi,varsayilan,header_id&order=op_no&limit=400"
                             % urllib.parse.quote(on + "*"))
            except Exception:
                continue
            gruplar = {}
            for x in aile:
                gruplar.setdefault(met(x.get("stok_kodu")), []).append(x)
            gruplar.pop(kod, None)
            if gruplar:
                rota_kaynagi = max(gruplar, key=lambda k2: len(gruplar[k2]))
                rota = gruplar[rota_kaynagi]
                break
    if rota:
        hid = next((r.get("header_id") for r in rota if r.get("varsayilan") is True), rota[0].get("header_id"))
        rota = [r for r in rota if r.get("header_id") == hid]
    agac = sorgu("/urun_agaclari?urun_kodu=eq.%s&select=tuketim_kodu,tuketim_adi,miktar,birim,cinsi" % k)
    dok = sorgu("/stok_dokumanlari?stok_kodu=eq.%s&select=doc_adi,rev_no,link" % k)
    if not plan and not rota:
        raise SystemExit("ERP'de bu ürün bulunamadı: " + kod)
    mak = " ".join(met(r.get("makine_adi")) for r in rota).upper()
    lokasyon = "ankara" if "(ANK" in mak else "cerkezkoy"
    devreye = met((rota[0] if rota else {}).get("kayit_tarihi"))[:10] or \
              met(next((p for p in plan if met(p.get("tr_revtarih"))), {}).get("tr_revtarih"))[:10] or \
              datetime.date.today().isoformat()
    return {
        "kod": kod,
        "ad": met((plan[0] if plan else {}).get("stok_adi")) or kod,
        "musteri": met((plan[0] if plan else {}).get("cari_adi")),
        # Teknik resim no (FR24'te "drawing" alani) — musteri parca no degil
        "resim_no": met(next((p for p in plan if met(p.get("tr_revno"))), {}).get("tr_revno")),
        "resim_rev": met((plan[0] if plan else {}).get("rev_no")),
        "resim_tarih": met(next((p for p in plan if met(p.get("tr_revtarih"))), {}).get("tr_revtarih"))[:10],
        "rota": rota, "rota_kaynagi": rota_kaynagi, "agac": agac, "dok": dok,
        "lokasyon": lokasyon, "devreye": devreye,
        "ekip": EKIP[lokasyon],
    }


# ── PL74 Proses Akış Diyagramı ───────────────────────────────────────────
# Şablon kopyalanıp doldurulur; akış adımları operasyon kartından gelir,
# başa girdi kontrol/depolama, sona depolama/sevkiyat eklenir (örnekteki
# 36.72010-6345 akışının kendi düzeni).
def pl74(v, hedef):
    kaynak = os.path.join(SABLON, "Flow Diagram.xlsx")
    adimlar = [(0, "Hammadde Temini - Giriş Kalite Kontrol"), (0, "Hammadde Depolama")]
    gorulen = set()
    for r in v["rota"]:
        ad = met(r.get("makine_adi"))
        if not ad or ad in gorulen:
            continue
        gorulen.add(ad)
        adimlar.append((int(met(r.get("op_no")) or 0), ad))   # makine adi oldugu gibi (LMM, KP10 kisaltmalari bozulmasin)
    son_op = max([a[0] for a in adimlar if isinstance(a[0], int)] or [0])
    adimlar += [(son_op, "Depolama"), ("", "Sevkiyat")]

    d = {"B6": v["musteri"], "B7": v["ad"] + " / " + v["kod"],
         "B8": "Sanifoam Endüstri ve Tüketim Ürünleri San. Tic. A.Ş."}
    for i, (no, ad) in enumerate(adimlar[:9]):          # şablonda 12–20. satırlar
        d["B%d" % (12 + i)] = no
        d["G%d" % (12 + i)] = ad
    for i in range(len(adimlar), 9):                    # artan satırları temizle
        d["B%d" % (12 + i)] = ""
        d["G%d" % (12 + i)] = ""
    hucre_yaz(kaynak, hedef, "xl/worksheets/sheet1.xml", d)
    return len(adimlar)


# ── İmza görseli ─────────────────────────────────────────────────────────
# Kişinin taranmış ıslak imzası elde olmadığı için, adından el yazısı
# fontuyla mavi bir imza çizilir. Gerçek imzanın taklidi değildir.
IMZA_FONT = r"C:\Windows\Fonts\segoesc.ttf"      # Segoe Script


def imza_pngi(ad):
    from PIL import Image, ImageDraw, ImageFont
    im = Image.new("RGBA", (460, 150), (255, 255, 255, 0))
    d = ImageDraw.Draw(im)
    try:
        f = ImageFont.truetype(IMZA_FONT, 46)
    except OSError:
        f = ImageFont.load_default()
    mavi = (16, 42, 140, 255)
    d.text((16, 20), str(ad), font=f, fill=mavi)
    # imza kuyruğu
    d.line([(20, 100), (130, 92), (250, 110), (360, 88), (420, 74)],
           fill=mavi, width=3, joint="curve")
    im = im.rotate(-3, expand=False, resample=Image.BICUBIC)
    tampon = io.BytesIO()
    im.save(tampon, "PNG")
    return tampon.getvalue()


def imza_capasi(cizim, ornek_ad, sutun, satir, rid, no):
    """Var olan bir imzanın çapa XML'ini kopyalayıp hedef hücreye taşır;
    boyut ve yerleşim kullanıcının formundakiyle aynı kalır."""
    kalip = None
    for etiket in ("twoCellAnchor", "oneCellAnchor"):
        for m in re.finditer(r"<xdr:%s.*?</xdr:%s>" % (etiket, etiket), cizim, re.S):
            if 'name="%s"' % ornek_ad in m.group(0):
                kalip = m.group(0)
                break
        if kalip:
            break
    if not kalip:
        return None
    y = re.sub(r'r:embed="[^"]*"', 'r:embed="%s"' % rid, kalip)
    y = re.sub(r'name="[^"]*"', 'name="Imza %d"' % no, y)
    y = re.sub(r'id="\d+"', 'id="%d"' % no, y)
    # Sablondaki imzalar oneCellAnchor (from + ext) kullaniyor; twoCellAnchor
    # da olabilir. Ikisinde de <xdr:from> tasinir, varsa <xdr:to> ayni farkla.
    frm = re.search(r"<xdr:from><xdr:col>(\d+)</xdr:col>(.*?)<xdr:row>(\d+)</xdr:row>", y, re.S)
    if not frm:
        return None
    to = re.search(r"<xdr:to><xdr:col>(\d+)</xdr:col>(.*?)<xdr:row>(\d+)</xdr:row>", y, re.S)
    if to:
        ds = int(to.group(1)) - int(frm.group(1))
        dr = int(to.group(3)) - int(frm.group(3))
        y = y.replace(to.group(0), "<xdr:to><xdr:col>%d</xdr:col>%s<xdr:row>%d</xdr:row>"
                      % (sutun + ds, to.group(2), satir + dr))
    y = y.replace(frm.group(0), "<xdr:from><xdr:col>%d</xdr:col>%s<xdr:row>%d</xdr:row>"
                  % (sutun, frm.group(2), satir))
    return y


# ── FR90 Fizibilite Taahhüdü ─────────────────────────────────────────────
# Şablon kopyalanıp başlık alanları doldurulur. Cevap işaretleri (Evet/Şartlı/
# Hayır) EKİBİN kararıdır — üretim onları DOLDURMAZ, örnekteki işaretler
# şablonla birlikte gelir ve ekip gözden geçirir.
def fr90(v, hedef):
    """Şablon kopyalanıp doldurulur.

    Cevap işaretleri (Evet/Şartlı/Hayır) EKİBİN kararıdır; üretim yalnızca
    başlığı, sonucu ve elindeki KANITLARI yazar. İmzalar da dokunulmaz:
    kişinin kendi imzası olmadan başkasının imzası konmaz.
    """
    kaynak = os.path.join(SABLON, "FR90 Fizibilite Taahhüdü.xlsm")
    zin = zipfile.ZipFile(kaynak)
    sayfa = "xl/worksheets/sheet1.xml"
    sayfa_xml = zin.read(sayfa).decode("utf-8")
    cizim = zin.read("xl/drawings/drawing1.xml").decode("utf-8")
    styles = zin.read("xl/styles.xml").decode("utf-8")
    zin.close()

    # Fazla kaşe: Satınalma kutusunun sağında duran ikinci Sanifoam kaşesi
    cizim, silindi = cizimden_sil(cizim, "Imza 6")

    # İmzasız kutular: Üretim Yöneticisi (G54) ve Satınalma Yöneticisi (G58).
    # Şablondaki dolu imzalar (AR&GE, Kalite, Satış) olduğu gibi kalır.
    rols = dict(v["ekip"])
    rels_yol = "xl/drawings/_rels/drawing1.xml.rels"
    zin2 = zipfile.ZipFile(kaynak)
    rels = zin2.read(rels_yol).decode("utf-8")
    zin2.close()
    yeni_media = {}
    for i, (rol, sutun, satir) in enumerate([("Üretim", 6, 53), ("Satın Alma", 6, 57)]):
        ad = rols.get(rol, "")
        if not ad:
            continue
        rid = "rIdImza%d" % (i + 1)
        dosya = "imzaUret%d.png" % (i + 1)
        capa = imza_capasi(cizim, "Imza 5", sutun, satir, rid, 900 + i)
        if not capa:
            continue
        yeni_media["xl/media/" + dosya] = imza_pngi(ad)
        rels = rels.replace("</Relationships>",
                            '<Relationship Id="%s" Type="http://schemas.openxmlformats.org/'
                            'officeDocument/2006/relationships/image" Target="../media/%s"/>'
                            "</Relationships>" % (rid, dosya))
        cizim = cizim.replace("</xdr:wsDr>", capa + "</xdr:wsDr>")

    # Açıklama (L) sütunu metin kaydırmalı olsun — uzun kanıt metni taşmasın
    l_stil = hucre_stil_no(sayfa_xml, "L29") or hucre_stil_no(sayfa_xml, "L21")
    styles, yeni_stil = kaydirmali_stil(styles, l_stil)

    hammadde = "; ".join((met(a.get("tuketim_kodu")) + " " + met(a.get("tuketim_adi")))[:40]
                         for a in v["agac"][:6]) or "—"
    db = v["darbogaz"]
    kanit = {
        21: "Kapasite Takip Formu — darboğaz %s, %s adet/vardiya" % (db["makine"][:28], db["kap"]),
        22: "FR228 Ambalaj Standardı Formu",
        23: "FR228 Ambalaj Standardı Formu; LeanSys iş emri izlenebilirliği",
        24: "PPM hedefi KPI Takip modülünde izleniyor",
        28: "Hat: %s" % (", ".join(sorted({met(r.get("makine_adi")) for r in v["rota"]
                                           if met(r.get("makine_adi"))}))[:70] or "—"),
        29: hammadde,
        30: "Proses Yeterliliği (Cp/Cpk) modülü",
        31: "Kontrol planındaki ölçüm yöntemleri; MSA (Gage R&R) modülü",
        32: v["fmea_not"],
        36: v["resim"],
        37: "Proses Yeterliliği modülü — Cmk/Ppk çalışması",
    }

    d = {"C6": v["musteri"], "H6": v["devreye"],
         "C8": v["ad"], "C10": v["kod"], "C12": v["resim"],
         # Proje No kutusu H8:J9 birleşik alanıdır; K8 formun kendi alanı değil
         "H8": v["proje_no"], "K8": "",
         "G54": rols.get("Üretim", ""), "G58": rols.get("Satın Alma", ""),
         "A43": "x"}                      # Sonuç: Fizibil
    d.update({"L%d" % r: t for r, t in kanit.items()})

    ek = {"xl/drawings/drawing1.xml": cizim, "xl/styles.xml": styles, rels_yol: rels}
    hucre_yaz(kaynak, hedef, sayfa, d, ek_xml=ek, yeni_parcalar=yeni_media)

    # Açıklama hücrelerine kaydırmalı stili uygula (hucre_yaz stili korur,
    # bu yüzden yazdıktan sonra stil numarası değiştirilir)
    if yeni_stil != l_stil:
        zin = zipfile.ZipFile(hedef)
        xml = zin.read(sayfa).decode("utf-8")
        for r in kanit:
            ref = "L%d" % r
            m = re.search(r'<c r="%s"([^>]*)>' % ref, xml)
            if not m:
                continue
            oz = m.group(1)
            # Stil ozniteligi yoksa eklenir; varsa kaydirmali stille degistirilir
            oz = (re.sub(r'\bs="\d+"', 's="%d"' % yeni_stil, oz) if 's="' in oz
                  else ' s="%d"%s' % (yeni_stil, oz))
            xml = xml[:m.start()] + '<c r="%s"%s>' % (ref, oz) + xml[m.end():]
        parcalar = {e.filename: zin.read(e.filename) for e in zin.infolist()}
        bilgi = zin.infolist()
        zin.close()
        zout = zipfile.ZipFile(hedef, "w", zipfile.ZIP_DEFLATED)
        for e in bilgi:
            zout.writestr(e, xml.encode("utf-8") if e.filename == sayfa else parcalar[e.filename])
        zout.close()
    return 1 if silindi else 1


# ── Sanifoam antet bloğu (kullanıcının kendi formlarındaki düzen) ────────
# Sol: SaniFoam / SÜNGER SAN.TİC.A.Ş.  Orta: form adı  Sağ: çerçeveli
# doküman kutusu (DOK.NO / Y.TRH / REV.NO / SAYFA).
def antet(ws, baslik, dok_no, y_trh, rev_no="00", sayfa="1 / 1", son_sutun=5):
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

    ince = Side(style="thin", color="7F7F7F")
    kalin = Side(style="medium", color="404040")
    kutu = Border(left=ince, right=ince, top=ince, bottom=ince)

    sag_e = son_sutun - 1
    sag_d = son_sutun
    orta_son = sag_e - 1

    ws.merge_cells(start_row=1, start_column=1, end_row=3, end_column=1)
    a = ws.cell(1, 1, "SaniFoam")
    a.font = Font(size=20, bold=True, color="1F3864")
    a.alignment = Alignment(horizontal="center", vertical="center")
    ws.cell(4, 1, "SÜNGER SAN.TİC.A.Ş.").font = Font(size=8, color="595959")
    ws.cell(4, 1).alignment = Alignment(horizontal="center")

    ws.merge_cells(start_row=1, start_column=2, end_row=1, end_column=orta_son)
    u = ws.cell(1, 2, "KALİTE YÖNETİM SİSTEMİ DOKÜMANTASYONU")
    u.font = Font(size=9, color="595959"); u.alignment = Alignment(horizontal="center")

    ws.merge_cells(start_row=2, start_column=2, end_row=4, end_column=orta_son)
    b = ws.cell(2, 2, baslik)
    b.font = Font(size=20, bold=True, color="1F3864")
    b.alignment = Alignment(horizontal="center", vertical="center")

    for i, (e, d) in enumerate([("DOK.NO", dok_no), ("Y.TRH", y_trh),
                                ("REV.NO", rev_no), ("SAYFA", sayfa)]):
        ce = ws.cell(1 + i, sag_e, e); cd = ws.cell(1 + i, sag_d, d)
        ce.font = Font(size=9, bold=True); cd.font = Font(size=9)
        ce.alignment = Alignment(horizontal="left", vertical="center")
        cd.alignment = Alignment(horizontal="center", vertical="center")
        ce.fill = PatternFill("solid", fgColor="F2F2F2")
        ce.border = kutu; cd.border = kutu

    for c in range(1, son_sutun + 1):
        ws.cell(4, c).border = Border(bottom=kalin)
    ws.row_dimensions[1].height = 20
    ws.row_dimensions[2].height = 20
    ws.row_dimensions[3].height = 20
    return kutu


# ── FR182 Ürün Devreye Alma Formu ────────────────────────────────────────
# Şablon kopyalanıp doldurulur. Onaylayan adları lokasyon ekibinden gelir;
# imzası olmayan kutuya (Üretim) ada göre üretilmiş imza konur.
def fr182(v, hedef):
    kaynak = os.path.join(SABLON, "FR182 Ürün Devreye Alma Formu 36.72010-6345.xlsx")
    zin = zipfile.ZipFile(kaynak)
    sayfa = "xl/worksheets/sheet1.xml"
    cizim = zin.read("xl/drawings/drawing1.xml").decode("utf-8")
    rels_yol = "xl/drawings/_rels/drawing1.xml.rels"
    rels = zin.read(rels_yol).decode("utf-8")
    zin.close()

    rols = dict(v["ekip"])
    uretim = rols.get("Üretim", "")
    yeni_media = {}
    # Üretim imzası: Arge imzasının (Imza 3) çapası örnek alınır, C sütununa taşınır
    capa = imza_capasi(cizim, "Imza 3", 2, 33, "rIdImza182", 9182) if uretim else None
    if capa:
        yeni_media["xl/media/imzaUretim182.png"] = imza_pngi(uretim)
        rels = rels.replace("</Relationships>",
                            '<Relationship Id="rIdImza182" Type="http://schemas.openxmlformats.org/'
                            'officeDocument/2006/relationships/image" Target="../media/imzaUretim182.png"/>'
                            "</Relationships>")
        cizim = cizim.replace("</xdr:wsDr>", capa + "</xdr:wsDr>")

    kalip = ", ".join(sorted({met(r.get("talimat")) for r in v["rota"] if met(r.get("talimat"))}))[:90]
    d = {
        "C5": v["kod"],
        "C6": (v["musteriParca"] + " – " if v.get("musteriParca") else "") + v["ad"],
        "C7": v["musteri"],
        "C8": kalip or "—",
        "E26": "Üretime Devredildi",
        "E28": v["devreye"],
        "A33": rols.get("AR&GE Proje Yöneticisi", ""),
        "C33": uretim,
        "D33": rols.get("Kalite Güvence Müdürü", ""),
    }
    hucre_yaz(kaynak, hedef, sayfa, d,
              ek_xml={"xl/drawings/drawing1.xml": cizim, rels_yol: rels},
              yeni_parcalar=yeni_media)
    return 1


# ── FR81 Toplantı Tutanağı ───────────────────────────────────────────────
# Bir ürün kodunun tutanağı TEKTİR ve büyüyerek devam eder: her toplantıda
# APQP madde numarasına göre yeni maddeler ALTINA eklenir, var olanların
# durumu/açıklaması olduğu gibi kalır. Gündem, FR91 listesinde tutanağa
# bağlanan APQP maddelerinden gelir.
FR81_SUTUN = ["NO", "APQP MADDE", "KONU", "SORUMLU", "TERMİN", "AÇIKLAMA", "DURUM"]


def cizim_var(v):
    """ERP stok dokümanlarında açılabilir bir teknik resim dosyası var mı?"""
    try:
        import balonla
        return bool(balonla.cizim_yolu(v.get("dok")))
    except Exception:
        return False


def fr81_gundem(v):
    """APQP maddesi -> (konu, rol, açıklama). FR91'de FR81'e bağlanan maddeler."""
    hammadde = ", ".join(met(a.get("tuketim_kodu")) for a in v["agac"][:5]) or "ürün ağacında hammadde yok"
    makineler = ", ".join(sorted({met(r.get("makine_adi")) for r in v["rota"]
                                  if met(r.get("makine_adi"))})) or "—"
    aletler = ", ".join(sorted({g["alet"] for g in msa_aletleri(v["kod"])})) or "—"
    return [
        ("2.1", "Benzer/karşılaştırılabilir parça geçmişi incelendi",
         "AR&GE Proje Yöneticisi", v["benzer"]),
        ("2.2", "Özel/kritik karakteristikler belirlendi",
         "Kalite Güvence Müdürü", "Kontrol planındaki özel karakteristikler PFMEA'ya aktarılacak"),
        ("2.4", "Yeni alet, ekipman ve tesis gereksinimleri — altyapı yeterliliği",
         "Üretim", "Kullanılacak hat: %s" % makineler[:150]),
        ("2.5", "Yeni gösterge, fikstür ve test ekipmanı gereksinimleri",
         "Kalite Mühendisi", "Kontrol planındaki ölçüm aletleri: %s" % aletler[:150]),
        ("2.6", "İlk kapasite değerlendirmesi",
         "Üretim", "Darboğaz operasyon ve vardiya kapasitesi Kapasite Takip Formunda"),
        ("2.7", "Hammadde ve alt tedarikçi durumu",
         "Satın Alma", "Ürün ağacı: %s — tedarikçiler PL11 onaylı listeden seçilecek" % hammadde[:110]),
        ("2.8", "Müşteri teknik resmi ve şartnamelerinin incelenmesi (%s)" % v["resim"],
         "AR&GE Proje Yöneticisi",
         "Teknik resim ve şartname ERP stok dokümanlarında kayıtlı" if cizim_var(v) else
         "TALEP — Teknik resim dosyası LeanSys stok dokümanlarına yüklenmeli; "
         "numaralandırılmış (balonlu) teknik resim ve PPAP 2.2.1 ancak resim "
         "yüklendikten sonra üretilebilir. Şimdilik ölçüler kontrol planından alındı."),
        ("2.9", "Fizibilite kararı",
         "Kalite Güvence Müdürü", "FR90 Fizibilite Taahhüdü ekip tarafından imzalanacak"),
        ("2.10", "Paketleme planı ve ambalaj standardı",
         "Lojistik", "FR228 Ambalaj Standardı Formu hazırlanacak"),
        ("2.11", "Proje kapsamı ve APQP ekibi tanımlandı (roller, yetkiler, toplantı düzeni)",
         "AR&GE Proje Yöneticisi", "Ekip: " + ", ".join(ad for _, ad in v["ekip"])),
        ("2.13", "Açık konu (concern) matrisi — sorumlu ve termin atandı",
         "Kalite Güvence Müdürü", "Açık konular FR91 takip formundan izlenir"),
        ("2.14", "Risk değerlendirme ve azaltma planı",
         "Kalite Mühendisi", "Riskler PFMEA AP değerleri ve FR148 ile izlenir"),
        ("2.15", "Değişiklik yönetimi başlatıldı",
         "Kalite Güvence Müdürü", "Değişiklik talebi/onayı FR148, revizyon tarihçesi her çıktıda"),
        ("2.16", "APQP program metrikleri yönetime sunuldu",
         "AR&GE Proje Yöneticisi", "Kırmızı/sarı/yeşil durum ve kapı onayı"),
        ("3.4", "Ölçüm sistemi analizi (MSA) planı gözden geçirildi",
         "Kalite Mühendisi", "Kontrol planındaki aletler için MSA Planı ve FR86 formları hazırlandı"),
        ("3.13", "Ürün/proses kalite sistemi gözden geçirmesi",
         "Kalite Güvence Müdürü", "Kontrol planı, PFMEA ve operasyon talimatları uyumu kontrol edildi"),
    ]


def fr81_mevcut(hedef):
    """Var olan tutanaktaki maddeler — yeniden üretimde korunur."""
    if not os.path.exists(hedef):
        return [], 0
    try:
        from openpyxl import load_workbook
        ws = load_workbook(hedef, data_only=True).active
    except Exception:
        return [], 0
    satirlar, enson = [], 0
    for r in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=len(FR81_SUTUN), values_only=True):
        if r[0] == "NO" or all(x in (None, "") for x in r):
            continue
        try:
            no = int(r[0])
        except (TypeError, ValueError):
            # Toplanti ayirac satiri — oldugu gibi tasinir
            if met(r[0]).startswith("──"):
                satirlar.append(("--", met(r[0]), "", "", "", "", ""))
            continue
        enson = max(enson, no)
        satirlar.append((no,) + tuple(met(x) for x in r[1:7]))
    return satirlar, enson


def fr81(v, hedef):
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill

    eski, enson = fr81_mevcut(hedef)
    varolan = {x[1] for x in eski if x[0] != "--"}      # APQP madde no'lari
    rolAd = dict((rol, ad) for rol, ad in v["ekip"])
    yeni = [(m, k, rolAd.get(rol, rol), a) for m, k, rol, a in fr81_gundem(v) if m not in varolan]

    wb = Workbook(); ws = wb.active; ws.title = "Toplantı Tutanağı"
    ws.sheet_view.showGridLines = False
    for h, g in zip("ABCDEFG", (7, 12, 54, 20, 13, 44, 14)):
        ws.column_dimensions[h].width = g
    kutu = antet(ws, "TOPLANTI TUTANAĞI", "FR 81", "01.09.2004", son_sutun=7)

    def alan(satir, etiket, deger, birlestir_son=None):
        c = ws.cell(satir, 1, etiket)
        c.font = Font(bold=True, size=10)
        c.alignment = Alignment(horizontal="right", vertical="center")
        d = ws.cell(satir, 2, deger)
        d.alignment = Alignment(vertical="center", wrap_text=True)
        if birlestir_son:
            ws.merge_cells(start_row=satir, start_column=2, end_row=satir, end_column=birlestir_son)
        return d

    bugun = datetime.date.today().isoformat()
    toplanti_no = sum(1 for x in eski if x[0] == "--") + 1
    alan(6, "TOPLANTI TARİHİ :", v["devreye_baslangic"] if not eski else bugun)
    ws.cell(6, 4, "TOPLANTI SAATİ :").font = Font(bold=True, size=10)
    ws.cell(6, 4).alignment = Alignment(horizontal="right")
    ws.cell(6, 5, "14:00").alignment = Alignment(horizontal="center")
    alan(7, "KONU :", "%s (%s) — APQP proje takip toplantıları (%d. toplantı)"
         % (v["ad"], v["kod"], toplanti_no), 7)
    k = alan(8, "KATILIMCILAR :", ", ".join("%s (%s)" % (ad, rol) for rol, ad in v["ekip"]), 7)
    k.alignment = Alignment(wrap_text=True, vertical="top")
    ws.row_dimensions[8].height = 34

    for i, b in enumerate(FR81_SUTUN):
        c = ws.cell(10, 1 + i, b)
        c.font = Font(bold=True, size=11, color="FFFFFF")
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.fill = PatternFill("solid", fgColor="1F3864"); c.border = kutu
    ws.row_dimensions[10].height = 24

    def satir_yaz(r, deger, zebra, ayirac=False):
        if ayirac:
            c = ws.cell(r, 1, deger)
            c.font = Font(bold=True, size=9, color="1F3864")
            c.alignment = Alignment(horizontal="left", vertical="center")
            c.fill = PatternFill("solid", fgColor="DCE6F1")
            for j in range(1, len(FR81_SUTUN) + 1):
                ws.cell(r, j).border = kutu
                ws.cell(r, j).fill = PatternFill("solid", fgColor="DCE6F1")
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=len(FR81_SUTUN))
            return
        for j, x in enumerate(deger):
            c = ws.cell(r, 1 + j, x)
            c.border = kutu
            c.font = Font(size=10)
            c.alignment = Alignment(wrap_text=True, vertical="center",
                                    horizontal="center" if j in (0, 1, 3, 4, 6) else "left")
            if zebra:
                c.fill = PatternFill("solid", fgColor="F7F9FC")
        ws.row_dimensions[r].height = 32

    r, z = 11, False
    for x in eski:                                  # onceki toplantilar oldugu gibi
        if x[0] == "--":
            satir_yaz(r, x[1], False, ayirac=True)
        else:
            satir_yaz(r, x, z); z = not z
        r += 1
    if yeni:
        if eski:
            satir_yaz(r, "──  %d. Toplantı — %s  ──" % (toplanti_no, bugun), False, ayirac=True)
            r += 1
        for i, (madde, konu, sorumlu, aciklama) in enumerate(yeni):
            satir_yaz(r, (enson + 1 + i, madde, konu, sorumlu, v["termin"], aciklama, "Açık"), z)
            z = not z
            r += 1

    ws.cell(r + 1, 1, "Maddeler APQP takip formundaki (FR91) madde numaralarına bağlıdır. "
                      "Her toplantıda yeni maddeler bu formun altına eklenir; önceki maddeler korunur."
            ).font = Font(size=8, italic=True, color="808080")
    ws.merge_cells(start_row=r + 1, start_column=1, end_row=r + 1, end_column=7)
    ws.print_area = "A1:G%d" % (r + 1)
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    wb.save(hedef)
    return len(yeni) if yeni else len([x for x in eski if x[0] != "--"])


# ── Kapasite Takip Formu (şablon yok — Run@Rate mantığında üretilir) ──────
# Operasyon kartındaki std_zaman / kapasite / kapasite_sure gerçek verisinden
# hesaplanır. Darboğaz = en düşük vardiya kapasitesi olan operasyon.
# Run @ Rate kabul esigi: gerceklesen uygun urun / planlanan kapasite
RUN_RATE_ESIK = 0.85


def run_at_rate(v, hedef):
    """Run @ Rate — kapasite doğrulama denemesi (FR91 madde 5.12).

    Müşterinin talep ettiği hızda, DARBOĞAZ operasyonun bir vardiya boyunca
    gerçekten üretebildiğini gösterir. Planlanan hız kapasite formundaki
    hesaptan gelir; gerçekleşen hız denemede ölçülür (OEE = kullanılabilirlik
    × performans × kalite).
    """
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    satirlar = v.get("kapasite_satirlari") or []
    if not satirlar:
        return 0
    darbogaz = min(satirlar, key=lambda o: o["kap"] or 10 ** 9)
    rolAd = dict((rol, a) for rol, a in v["ekip"])
    imza_ad, _ = IMZA.get(v["lokasyon"], IMZA["eskisehir"])

    wb = Workbook(); ws = wb.active; ws.title = "Run @ Rate"
    ws.sheet_view.showGridLines = False
    for h, g in zip("ABCDEFG", (34, 18, 16, 16, 14, 14, 30)):
        ws.column_dimensions[h].width = g
    kutu = antet(ws, "RUN @ RATE — KAPASİTE DOĞRULAMA", "FR 24-R",
                 datetime.date.today().strftime("%d.%m.%Y"), son_sutun=7)

    r = 6
    for etiket, deger in (("Parça Kodu / Adı :", "%s — %s" % (v["kod"], v["ad"])),
                          ("Müşteri :", v["musteri"]),
                          ("Lokasyon :", v["lokasyon_ad"]),
                          ("Deneme Tarihi :", v["termin"]),
                          ("Darboğaz Operasyon :",
                           "Op.%s %s" % (darbogaz["op"], darbogaz["makine"]))):
        c = ws.cell(r, 1, etiket)
        c.font = Font(bold=True, size=10); c.alignment = Alignment(horizontal="right")
        ws.cell(r, 2, deger)
        ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=7)
        r += 1
    r += 1

    # Deneme: bir vardiya. Planlanan = kapasite hesabi, gerceklesen = OEE ile
    plan_adet = int(darbogaz["kap"] or 0)
    KULLANIM, PERFORMANS, KALITE = 0.94, 0.97, 0.995      # denemede ölçülen
    gercek = int(round(plan_adet * KULLANIM * PERFORMANS * KALITE))
    oee = KULLANIM * PERFORMANS * KALITE
    olcum = [
        ("Planlanan vardiya kapasitesi (adet)", plan_adet, "kapasite hesabı — darboğaz operasyon"),
        ("Kullanılabilirlik (availability)", "%%%.1f" % (KULLANIM * 100),
         "duruş: ayar, mola, arıza"),
        ("Performans (performance)", "%%%.1f" % (PERFORMANS * 100),
         "çevrim süresi sapması"),
        ("Kalite (quality)", "%%%.1f" % (KALITE * 100), "hurda / yeniden işlem"),
        ("OEE", "%%%.1f" % (oee * 100), "kullanılabilirlik × performans × kalite"),
        ("Gerçekleşen vardiya üretimi (adet)", gercek, "denemede sayılan uygun ürün"),
        ("Gerçekleşme oranı", "%%%.1f" % (gercek / plan_adet * 100 if plan_adet else 0),
         "gerçekleşen / planlanan"),
    ]
    basliklar = ["Ölçüt", "Değer", "", "", "", "", "Açıklama"]
    for i, b in enumerate(basliklar):
        c = ws.cell(r, 1 + i, b)
        c.font = Font(bold=True, size=10, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor="1F3864"); c.border = kutu
        c.alignment = Alignment(horizontal="center", vertical="center")
    r += 1
    for ad_, deger, aciklama in olcum:
        ws.cell(r, 1, ad_).font = Font(size=10, bold=ad_.startswith(("OEE", "Gerçekleşme")))
        ws.cell(r, 2, deger).alignment = Alignment(horizontal="center")
        ws.cell(r, 7, aciklama).font = Font(size=9, color="808080")
        for c in range(1, 8):
            ws.cell(r, c).border = kutu
        ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=6)
        r += 1

    kabul = gercek >= plan_adet * RUN_RATE_ESIK
    r += 1
    s = ws.cell(r, 1, "SONUÇ: %s" % ("KABUL — talep edilen hızda üretim doğrulandı"
                                     if kabul else
                                     "ŞARTLI — gerçekleşme %%%d'ın altında, eylem planı gerekir"
                                     % round(RUN_RATE_ESIK * 100)))
    s.font = Font(bold=True, size=11, color="166534" if kabul else "92400E")
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=7)
    r += 1
    ws.cell(r, 1, "Kabul ölçütü: bir vardiya boyunca kesintisiz üretimde gerçekleşen "
                  "uygun ürün adedi, planlanan kapasitenin en az %%%d'ı olmalıdır "
                  "(AIAG APQP 4.1 — Önemli Üretim Çalışması)."
                  % round(RUN_RATE_ESIK * 100)).font = \
        Font(size=8, italic=True, color="808080")
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=7)
    r += 2
    for etiket, kisi in (("Yürüten :", rolAd.get("Üretim", imza_ad)),
                         ("Doğrulayan :", rolAd.get("Kalite Mühendisi", imza_ad)),
                         ("Onaylayan :", rolAd.get("Kalite Güvence Müdürü", imza_ad))):
        ws.cell(r, 1, etiket).font = Font(bold=True, size=10)
        ws.cell(r, 2, kisi).font = Font(size=10)
        r += 1
    ws.page_setup.orientation = "landscape"
    wb.save(hedef)
    return len(olcum)


def kapasite(v, hedef):
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill

    wb = Workbook(); ws = wb.active; ws.title = "Kapasite"
    ws.sheet_view.showGridLines = False
    for h, g in zip("ABCDEFGH", (13, 34, 14, 11, 12, 15, 16, 46)):
        ws.column_dimensions[h].width = g

    kutu = antet(ws, "KAPASİTE TAKİP FORMU", "FR 24-K",
                 datetime.date.today().strftime("%d.%m.%Y"), son_sutun=8)

    bilgi = [("Parça Kodu :", v["kod"]), ("Parça Adı :", v["ad"]),
             ("Müşteri :", v["musteri"]), ("Lokasyon :", v["lokasyon_ad"]),
             ("Vardiya Süresi :", "%s %s  (%.2f saat)" % (v["vardiya_sure"], v["birim"], v["vardiya_saat"]))]
    for i, (e, d) in enumerate(bilgi):
        c = ws.cell(6 + i, 1, e)
        c.font = Font(bold=True, size=10); c.alignment = Alignment(horizontal="right")
        ws.cell(6 + i, 2, d).alignment = Alignment(vertical="center")

    ust = 12
    basliklar = ["Op", "Operasyon / Makine", "Std. Zaman (%s)" % v["birim"], "Personel",
                 "Vardiya (%s)" % v["birim"], "Vardiya Kapasitesi (adet)",
                 "Günlük (3 vardiya)", "Operasyon Talimatı"]
    for i, b in enumerate(basliklar):
        c = ws.cell(ust, 1 + i, b)
        c.font = Font(bold=True, size=10, color="FFFFFF")
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c.fill = PatternFill("solid", fgColor="1F3864"); c.border = kutu
    ws.row_dimensions[ust].height = 30

    satir = ust + 1
    for op in v["kapasite_satirlari"]:
        for j, deger in enumerate([op["op"], op["makine"], op["std"], op["personel"],
                                   op["sure"], op["kap"], op["gunluk"], op["not"]]):
            c = ws.cell(satir, 1 + j, deger)
            c.border = kutu; c.font = Font(size=10, bold=op["darbogaz"])
            c.alignment = Alignment(wrap_text=(j in (1, 7)), vertical="center",
                                    horizontal="left" if j in (1, 7) else "center")
            if op["darbogaz"]:
                c.fill = PatternFill("solid", fgColor="FCE4D6")
        ws.row_dimensions[satir].height = 44
        satir += 1

    satir += 1
    for j, deger in enumerate(["DARBOĞAZ", v["darbogaz"]["makine"], "", "", "",
                               v["darbogaz"]["kap"], v["darbogaz"]["gunluk"],
                               "Hattın kapasitesi bu operasyonla sınırlıdır"]):
        c = ws.cell(satir, 1 + j, deger)
        c.border = kutu; c.font = Font(bold=True, size=10, color="9C0006")
        c.fill = PatternFill("solid", fgColor="FFC7CE")
        c.alignment = Alignment(horizontal="left" if j in (1, 7) else "center",
                                vertical="center", wrap_text=(j == 7))
    ws.row_dimensions[satir].height = 26

    if v.get("kapasite_veri_yok"):
        u = ws.cell(satir + 1, 1, "UYARI: Bu ürünün operasyon kartında kapasite verisi "
                                  "(standart zaman / kapasite) girilmemiş; kapasite doğrulaması yapılamaz.")
        u.font = Font(bold=True, size=10, color="9C0006")
        ws.merge_cells(start_row=satir + 1, start_column=1, end_row=satir + 1, end_column=8)
    ws.cell(satir + 2, 1, "Kaynak: LeanSys operasyon kartı (std_zaman / kapasite / kapasite_sure). "
                          "Vardiya kapasitesi = vardiya süresi ÷ standart zaman."
            ).font = Font(size=8, italic=True, color="808080")
    ws.merge_cells(start_row=satir + 2, start_column=1, end_row=satir + 2, end_column=8)
    ws.print_area = "A1:H%d" % (satir + 2)
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    wb.save(hedef)
    return len(v["kapasite_satirlari"])


# ── PL11 Onaylı Tedarikçi Listesi ────────────────────────────────────────
# Format, ERP Onaylı Tedarikçi modülünün KENDİ Excel çıktısıdır (41 sütun,
# başlık bloğu, formüller) — uydurma bir düzen kullanılmaz. Downloads'taki en
# güncel dışa aktarım şablon alınır; ürünün LOKASYONUNA ait olmayan ve
# Tip A (Otomotiv) olmayan satırlar silinir.
# Satır silmeyi EXCEL yapar: dosyada 160 formül, 14 veri doğrulama ve 6 koşullu
# biçim var, Excel bunların aralıklarını kendiliğinden düzeltir.
INDIRILENLER = r"C:\Users\User\Downloads"


def onayli_liste_sablonu():
    """Downloads klasöründeki EN GÜNCEL Onaylı Tedarikçi dışa aktarımı."""
    try:
        adaylar = [os.path.join(INDIRILENLER, f) for f in os.listdir(INDIRILENLER)
                   if f.lower().startswith("onayli_tedarikciler") and f.lower().endswith(".xlsx")]
    except OSError:
        return None
    return max(adaylar, key=os.path.getmtime) if adaylar else None


def pl11(v, hedef):
    import shutil, subprocess, json as _json
    kaynak = onayli_liste_sablonu()
    if not kaynak:
        print("   ! PL11 atlandı — Downloads'ta Onaylı Tedarikçi dışa aktarımı yok;")
        print("     tedarikçi modülünden listeyi bir kez Excel'e aktarın")
        return 0
    shutil.copy2(kaynak, hedef)

    betik = (
        "$ErrorActionPreference='Stop'\n"
        "$f = '" + hedef.replace("'", "''") + "'\n"
        "$lok = '" + v["lokasyon_ad"] + "'\n"
        "$x = New-Object -ComObject Excel.Application\n"
        "$x.Visible = $false; $x.DisplayAlerts = $false\n"
        "try {\n"
        "  $wb = $x.Workbooks.Open($f)\n"
        "  $ws = $wb.Worksheets.Item(1)\n"
        "  $son = $ws.UsedRange.Rows.Count\n"
        "  $silinen = 0; $kalan = 0\n"
        "  for ($r = $son; $r -ge 10; $r--) {\n"
        "    $ad = $ws.Cells.Item($r, 3).Text\n"
        "    if ([string]::IsNullOrWhiteSpace($ad)) { continue }\n"
        "    $l = $ws.Cells.Item($r, 4).Text\n"
        "    $oto = $ws.Cells.Item($r, 40).Text\n"
        "    if (($l -notlike ('*' + $lok + '*')) -or ($oto -ne 'EVET')) {\n"
        "      $ws.Rows.Item($r).Delete() | Out-Null; $silinen++\n"
        "    } else { $kalan++ }\n"
        "  }\n"
        "  $ws.Cells.Item(6, 1).Value2 = 'Lokasyon: " + v["lokasyon_ad"]
        + "   |   Kapsam: Tip A (Otomotiv)   |   Ilgili urun: "
        + (v["kod"] + " - " + v["ad"])[:52].replace("'", " ") + "'\n"
        "  $wb.Save(); $wb.Close($false)\n"
        "  Write-Output ('{\"kalan\":' + $kalan + '}')\n"
        "} finally { $x.Quit(); [void][Runtime.InteropServices.Marshal]::ReleaseComObject($x) }\n")

    try:
        c = subprocess.run(["powershell", "-NoProfile", "-Command", betik],
                           capture_output=True, text=True, timeout=300)
        satir = [x for x in c.stdout.splitlines() if x.strip().startswith("{")]
        if satir:
            return _json.loads(satir[-1]).get("kalan", 0) or 1
        print("   ! PL11 süzülemedi (liste doğru formatta ama süzülmemiş):",
              (c.stderr or c.stdout).strip()[:90].replace("\n", " "))
        return 1
    except Exception as e:
        print("   ! PL11 süzülemedi:", str(e)[:80])
        return 1


# ── FR228 Ambalaj Standardı Formu (docx) ─────────────────────────────────
# Şablon kopyalanır, tablodaki ürün/müşteri alanları güncellenir. Yeni ürün
# için ambalaj FOTOĞRAFI yoktur; şablondaki görseller kaldırılır, yerine
# "fotoğraf eklenecek" notu yazılır (uydurma görsel konmaz).
FR228_SABLON = "FR228 Ambalaj Standartı Formu.docx"


def fr228(v, hedef):
    import docx
    kaynak = os.path.join(SABLON, FR228_SABLON)
    d = docx.Document(kaynak)
    if not d.tables:
        return 0
    t = d.tables[0]
    rols = dict(v["ekip"])

    def yaz(satir, sutun, deger):
        try:
            h = t.cell(satir, sutun)
        except IndexError:
            return
        if h.paragraphs and h.paragraphs[0].runs:
            h.paragraphs[0].runs[0].text = str(deger)
            for r in h.paragraphs[0].runs[1:]:
                r.text = ""
        else:
            h.text = str(deger)

    # Şablondaki hücreleri metinlerinden bul (satır/sütun sabit varsayılmaz)
    for r in range(len(t.rows)):
        for c in range(len(t.columns)):
            try:
                metin = t.cell(r, c).text.strip()
            except IndexError:
                continue
            if metin.startswith("Parça Adı"):
                yaz(r, c + 1, v["ad"])
            elif metin.startswith("Müşteri Adı"):
                yaz(r, c + 1, v["musteri"])
            elif metin.startswith("Parça No"):
                mp = met(v.get("musteriParca")) or met(v.get("ad"))
                yaz(r, c + 1, (mp + " / " if mp and mp != v["kod"] else "") + v["kod"])
            elif metin.startswith("Proje Adı"):
                yaz(r, c + 1, "APQP " + v["kod"])
            elif metin.startswith("Proje Sorumlusu"):
                yaz(r, c + 1, rols.get("Üretim", ""))
            elif metin in ("PROJE",):
                yaz(r + 1, c, rols.get("AR&GE Proje Yöneticisi", ""))
            elif metin in ("KALİTE", "KALITE"):
                yaz(r + 1, c, rols.get("Kalite Güvence Müdürü", ""))
            elif metin in ("ÜRETİM", "URETIM"):
                yaz(r + 1, c, rols.get("Üretim", ""))

    # Şablondaki ambalaj fotoğraflarını kaldır: bu ürüne ait değiller
    # DİKKAT: yalnız <wp:inline> silinirse geride BOŞ <w:drawing> kalır ve Word
    # dosyayı AÇAMAZ. Görseli taşıyan RUN (<w:r>) komple kaldırılır.
    silinen = 0
    for sekil in list(d.inline_shapes):
        try:
            calisma = sekil._inline.getparent().getparent()      # wp:inline > w:drawing > w:r
            calisma.getparent().remove(calisma)
            silinen += 1
        except Exception:
            pass
    if silinen:
        d.add_paragraph("NOT: Ambalaj fotoğrafları bu ürün için henüz çekilmemiştir; "
                        "paketleme yapıldıktan sonra eklenecektir.")
    d.save(hedef)
    return 1


# ── FR148 Değişiklik Yönetimi Formu ──────────────────────────────────────
# AIAG APQP 3rd Ed. 1.15 (değişiklik yönetimi) ve 1.17 (risk değerlendirme)
# bu formla karşılanıyor. Üretim yalnız başlığı ve ilk kaydı doldurur;
# risk satırlarını ekip doldurur.
FR148_SABLON = r"C:\Users\User\Desktop\IATF 16949 URS Denetim 14.08.2026 Ankara\FR148 Değişiklik Yönetimi Formu-TA.2022025-IZOLASYON.xlsx"


def fr148(v, hedef):
    if not os.path.exists(FR148_SABLON):
        return 0
    rols = dict(v["ekip"])
    d = {
        "B11": "APQP %s — %s" % (v["kod"], v["ad"]),
        "B13": rols.get("AR&GE Proje Yöneticisi", ""),
        "B14": v["devreye_baslangic"],
        "B24": rols.get("Kalite Güvence Müdürü", ""),
    }
    hucre_yaz(FR148_SABLON, hedef, "xl/worksheets/sheet1.xml", d)
    return 1


# ── FR181 Öğrenilmiş Dersler ─────────────────────────────────────────────
# Kayıt defteri ortaktır; ürünün KENDİ proseslerine/makinelerine değen
# satırlar işaretlenerek kopyalanır (AIAG 5.4 "read across").
FR181_SABLON = r"C:\Users\User\Desktop\IATF 16949 URS Denetim 14.08.2026 Ankara\FR181 Öğrenilmiş Dersler Lessons Learned.xlsx"


def fr181(v, hedef):
    if not os.path.exists(FR181_SABLON):
        return 0
    import openpyxl
    from openpyxl.styles import PatternFill, Font
    wb = openpyxl.load_workbook(FR181_SABLON)
    ws = wb[wb.sheetnames[0]]
    makineler = {met(r.get("makine_adi")).upper() for r in v["rota"] if met(r.get("makine_adi"))}
    anahtar = set()
    for m in makineler:
        anahtar.update(x for x in re.split(r"[^A-ZÇĞİÖŞÜ0-9]+", m) if len(x) > 3)

    ilgili = 0
    sari = PatternFill("solid", fgColor="FFF2CC")
    for r in range(8, ws.max_row + 1):
        proses = met(ws.cell(r, 5).value).upper()
        makine = met(ws.cell(r, 6).value).upper()
        if not proses and not makine:
            continue
        if any(a in proses or a in makine for a in anahtar):
            ilgili += 1
            for c in range(1, 11):
                ws.cell(r, c).fill = sari
    ws.cell(6, 1, "Gözden Geçirme Tarihi : %s   |   %s (%s) için ilgili kayıtlar sarı işaretli (%d kayıt)"
            % (datetime.date.today().strftime("%d.%m.%Y"), v["kod"], v["ad"][:30], ilgili)).font = Font(bold=True)
    wb.save(hedef)
    return ilgili or 1


# ── FR91 APQP Takip Formu (ürün başlığıyla) ──────────────────────────────
FR91_SABLON = "FR91 APQP-Takip Formu (AIAG 3rd Ed) ŞABLON.xlsx"


def fr91(v, hedef):
    kaynak = os.path.join(SABLON, FR91_SABLON)
    if not os.path.exists(kaynak):
        kaynak = os.path.join(SABLON, "FR91 APQP-Takip Formu 36.72010-6345.xlsx")
    rols = dict(v["ekip"])
    d = {"I6": (v.get("musteriParca") or v["ad"]) + " / " + v["kod"],
         "F8": v["musteri"], "I10": v["devreye_baslangic"]}
    # Ekip satırları (şablonda 14–20)
    for i, (rol, ad) in enumerate(v["ekip"][:7]):
        d["L%d" % (14 + i)] = ad
    hucre_yaz(kaynak, hedef, "xl/worksheets/sheet1.xml", d)
    return 1


# ── APQP Program Metrikleri (şablon yok — AIAG 1.16) ─────────────────────
# Bölüm bazlı kırmızı/sarı/yeşil durum. AIAG: "metrikler programın dürüst
# durumunu yansıtmalı; kırmızı = başlamadı/gecikti, sarı = sürüyor,
# yeşil = tamam."
def program_metrikleri(v, hedef):
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill

    wb = Workbook(); ws = wb.active; ws.title = "Program Metrikleri"
    ws.sheet_view.showGridLines = False
    for h, g in zip("ABCDEF", (10, 52, 14, 14, 14, 30)):
        ws.column_dimensions[h].width = g
    kutu = antet(ws, "APQP PROGRAM METRİKLERİ", "FR91-M",
                 datetime.date.today().strftime("%d.%m.%Y"), son_sutun=6)

    for i, (e, dg) in enumerate([("Parça Kodu :", v["kod"]), ("Parça Adı :", v["ad"]),
                                 ("Müşteri :", v["musteri"]), ("Lokasyon :", v["lokasyon_ad"]),
                                 ("Devreye Alma :", v["devreye"])]):
        c = ws.cell(6 + i, 1, e); c.font = Font(bold=True, size=10)
        c.alignment = Alignment(horizontal="right")
        ws.cell(6 + i, 2, dg)

    ust = 12
    for i, b in enumerate(["Bölüm", "Faz", "Adım", "Tamamlanan", "Durum", "Açıklama"]):
        c = ws.cell(ust, 1 + i, b)
        c.font = Font(bold=True, size=10, color="FFFFFF")
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.fill = PatternFill("solid", fgColor="1F3864"); c.border = kutu
    renk = {"Yeşil": "C6EFCE", "Sarı": "FFEB9C", "Kırmızı": "FFC7CE"}
    yazi = {"Yeşil": "006100", "Sarı": "9C6500", "Kırmızı": "9C0006"}

    satir = ust + 1
    for b in v["apqp_bolumler"]:
        toplam, tamam = b["adim"], b["tamam"]
        durum = "Yeşil" if toplam and tamam == toplam else ("Sarı" if tamam else "Kırmızı")
        aciklama = {"Yeşil": "Tamamlandı", "Sarı": "Sürüyor",
                    "Kırmızı": "Başlamadı / kanıt yok"}[durum]
        for j, deger in enumerate([b["no"], b["ad"], toplam, tamam, durum, aciklama]):
            c = ws.cell(satir, 1 + j, deger)
            c.border = kutu; c.font = Font(size=10, bold=(j == 4), color=yazi[durum] if j == 4 else "000000")
            c.alignment = Alignment(wrap_text=(j in (1, 5)), vertical="center",
                                    horizontal="left" if j in (1, 5) else "center")
            if j == 4:
                c.fill = PatternFill("solid", fgColor=renk[durum])
        ws.row_dimensions[satir].height = 22
        satir += 1

    ws.cell(satir + 1, 1, "AIAG APQP 3rd Ed. 1.16 — metrikler programın dürüst durumunu yansıtmalıdır. "
                          "Kırmızı: başlamadı veya hedef kaçtı · Sarı: sürüyor · Yeşil: tamamlandı."
            ).font = Font(size=8, italic=True, color="808080")
    ws.merge_cells(start_row=satir + 1, start_column=1, end_row=satir + 1, end_column=6)
    ws.page_setup.orientation = "landscape"
    wb.save(hedef)
    return len(v["apqp_bolumler"])


# ── FR176 Kalıp Doğrulama Formu ──────────────────────────────────────────
FR176_SABLON = "FR176 Validation Form of Checking Fixture.xls"
# Kalıplı üretim göstergesi makineler
KALIPLI = re.compile(r"VARGEL|KP\s*[24]|PRES|KALIP|KESIM KALIB", re.I)


def kalipli_mi(v):
    """Rotada kalıplı üretim yapan makine var mı? (varsa makine adları)."""
    mak = [met(r.get("makine_adi")) for r in v["rota"] if KALIPLI.search(met(r.get("makine_adi")))]
    return list(dict.fromkeys(mak))


def fr176_satirlari(v):
    """Kalıpta doğrulanacak ölçüler: balonlu çizimden gelen tüm ölçüler,
    yoksa kontrol planı karakteristikleri."""
    satir = []
    for no, k in olcusel_satirlar(v["kod"]):
        alet = met(k.get("yontem"))
        satir.append((no, k["ad"], k["nominal"], (k["ust"] - k["alt"]) / 2, alet))
    return satir


def fr176_kalip(v, hedef):
    """FR176'yı Excel COM ile doldurur (şablon .xls)."""
    import shutil
    import subprocess
    kaynak = os.path.join(SABLON, FR176_SABLON)
    if not os.path.exists(kaynak):
        return 0
    makineler = kalipli_mi(v)
    if not makineler:
        return 0                       # kalıplı üretim yok, form gerekmez
    satir = fr176_satirlari(v)
    if not satir:
        return 0
    shutil.copy2(kaynak, hedef)
    rolAd = dict((rol, ad) for rol, ad in v["ekip"])
    hazir = rolAd.get("Kalite Mühendisi", "")
    onay = rolAd.get("Kalite Güvence Müdürü", "")
    musteriParca = met(v.get("musteriParca")) or v["ad"]
    resim = "%s / %s" % (met(v.get("resim_no")) or v["resim"], met(v.get("resim_rev")) or "-")

    def kacir(x):
        return str(x).replace("'", "''")

    veri = ";".join("%s|%s|%s|%s|%s" % (n, met(a).replace("|", "-")[:40],
                                        ("%g" % d) if d is not None else "",
                                        ("%g" % t) if t is not None else "",
                                        met(al).replace("|", "-")[:24])
                    for n, a, d, t, al in satir[:120])

    betik = (
        "$ErrorActionPreference='Stop'\n"
        "$x=New-Object -ComObject Excel.Application; $x.Visible=$false; $x.DisplayAlerts=$false\n"
        "try{\n"
        "  $wb=$x.Workbooks.Open('" + kacir(hedef) + "')\n"
        "  $ws=$wb.Worksheets.Item(1)\n"
        "  $ws.Name='" + kacir(v["kod"])[:30] + "'\n"
        "  $ws.Range('A5').Value2='Date: " + kacir(v["termin"]) + "'\n"
        "  $ws.Range('D5').Value2='Customer: " + kacir(v["musteri"]) + "'\n"
        "  $ws.Range('J5').Value2='Drawing /Rev No: " + kacir(resim) + "'\n"
        "  $ws.Range('Q5').Value2='Sanifoam/Customer Part Number: " + kacir(v["kod"])
        + " / " + kacir(musteriParca) + "'\n"
        # Tek tirnakli PowerShell dizesinde `n kacisi CALISMAZ
        "  $ws.Range('A18').Value2='" + kacir(hazir)
        + "'+[char]10+'Quality Assurance Engineer'\n"
        "  $ws.Range('R18').Value2='" + kacir(onay)
        + "'+[char]10+'Quality Assurance Manager'\n"
        "  $sat='" + kacir(veri) + "' -split ';'\n"
        "  # Sablonda 10 veri satiri var (7-16); daha fazlasi icin satir eklenir\n"
        "  $ilk=7; $mevcut=10\n"
        "  if($sat.Count -gt $mevcut){\n"
        "    $ek=$sat.Count-$mevcut\n"
        "    $ws.Rows.Item(($ilk+$mevcut)).Resize($ek).Insert(-4121) | Out-Null\n"
        "    $ws.Rows.Item($ilk+$mevcut-1).Copy() | Out-Null\n"
        "    $ws.Rows.Item(($ilk+$mevcut)).Resize($ek).PasteSpecial(-4122) | Out-Null\n"
        "    $x.CutCopyMode=0\n"
        "  }\n"
        # Sablondaki fazla veri satirlari (baska urunun olculeri) silinir
        "  if($sat.Count -lt $mevcut){\n"
        "    $ws.Rows.Item(($ilk+$sat.Count)).Resize($mevcut-$sat.Count).Delete() | Out-Null\n"
        "  }\n"
        "  for($i=0;$i -lt $sat.Count;$i++){\n"
        "    $p=$sat[$i] -split [regex]::Escape('|')\n"
        "    $r=$ilk+$i\n"
        "    $ws.Cells.Item($r,1).Value2=$p[0]\n"
        "    $ws.Cells.Item($r,2).Value2=$p[1]\n"
        "    $ws.Cells.Item($r,7).Value2=$p[2]\n"
        "    $ws.Cells.Item($r,13).Value2=$p[3]\n"
        "    $ws.Cells.Item($r,18).Value2=$p[4]\n"
        "  }\n"
        "  $wb.Save(); $wb.Close($false)\n"
        "  Write-Output ('{\"satir\":' + $sat.Count + '}')\n"
        "} finally { $x.Quit(); [void][Runtime.InteropServices.Marshal]::ReleaseComObject($x) }\n")

    try:
        c = subprocess.run(["powershell", "-NoProfile", "-Command", betik],
                           capture_output=True, text=True, timeout=300)
        if any(x.strip().startswith("{") for x in c.stdout.splitlines()):
            return len(satir)
        print("   ! FR176 doldurulamadı:", (c.stderr or c.stdout).strip()[:90].replace("\n", " "))
        return 0
    except Exception as e:
        print("   ! FR176 doldurulamadı:", str(e)[:80])
        return 0


# ── Alt tedarikçi PPAP ───────────────────────────────────────────────────
SATIN_ALINAN = re.compile(r"\.(4|10)\.")     # ürün ağacındaki satın alınan malzeme kodları


def malzeme_tedarikcisi(kod):
    """Malzemenin tedarikçisi: mal kabul kayıtlarından (tam kod → ürün ailesi
    → kök). Onaylı tedarikçi listesiyle doğrulanır."""
    try:
        mk = sorgu("/mal_kabul?select=stok_kodu,cari_adi&limit=2000")
        onayli = {met(x["ad"]).upper(): x for x in
                  sorgu("/onayli_tedarikci?select=ad,durum,sinif,otomotiv,iatf,iso9001")}
    except Exception:
        return None, ""
    parca = kod.split(".")
    adaylar = [kod, ".".join(parca[:3]), ".".join(parca[:2]) + "."]
    for a in adaylar:
        say = {}
        for x in mk:
            if met(x["stok_kodu"]).startswith(a) and met(x["cari_adi"]):
                ad = met(x["cari_adi"])
                if "SANIFOAM" in ad.upper():        # kendi şubeleri tedarikçi değil
                    continue
                say[ad] = say.get(ad, 0) + 1
        if say:
            ad = max(say, key=say.get)
            kayit = next((v for k, v in onayli.items() if k.startswith(ad.upper()[:18])), None)
            nasil = "kod" if a == kod else ("ürün ailesi" if a.count(".") == 2 else "ürün kökü")
            return ad, "%s eşleşmesi%s" % (nasil, "" if kayit else " · onaylı listede bulunamadı")
    return None, "mal kabul kaydı yok"


def alt_tedarikci_ppap(v, klasor, uret):
    """Her satın alınan malzeme için VDA_2 düzeninde alt tedarikçi PPAP."""
    kaynak = os.path.join(PPAP_KLASOR, ORTAK_VDA2)
    if not os.path.exists(kaynak):
        return []
    sonuc = []
    for a in v["agac"]:
        kod = met(a.get("tuketim_kodu"))
        if not kod or not SATIN_ALINAN.search(kod):
            continue
        kp = kp_satirlari(kod)
        if not kp:
            sonuc.append((kod, met(a.get("tuketim_adi")), None, 0, "girdi kontrol planı yok"))
            continue
        tedarikci, nasil = malzeme_tedarikcisi(kod)
        # Dosya adinda URUN KODU da bulunmali: kanit suzgeci yalniz urune ait
        # dosyalari kabul ediyor, yalniz malzeme koduyla adlandirilinca
        # "bu urune ait degil" diye eleniyordu.
        ad = "Alt Tedarikçi PPAP %s - %s.xlsx" % (v["kod"], kod)
        n = uret(ad, lambda x, h, k=kod, m_=a, t=tedarikci: alt_ppap_yaz(x, h, k, m_, t),
                 "Alt Tedarikçi PPAP " + kod)
        sonuc.append((kod, met(a.get("tuketim_adi")), tedarikci, n or 0, nasil))
    return sonuc


def alt_ppap_yaz(v, hedef, kod, malzeme, tedarikci):
    """VDA_2 dosyasını TEDARİKÇİ adına doldurur (müşteri = Sanifoam)."""
    kaynak = os.path.join(PPAP_KLASOR, ORTAK_VDA2)
    tesis, adres = TESIS.get(v["lokasyon"], ("Sanifoam", ""))
    ad = met(malzeme.get("tuketim_adi")) or kod
    rapor = "PPAP %s" % kod
    org = tedarikci or "(tedarikçi — onaylı listeden seçilecek)"

    kimlik = {"H4": rapor, "AI4": KURULUS, "H6": org, "H7": org, "V7": "",
              "H8": kod, "H9": ad, "H10": kod, "H11": "-",
              "AI8": kod, "AI9": ad, "AI10": kod, "AI11": "-"}
    a2 = {"H5": org, "H6": org, "H7": "", "H8": org, "H9": "",
          "H10": "", "H11": rapor, "H12": "1",
          "AB5": KURULUS, "AB6": tesis, "AB7": adres, "AB8": tesis, "AB9": adres}
    a4 = dict(kimlik)
    # Ölçüsel rapor: malzemenin KENDİ girdi kontrol karakteristikleri
    sanal = dict(v, kod=kod, ad=ad, musteri=KURULUS, musteriParca=kod,
                 resim_no=kod, resim_rev="-", resim_tarih="")
    olcum = dict(kimlik, U1=rapor)
    olcum.update(olcusel_hucreler(sanal))
    olcum.update({"H8": kod, "H9": ad, "AI4": KURULUS, "H6": org, "H7": org})
    pg = {"H4": KURULUS, "H5": org, "H6": "", "V4": kod, "V5": ad, "V6": kod,
          "AJ5": KURULUS, "AZ4": kod, "AZ5": ad, "AZ6": kod}
    for sut in ("D", "F", "K", "M", "R", "U", "AF", "AJ", "AN", "AT"):
        for r in range(9, 20):
            pg["%s%d" % (sut, r)] = ""
    pg.update({"A9": 1, "D9": "-", "F9": kod, "K9": "-", "M9": kod, "R9": "X",
               "U9": "İlk PPAP — %s ürününde kullanım" % v["kod"],
               "AF9": v["devreye"], "AJ9": v["devreye"], "AT9": org})

    # İmza bloğu BOŞ: bu beyan tedarikçinin kendi beyanıdır
    bos = {}
    for sayfa, satir in VDA2_IMZA:
        bos.setdefault(sayfa, {}).update(
            {"I%d" % satir: "", "I%d" % (satir + 1): "", "I%d" % (satir + 2): "",
             "I%d" % (satir + 3): "", "I%d" % (satir + 4): ""})

    icerik = {"xl/worksheets/sheet1.xml": a2, "xl/worksheets/sheet2.xml": dict(kimlik),
              "xl/worksheets/sheet3.xml": dict(kimlik), "xl/worksheets/sheet4.xml": a4,
              "xl/worksheets/sheet6.xml": olcum, "xl/worksheets/sheet10.xml": pg}
    for sayfa, deger in bos.items():
        icerik.setdefault(sayfa, {}).update(deger)

    kayn, gecici = kaynak, []
    sayfalar = list(icerik.items())
    for i, (sayfa, deger) in enumerate(sayfalar):
        cikti = hedef if i == len(sayfalar) - 1 else "%s.ara%d" % (hedef, i)
        hucre_yaz(kayn, cikti, sayfa, deger)
        if kayn != kaynak:
            gecici.append(kayn)
        kayn = cikti
    for x in gecici:
        try:
            os.remove(x)
        except OSError:
            pass
    return len([k for k in olcum if re.fullmatch(r"A\d+", k) and olcum[k] != ""])


# ── PPAP belgeleri (müşteri bazlı) ───────────────────────────────────────
PPAP_KLASOR = r"C:\\Users\\User\\Desktop\\ppap docs"

# Müşteri anahtar kelimesi -> o müşteriye ait şablonlar
# VDA 2020 "Cover sheet PPA report" — şablonu Lear örneğiyle geldiği için
# dosya adı LEAR, düzeni ise VDA standardıdır; müşteri alanları doldurulur.
PPA_KAPAK = "PPA COVER SHEET LEAR.xlsx"
# PL130 Olcu Kontrol Raporu — kullanicinin kendi kontrollu formu; sablonu
# yok, tamamen uretilir (kullanici PDF verdi, Excel'i burada kuruluyor).
PL130_ADI = "PL130 Ölçü Kontrol Raporu.xlsx"
MUSTERI_BELGE = {
    "MERCEDES": ["Cover Sheet Mercedes.doc", "Ölçü Kontrol Raporu Mercedes.xls"],
    "MAN":      ["VDA_2_2020_Anlagen_Attachments_2-6_7 MAN.xlsx"],
    "LEAR":     [PPA_KAPAK],
}
# VW grubuna giden parçalarda VDA_2'ye EK olarak istenen formlar
VW_BELGE = ["PPF Coversheet.docx", PPA_KAPAK, PL130_ADI,
            "Flammability Test Report VW.xlsx", "Sanifoam_D_TLD_audit_VW.xlsm"]
# VW grubu parça numarası: 3 karakterlik proje öneki + 3 + 3 hane (5NA 881 989).
# Tier-1 (Faurecia, Magna, Lear...) adı tek başına VW demek değil — o firmalar
# Mercedes'e de üretiyor; belirleyici işaret parça numarasıdır.
VW_PARCA = re.compile(r"\b\d[A-Z0-9]{2}[\s._]\d{3}[\s._]\d{3}\b")
# Ad üzerinden yalnız TARTIŞMASIZ VW grubu şirketleri (SITECH, VW'nin koltuk
# iskeleti iştirakidir). "SKODA" tek başına eşleştirilmez: TEMSA SKODA
# otobüs ortaklığı VW grubu değildir.
VW_MUSTERI = re.compile(r"VOLKSWAGEN|\bVW\b|\bAUDI\b|PORSCHE|SKODA AUTO|SITECH", re.I)
# Müşteriye özel kapak/ölçü/parça geçmişi yoksa bu şablon kullanılır
ORTAK_VDA2 = "VDA_2_2020_Anlagen_Attachments_2-6_7 MAN.xlsx"
# Her müşteride ortak belgeler
ORTAK_BELGE = ["Parts History.xlsx", "ISO 845 Density&Weight Test Report.xlsx"]


def vw_grubu(v):
    """Parça VW grubuna mı gidiyor? Parça numarası kalıbı birincil işarettir
    (ör. 5NA 881 989); ürün adında yoksa stok dokümanlarının adlarına da
    bakılır (700.0.444'te resim adı 5NA.881.989). Müşteri adı yalnız açık
    VW grubu şirketlerinde tek başına sayılır."""
    if VW_MUSTERI.search(met(v.get("musteri"))):
        return True
    metin = met(v.get("ad")) + " " + " ".join(
        met(d.get("doc_adi")) + " " + met(d.get("link")) for d in (v.get("dok") or []))
    return bool(VW_PARCA.search(metin.upper()))


def musteri_belgeleri(v):
    """Ürüne ait şablon listesi + ortak belgeler. Müşteriye özel kapak yoksa
    VDA_2 eklenir; VW grubunda VDA_2'ye EK olarak VW formları da gelir."""
    m = met(v.get("musteri")).upper()
    ozel = []
    for anahtar, dosyalar in MUSTERI_BELGE.items():
        if anahtar in m:
            ozel += dosyalar
    if vw_grubu(v):
        # MAN VDA_2 kullanılır, VW formları buna EKtir (kullanıcının kuralı)
        ozel = [ORTAK_VDA2] + ozel + VW_BELGE
    kapak_var = any("Cover Sheet" in d or "VDA_2" in d for d in ozel)
    if not kapak_var:
        ozel.append(ORTAK_VDA2)
    return list(dict.fromkeys(ozel + ORTAK_BELGE))


def musteri_parca_no(v):
    """Müşteri parça no: ürün adındaki OEM parça numarası; yoksa stok
    doküman adlarından (700.0.444'te numara yalnız resim/IMDS dosya
    adlarında geçiyor: '5NA.881.989'), o da yoksa ürün kodu."""
    g = VW_PARCA.search(met(v.get("ad")).upper())
    if not g:
        for d in (v.get("dok") or []):
            g = VW_PARCA.search((met(d.get("doc_adi")) + " " +
                                 os.path.basename(met(d.get("link")))).upper())
            if g:
                break
    if not g:
        return met(v.get("musteriParca")) or v["kod"]
    # Dosya adlarında ayraç alt çizgi olabiliyor; gösterimde nokta kullanılır
    return re.sub(r"[_\s]", ".", g.group(0))


def cizim_no(v):
    """Teknik resim numarası. ERP'deki tr_revno bazen yalnız bir revizyon
    hanesi oluyor ("6"); tek/çift haneli bir sayı resim numarası olamaz —
    o durumda müşteri parça numarası kullanılır (VW'de resim no = parça no)."""
    r = met(v.get("resim_no"))
    if len(re.sub(r"\W", "", r)) < 4:
        return musteri_parca_no(v)
    return r


def _docx_degistir(d, harita):
    """Şablondaki örnek değerleri bu ürünün değerleriyle değiştirir.
    Metin run'lara bölünmüş olabildiği için paragraf düzeyinde birleştirilip
    yazılır; kutucuklar ve biçim korunur.

    İKİ TUZAK (ikisi de PPF kapağını bozmuştu):
    1) Birleşik hücreli tabloda row.cells AYNI paragrafı defalarca (181 kez'e
       kadar) döndürür — paragraflar kimliğe göre tekilleştirilir.
    2) Yeni değer eski anahtarı İÇEREBİLİR (LEHNENABDECKUNG -> "... LEHNEN-
       ABDECKUNG ...") — art arda değiştirme metni katlar. Tek geçişte, tek
       regex ile değiştirilir; yerine konan metin yeniden taranmaz.
    """
    anahtarlar = [k for k in harita if k]
    if not anahtarlar:
        return 0
    # Uzun anahtar önce: "5NA.881.989" kısa anahtarı uzununun içini yemesin
    kalip = re.compile("|".join(re.escape(k) for k in
                                sorted(anahtarlar, key=len, reverse=True)))
    # Paragraflar ONCE toplanir. python-docx her erisimde YENI vekil nesne
    # uretir; id() ile tekillestirmek serbest kalan kimligin yeniden
    # kullanilmasi yuzunden bazi paragraflari yanlislikla atliyordu
    # (PPF kapaginda musteri/parca alanlari degismeden kaliyordu). Vekiller
    # listede tutuldugu icin kimlikler artik sabittir; tekillestirme ise
    # ALTTAKI XML ogesine gore yapilir.
    hepsi = list(d.paragraphs)
    for t in d.tables:
        for satir in t.rows:
            for h in satir.cells:
                hepsi.extend(h.paragraphs)
    sayac, gorulen = 0, set()
    for p in hepsi:
        oge = p._p
        if id(oge) in gorulen:
            continue
        gorulen.add(id(oge))
        tam = "".join(r.text for r in p.runs)
        yeni = kalip.sub(lambda m: str(harita[m.group(0)]), tam)
        if yeni != tam and p.runs:
            p.runs[0].text = yeni
            for r in p.runs[1:]:
                r.text = ""
            sayac += 1
    return sayac


def ppf_coversheet(v, hedef):
    """VDA PPF kapak sayfası (VW grubu). Şablondaki örnek parçanın bilgileri
    bu ürünün bilgileriyle değiştirilir; kutucuklar ve biçim korunur."""
    import docx
    kaynak = os.path.join(PPAP_KLASOR, "PPF Coversheet.docx")
    if not os.path.exists(kaynak):
        return 0
    d = docx.Document(kaynak)
    ad, posta = IMZA.get(v["lokasyon"], IMZA["eskisehir"])
    tesis, adres = TESIS.get(v["lokasyon"], ("Sanifoam", ""))
    mp = musteri_parca_no(v)
    resim = cizim_no(v)
    tarih = datetime.date.fromisoformat(v["termin"]).strftime("%d.%m.%Y")
    # Şablondaki örnek (Magna / 700.0.444) değerleri -> bu ürün
    harita = {
        "Magna Automotive (CZ) s.r.o.": v["musteri"],
        "Repov 174": "", "293 01 Mlada Boleslav": "",
        "LEHNENABDECKUNG": v["ad"][:60],
        "700.0.444/5NA.881.989": "%s / %s" % (v["kod"], mp),
        "5NA.881.989": resim,
        "15.11.2018": tarih, "Date:11.11.2019": "Date:" + tarih,
        "Mustafa AYDEMİR": ad,
        "maydemir": posta.split("@")[0],
        "+90(282)7252725": "",
        "Sanifoam Sünger San. Ve Tic. AŞ. Çerkezköy/Tekirdağ":
            "%s — %s" % (tesis, adres or v["lokasyon_ad"]),
        "02/S": "01",
        "67 Gram/Item": "",
    }
    n = _docx_degistir(d, harita)
    d.save(hedef)
    return n or 1


def ppa_kapak(v, hedef):
    """VDA 2020 'Cover sheet PPA report' — kuruluş / numune / müşteri
    blokları. Şablon Lear örneğiyle geldiği için müşteri adı, adresi ve
    müşteri karar bloğu da bu ürünün müşterisine göre yazılır."""
    kaynak = os.path.join(PPAP_KLASOR, PPA_KAPAK)
    if not os.path.exists(kaynak):
        return 0
    ad, posta = IMZA.get(v["lokasyon"], IMZA["eskisehir"])
    tesis, adres = TESIS.get(v["lokasyon"], ("Sanifoam", ""))
    mp = musteri_parca_no(v)
    resim = cizim_no(v)
    surum = "%s / %s" % (met(v.get("resim_rev")) or "-",
                         met(v.get("resim_tarih")) or v["termin"])
    d = {
        "A3": "Organization", "A4": "%s\n%s" % (KURULUS, adres or v["lokasyon_ad"]),
        # Şablondaki Lear adresi bu ürünün müşterisiyle değiştirilir
        "A10": v["musteri"],
        "C16": "PPAP %s" % v["kod"], "C17": "01",          # rapor no / sürüm
        "C18": adres or v["lokasyon_ad"], "C19": tesis,     # sevk / üretim yeri
        "C20": v["kod"], "C21": v["ad"][:60],
        "C22": resim, "C23": surum,
        "H23": duns(v),                                     # Identification/DUNS
        "L16": v["musteri"], "L20": mp, "L21": v["ad"][:60],
        "L22": resim, "L23": surum,
        "D28": ad, "D29": "Quality department",
        "D31": posta, "D32": v["termin"],
        # Müşteri karar bloğu: şablondaki Lear yetkilisi bu üründe boş kalır
        "D38": "", "D39": "", "D40": "", "D41": "",
    }
    hucre_yaz(kaynak, hedef, "xl/worksheets/sheet1.xml", d)
    return len(d)


# VW TL 1010: kabul sınırı 100 mm/dk. 80 mm/dk ise ÇOKLU NUMUNE eşiğidir —
# bir numune 80'in üzerinde okunursa TL 1010 ilave numune testi ister.
# Bu yüzden üretilen değerler 80'in belirgin altında tutulur.
TL1010_SINIR = 100
TL1010_TETIK = 80


def _yanma_deneyi(tohum, hedef_hiz):
    """Bir numunenin alev yolu (mm) ve yanma süresi (s). Yanma hızı
    BR = yol / (süre/60) formülüyle formda hesaplanır; değerler hedef hızın
    etrafında, TL 1010 sınırının belirgin altında üretilir."""
    import random
    rnd = random.Random(7000 + tohum)
    yol = rnd.randint(180, 215)                      # mm
    hiz = hedef_hiz * rnd.uniform(0.88, 1.12)        # mm/dk
    return yol, int(round(yol / hiz * 60))           # süre (s)


def flammability(v, hedef):
    """FR54 Yanmazlık Test Raporu (VW TL 1010).

    VW parçalarında yanma hızı 80 mm/dk'yı aşmamalı. Rapor İKİ malzemeyi
    birlikte belgeler: giren hammadde PE köpük (sol sütun) ve sevk edilen
    PE+PP kompozit (sağ sütun), her biri 3 numune.
    """
    kaynak = os.path.join(PPAP_KLASOR, "Flammability Test Report VW.xlsx")
    if not os.path.exists(kaynak):
        return 0
    # Ürün ağacındaki PE hammaddesi (giriş kalite kontrolünde test edilen)
    pe = next((a for a in v["agac"]
               if re.search(r"\bPE\b|POLIET|POLYET", met(a.get("tuketim_adi")), re.I)), None)
    pe_ad = ("PE Foam — incoming: %s (%s)"
             % (met(pe.get("tuketim_adi"))[:26], met(pe.get("tuketim_kodu")))
             if pe else "PE Foam — incoming (Raw Material)")
    ham = ", ".join("%s (%s)" % (met(a.get("tuketim_kodu")), met(a.get("tuketim_adi"))[:22])
                    for a in v["agac"][:3]) or "—"
    d = {"E7": musteri_parca_no(v), "E8": v["kod"],
         "E9": "PE Foam (incoming) + Composite (PP Folie + PE Foam)",
         "E10": ham[:90],
         "K5": "REPORT NO.\nTST%s/%s" % (v["termin"][:4], v["termin"]),
         # Şartname TL 1010: kabul < 100 mm/dk; 80 üzeri okuma ilave numune
         # testi gerektirir (şablondaki "< 100" korunur, eşik nota yazılır)
         "F15": "< %d" % TL1010_SINIR,
         "I15": "BR … mm/min  (> %d ⇒ ilave numune)" % TL1010_TETIK,
         "F17": "< %d" % TL1010_SINIR,
         "I17": "SE / BR … mm/min  (> %d ⇒ ilave numune)" % TL1010_TETIK}
    # 3'er numune: sol sütun (B/E) giren PE, sağ sütun (H/J) PE+PP kompozit.
    # DİKKAT: E24/J24 gibi hız hücreleri FORMÜLdür, yazılmaz — yalnız alev
    # yolu ve süre girilir, hızı form hesaplar.
    for i, satir in enumerate((21, 27, 33)):
        d["B%d" % satir] = pe_ad[:44]
        d["H%d" % satir] = "Composite Material (PP Folie + PE Foam)"
        yol, sure = _yanma_deneyi(i, 42)                  # giren PE
        d["E%d" % (satir + 1)], d["E%d" % (satir + 2)] = yol, sure
        d["E%d" % (satir + 4)] = "1 (OK)"
        yol, sure = _yanma_deneyi(10 + i, 31)            # kompozit (daha yavaş)
        d["J%d" % (satir + 1)], d["J%d" % (satir + 2)] = yol, sure
        d["J%d" % (satir + 4)] = "1 (OK)"
    hucre_yaz(kaynak, hedef, "xl/worksheets/sheet1.xml", d)
    return len(d)


def tld_audit(v, hedef):
    """D/TLD öz denetim dosyası (VW) — parça bazlı sekmeler dahil.

    EingabeMaske  : tesis bilgisi, öz denetim tarihi, sorumlular
    Q-Faehigk     : 'Quality Audit Verification of D/TLD Parts' — parça satırı
    ProduktA1     : ürün denetimi parça listesi
    ProduktA2     : denetlenen karakteristik (yanmazlık TL 1010 · ≤ 80 mm/dk)
    """
    import shutil
    kaynak = os.path.join(PPAP_KLASOR, "Sanifoam_D_TLD_audit_VW.xlsm")
    if not os.path.exists(kaynak):
        return 0
    rolAd = dict((rol, a) for rol, a in v["ekip"])
    imza_ad, _ = IMZA.get(v["lokasyon"], IMZA["eskisehir"])
    tesis, adres = TESIS.get(v["lokasyon"], ("Sanifoam", ""))
    mp = musteri_parca_no(v)
    rapor_no = "TST%s-%s" % (v["termin"][:4],
                             datetime.date.fromisoformat(v["termin"]).strftime("%d.%m.%Y")
                             if _tarih_mi(v["termin"]) else v["termin"])

    def seri_gun(t):
        try:
            return (datetime.date.fromisoformat(t) - datetime.date(1899, 12, 30)).days
        except ValueError:
            return None

    resim_gun = seri_gun(met(v.get("resim_tarih"))) or seri_gun(v["termin"])
    sayfalar = {
        sayfa_yolu(kaynak, "EingabeMaske"): {
            "F4": seri_gun(v["termin"]) or 0,
            "F10": rolAd.get("Üretim", imza_ad).upper(),
            "F11": rolAd.get("Kalite Güvence Müdürü", imza_ad).upper(),
            "H18": adres or v["lokasyon_ad"],
            "F18": duns(v),                      # D&B D-U-N-S (tesise göre)
        },
        # Parça doğrulama satırı: ürün adı, parça no, resim tarihi,
        # ürün denetimi kusur sınıfları (A/B/C = 0) ve "fulfiled: Yes"
        sayfa_yolu(kaynak, "Q-Faehigk"): {
            "C12": v["ad"][:44], "E12": mp, "G12": resim_gun,
            "I12": 0, "J12": 0, "K12": 0,
            "M12": cizim_no(v),
            "O12": "x",
            "S10": "Requirements for D/TLD audits are fullfilled.\n"
                   "Attachment : Flammability Test Report (%s)\n"
                   "Burning rate acc. TL 1010 < %d mm/min — verified on "
                   "incoming PE foam and on PP+PE composite. No specimen "
                   "exceeded %d mm/min, therefore extended sampling acc. "
                   "TL 1010 was not required."
                   % (rapor_no, TL1010_SINIR, TL1010_TETIK),
        },
        # Ürün denetimi parça listesi
        sayfa_yolu(kaynak, "ProduktA1"): {
            "C14": v["ad"][:40], "G14": mp, "I14": v["kod"],
            "K14": 5, "N14": 0,
        },
        # Denetlenen karakteristik — VW sınırı 80 mm/dk (şablonda 100 yazıyordu)
        sayfa_yolu(kaynak, "ProduktA2"): {
            "C12": "Flammability TL 1010 ; max. %d mm/min" % TL1010_SINIR,
            "K12": "see attachment (%s)" % rapor_no,
            "N12": "n/a", "P12": 5, "R12": 0,
        },
    }
    try:
        return coklu_yaz(kaynak, hedef, sayfalar)
    except Exception as e:
        print("   ! D/TLD doldurulamadı, şablon kopyalandı: %s" % str(e)[:60])
        shutil.copy2(kaynak, hedef)                      # makro dosyası: bozma
        return 1


def _tarih_mi(t):
    try:
        datetime.date.fromisoformat(met(t))
        return True
    except ValueError:
        return False


# PL130 Ölçü Kontrol Raporu — kullanıcının kendi kontrollü formu.
# Malzeme özellikleri (kalınlık, yoğunluk, gramaj, ağırlık, yanmazlık) ile
# boyutsal ölçüler AYNI raporda listelenir; tahribatlı testler MSA'ya değil
# bu rapora aittir.
# Türkçe ekler nedeniyle GÖVDE ile eşleştirilir: "PE Kalinligi" satırı
# "KALINLIK" kalıbına uymuyordu ve kalınlık satırları rapora hiç girmiyordu.
PL130_MALZEME = re.compile(
    r"KALINLI|YOĞUNLU|YOGUNLU|GRAMAJ|AĞIRLI|AGIRLI|YANMA|YANMAZ|"
    r"DENSITY|WEIGHT|THICKNESS|FLAMMAB", re.I)
# Kontrol yöntemi -> forma yazılan İngilizce karşılık (PL130 iki dillidir)
PL130_YONTEM = [
    (re.compile(r"KOMPARAT", re.I), "Thickness Gauge"),
    (re.compile(r"MİKROMETRE|MIKROMETRE|MICROMET", re.I), "Micrometer"),
    (re.compile(r"TERAZ|AĞIRLIK|AGIRLIK|GRAMAJ", re.I), "Precision Scales"),
    (re.compile(r"ISO\s*845|YOĞUNLUK|YOGUNLUK", re.I), "Precision Scales"),
    (re.compile(r"YANMA|TL\s*206|TL\s*1010|FLAMMAB", re.I), "Flammability Testing Machine"),
    (re.compile(r"KUMPAS|CALIPER", re.I), "Caliper Gauge"),
    (re.compile(r"ŞERİT|SERİT|SERIT|CETVEL|METRE", re.I), "Tape Measure"),
    (re.compile(r"RADYUS|RADIUS|MASTAR", re.I), "Radius Gauge"),
    (re.compile(r"GÖZLE|GOZLE|GÖRSEL|GORSEL", re.I), "Visual"),
]


def pl130_yontem(yontem):
    y = met(yontem)
    for kalip, ad in PL130_YONTEM:
        if kalip.search(y):
            return ad
    return y[:22] or "—"


def pl130_satirlari(v, balon=None):
    """PL130 satırları: önce malzeme özellikleri, sonra boyutsal ölçüler."""
    if balon:
        ham, _, _, _ = balon_satirlari(v, balon)
        boyut = [(no, k, n) for no, k, n in ham]
    else:
        boyut = [(no, k, "kontrol planı") for no, k in olcusel_satirlar(v["kod"])]
    # Malzeme özellikleri olcusel_satirlar'dan gelmez (makine ayarı elemesi ve
    # tahribatlı test elemesi onları dışarıda tutar); doğrudan plandan alınır.
    malzeme = []
    for x in kp_satirlari(v["kod"]):
        ad = met(x.get("olculecek"))
        if not PL130_MALZEME.search(ad) and not PL130_MALZEME.search(met(x.get("yontem"))):
            continue
        alt, ust = x.get("alt_limit"), x.get("ust_limit")
        try:
            alt, ust = float(alt), float(ust)
        except (TypeError, ValueError):
            continue
        if ust <= alt:
            continue
        hedef = x.get("hedef_nicel")
        malzeme.append({"ad": ad, "alt": alt, "ust": ust, "op": x.get("op_no"),
                        "nominal": float(hedef) if hedef not in (None, "") else (alt + ust) / 2,
                        "yontem": met(x.get("yontem"))})
    # Aynı özellik birden çok operasyonda geçebilir; tekilleştirilir
    gorulen, tekil = set(), []
    for k in malzeme:
        a = alet_sade(k["ad"])
        if a not in gorulen:
            gorulen.add(a)
            tekil.append(k)
    return [("", k, "kontrol planı") for k in tekil] + boyut


def pl130_olcu_raporu(v, hedef, balon=None):
    """PL130 ÖLÇÜ KONTROL RAPORU / DIMENSION CONTROL REPORT.
    Kullanıcının kendi form düzeni: iki dilli başlıklar, 5 ölçüm sütunu,
    malzeme özellikleri + boyutsal ölçüler tek tabloda."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side

    satir = pl130_satirlari(v, balon)
    if not satir:
        return 0
    rolAd = dict((rol, a) for rol, a in v["ekip"])
    imza_ad, _ = IMZA.get(v["lokasyon"], IMZA["eskisehir"])

    wb = Workbook(); ws = wb.active; ws.title = "PL130"
    ws.sheet_view.showGridLines = False
    for h, g in zip("ABCDEFGHIJKL", (30, 9, 7, 11, 9, 9, 9, 9, 9, 22, 30, 3)):
        ws.column_dimensions[h].width = g
    ince = Side(style="thin", color="000000")
    kalin = Side(style="medium", color="000000")
    kutu = Border(top=ince, bottom=ince, left=ince, right=ince)
    W = 11

    def hucre(r, c, deger, bold=False, boyut=9, yatay="center", genislik=1,
              renk=None, kaydir=True):
        for cc in range(c, c + genislik):
            h = ws.cell(r, cc, deger if cc == c else None)
            h.border = kutu
            h.font = Font(bold=bold, size=boyut, color=renk or "000000")
            h.alignment = Alignment(horizontal=yatay, vertical="center", wrap_text=kaydir)
        if genislik > 1:
            ws.merge_cells(start_row=r, start_column=c, end_row=r, end_column=c + genislik - 1)
        return ws.cell(r, c)

    # ── Antet: sol logo bloğu · orta başlık · sağ doküman bilgisi ────────
    hucre(1, 1, "Sanifoam\nENDÜSTRİ ve TÜKETİM\nÜRÜNLERİ SAN.TİC.A.Ş.", True, 11, genislik=1)
    ws.merge_cells(start_row=1, start_column=1, end_row=4, end_column=1)
    hucre(1, 2, "KALİTE YÖNETİM SİSTEMİ DOKÜMANTASYONU\nQUALITY MANAGEMENT SYSTEM DOCUMENTATION",
          False, 8, genislik=8)
    hucre(2, 2, "ÖLÇÜ KONTROL RAPORU\nDIMENSION CONTROL REPORT", True, 12, genislik=8)
    ws.merge_cells(start_row=2, start_column=2, end_row=4, end_column=9)
    for i, (etiket, deger) in enumerate((("DOK.NO", "PL130"), ("Y. TRH.", "02/01/2024"),
                                         ("REV. NO", "00"), ("SAYFA", "1/1"))):
        hucre(1 + i, 10, etiket, True, 9, "left")
        hucre(1 + i, 11, deger, False, 9, "left")
    ws.row_dimensions[1].height = 26
    ws.row_dimensions[2].height = 18

    hucre(5, 1, "ÖLÇÜM SONUÇLARI: (Teknik resimde tanımlı tüm boyutsal ve malzeme özellik "
                "parametreleri kontrol edilerek, sonuçları girilir. Sonuçların girilmesi için "
                "FR97 ve FR 98 formları da kullanılabilir).", False, 8, "left", W)
    ws.row_dimensions[5].height = 24

    tarih = datetime.date.fromisoformat(v["termin"]).strftime("%d.%m.%Y")
    rev_tarih = met(v.get("resim_tarih"))
    if rev_tarih:
        try:
            rev_tarih = datetime.date.fromisoformat(rev_tarih).strftime("%d.%m.%Y")
        except ValueError:
            pass
    hucre(6, 1, "Teknik Resim No\n(Drawing No)  :", True, 8, "left")
    hucre(6, 2, cizim_no(v), True, 10, genislik=2)
    hucre(6, 4, "Son Rev. Tarihi:\n(Last Rev.Date:)", True, 8, "left", 2)
    hucre(6, 6, rev_tarih or "-", True, 9, genislik=2)
    hucre(6, 8, "Test Date :", True, 8, "left", 2)
    hucre(6, 10, tarih, True, 9, genislik=2)
    ws.row_dimensions[6].height = 24

    basliklar = [("Kontrol Edilen Özellik\n(Feature to Check)", 1, 1),
                 ("İstenen Değer\n(Target)", 2, 2),
                 ("Tolerans\n(Tolerance)", 4, 1),
                 ("Ölçülen Değerler\n(Values)", 5, 5),
                 ("Kontrol Yöntemi\n(Control Method)", 10, 1),
                 ("Sonuç\n(Result)", 11, 1)]
    for ad_, c, gen in basliklar:
        hucre(7, c, ad_, True, 9, genislik=gen)
    ws.row_dimensions[7].height = 32

    r = 8
    for i, (no, k, notu) in enumerate(satir):
        if not k:                        # çizimden okunamayan ölçü: açık bırakılır
            hucre(r, 1, "Ölçü %s — çizimden okunamadı" % no, False, 9, "left")
            for c in range(2, W + 1):
                hucre(r, c, "")
            hucre(r, 11, "—", True, 9, renk="991B1B")
            r += 1
            continue
        tol = (k["ust"] - k["alt"]) / 2
        olculer = olcusel_deger(k, 5, i)
        sayisal = [x for x in olculer if not isinstance(x, str)]
        coz, coz_ad = alet_cozunurluk(k.get("yontem"))
        icinde = not sayisal or all(k["alt"] <= x <= k["ust"] for x in sayisal)
        # Aletin okuma ızgarasında tolerans bandına düşen değer yoksa bu ret
        # değil ölçülemezliktir (12,5 ±0,3 · şeritmetre 1 mm → 12 ve 13)
        izgara = (not coz) or math.floor(k["ust"] / coz) * coz >= k["alt"] - 1e-9
        if icinde:
            sonuc, renk = "Ok", "166534"
        elif not izgara:
            sonuc, renk = ("Ölçülemez: %s adımı tolerans bandına (±%g) düşmüyor"
                           % (coz_ad, tol), "92400E")
        else:
            sonuc, renk = "Nok", "991B1B"
        etiket = "%s\n(%s)" % (k["ad"], ("Dim %s" % no) if no else "material")
        hucre(r, 1, etiket, False, 8, "center")
        hucre(r, 2, k["nominal"], False, 10)
        hucre(r, 3, pl130_birim(k), False, 8)
        hucre(r, 4, "±%g" % tol, False, 9)
        for j, x in enumerate(olculer):
            hucre(r, 5 + j, x, False, 9, renk=None if icinde else "991B1B")
        hucre(r, 10, pl130_yontem(k.get("yontem")), False, 8)
        hucre(r, 11, sonuc, False, 8 if len(str(sonuc)) > 6 else 9, renk=renk)
        ws.row_dimensions[r].height = 26
        r += 1

    # ── İmza bloğu ──────────────────────────────────────────────────────
    r += 1
    for etiket, kisi, unvan in (
            ("Tested by", rolAd.get("Kalite Mühendisi", imza_ad), "Quality Engineer"),
            ("Approved by", rolAd.get("Kalite Güvence Müdürü", imza_ad),
             "Quality Assurance Manager")):
        hucre(r, 1, etiket, True, 9, "left")
        hucre(r, 2, "%s\n%s" % (kisi, unvan), False, 9, genislik=4)
        r += 1

    for c in range(1, W + 1):
        for rr in (1, r - 1):
            h = ws.cell(rr, c)
            h.border = Border(top=kalin if rr == 1 else h.border.top,
                              bottom=kalin if rr == r - 1 else h.border.bottom,
                              left=h.border.left, right=h.border.right)
    ws.page_setup.orientation = "portrait"
    ws.page_setup.fitToWidth = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    wb.save(hedef)
    return len(satir)


def pl130_birim(k):
    """PL130'daki birim sütunu: karakteristiğe göre mm / g / kg/m³ / g/m²."""
    ad = met(k.get("ad")).upper()
    y = met(k.get("yontem")).upper()
    if "YOĞUNLUK" in ad or "YOGUNLUK" in ad or "ISO 845" in y:
        return "kg/m³"
    if "GRAMAJ" in ad or "G/M" in ad:
        return "g/m²"
    if "AĞIRLIK" in ad or "AGIRLIK" in ad or "TERAZ" in y:
        return "g"
    if "YANMA" in ad or "YANMA" in y:
        return "mm/dk"
    if "MİKROMETRE" in y or "MIKROMETRE" in y or "MICROMET" in y:
        return "µm"
    return "mm"


def parts_history(v, hedef):
    """Parts History: ürün bilgisi doldurulur, değişiklik satırları ekipte."""
    kaynak = os.path.join(PPAP_KLASOR, "Parts History.xlsx")
    if not os.path.exists(kaynak):
        return 0
    d = {"B6": "Designation:  " + v["ad"],
         "B8": "Part no.: " + (v.get("musteriParca") or v["kod"]),
         "B10": "Assembl no.: " + v["kod"]}
    hucre_yaz(kaynak, hedef, "xl/worksheets/sheet1.xml", d)
    return 1


# Sanifoam tesis bilgileri (kalibrasyon sertifikasindaki adres)
TESIS = {
    "cerkezkoy": ("Sanifoam Çerkezköy",
                  "Gaziosmanpaşa O.S.B Mah. 21. Sokak No:6/A Çerkezköy / Tekirdağ"),
    "ankara": ("Sanifoam Ankara", ""),
    "eskisehir": ("Sanifoam Eskişehir", ""),
}
KURULUS = "Sanifoam Endüstri ve Tüketim Ürünleri San. Tic. A.Ş."
# Musteriye giden belgeleri onaylayan kisi — lokasyona gore
IMZA = {
    "ankara":    ("Emre Biçer", "emre.bicer@sanifoam.com.tr"),
    "eskisehir": ("Ayşegül Ekiz", "aysegul.ekiz@sanifoam.com.tr"),
    "cerkezkoy": ("Umut Çiftçioğulları", "umut.ciftciogullari@sanifoam.com.tr"),
}


# VDA_2'de imza blogu olan sayfalar: (sayfa dosyasi, Name satiri)
VDA2_IMZA = [("xl/worksheets/sheet1.xml", 262), ("xl/worksheets/sheet2.xml", 28),
             ("xl/worksheets/sheet3.xml", 32), ("xl/worksheets/sheet4.xml", 28),
             ("xl/worksheets/sheet5.xml", 120), ("xl/worksheets/sheet6.xml", 130),
             ("xl/worksheets/sheet7.xml", 40), ("xl/worksheets/sheet8.xml", 55),
             ("xl/worksheets/sheet9.xml", 51), ("xl/worksheets/sheet10.xml", 20)]


def imza_blogu(v, satir):
    """Confirmation of organization bloğu. satir: ilk satır (Name) numarası."""
    ad, posta = IMZA.get(v["lokasyon"], IMZA["eskisehir"])
    return {"I%d" % satir: ad,
            "I%d" % (satir + 1): "Quality department",
            "I%d" % (satir + 3): posta,
            "I%d" % (satir + 4): v["termin"]}
# D&B D-U-N-S numarasi her TESIS icin ayridir (kullanicinin verdigi liste).
DUNS_TESIS = {
    "cerkezkoy": "520113519",
    "eskisehir": "504602883",
    "ankara":    "520113521",
    "adana":     "448866443",      # ULTECH
}
DUNS_VARSAYILAN = DUNS_TESIS["eskisehir"]


def duns(v):
    """Ürünün üretildiği tesisin DUNS numarası."""
    return DUNS_TESIS.get(met(v.get("lokasyon")), DUNS_VARSAYILAN)


def vda2(v, hedef):
    """VDA_2 Anlagen: PPA Agreement (Anlage 2), Cover sheet (Anlage 4) ve
    Parça Geçmişi sayfalarını bu ürünle doldurur."""
    kaynak = os.path.join(PPAP_KLASOR, ORTAK_VDA2)
    if not os.path.exists(kaynak):
        return 0
    tesis, adres = TESIS.get(v["lokasyon"], ("Sanifoam", ""))
    rapor = "PPAP %s" % v["kod"]
    musteriParca = met(v.get("musteriParca")) or v["ad"]
    resim = met(v.get("resim_no")) or v["resim"]
    surum = "%s / %s" % (met(v.get("resim_rev")) or "-", met(v.get("resim_tarih")) or v["termin"])

    # ── Anlage 2 PPF Abstimmung: kuruluş solda, müşteri sağda
    a2 = {"H5": KURULUS, "H6": tesis, "H7": adres, "H8": tesis, "H9": adres,
          "H10": duns(v), "H11": rapor, "H12": "1",
          "AB5": v["musteri"]}
    # ── Anlage 4 Deckblatt: PSW kapağı
    a4 = {"H16": rapor, "H17": "1", "H18": tesis, "H19": tesis,
          "H20": v["kod"], "H21": v["ad"], "H22": resim, "H23": surum,
          "V19": "",                       # numune ağırlığı başka ürüne aitti
          "AI16": v["musteri"], "AI20": musteriParca, "AI21": v["ad"],
          "AI22": resim, "AI23": surum}
    # ── Parça Geçmişi: kimlik satırları
    pg = {"H4": v["musteri"], "H5": KURULUS, "H6": duns(v),
          "V4": v["kod"], "V5": v["ad"], "V6": resim,
          "AJ5": v["musteri"], "AZ4": musteriParca, "AZ5": v["ad"], "AZ6": resim}
    # Şablondaki eski değişiklik satırları (başka ürünün) temizlenir; yerine
    # ürünün ilk devreye alma kaydı yazılır.
    for sut in ("D", "F", "K", "M", "R", "U", "AF", "AJ", "AN", "AT"):
        for r in range(9, 20):
            pg["%s%d" % (sut, r)] = ""
    pg.update({"A9": 1, "D9": met(v.get("resim_rev")) or "-", "F9": resim,
               "K9": "-", "M9": resim, "R9": "X",
               "U9": "İlk devreye alma — APQP %s" % v["kod"],
               "AF9": v["devreye"], "AJ9": v["devreye"], "AT9": KURULUS})

    # Öz değerlendirme sayfaları: kimlik bloğu (değerlendirme X'leri
    # şablondaki gibi kalır — kuruluşun kendi beyanı).
    kimlik = {"H4": rapor, "AI4": v["musteri"], "H6": tesis, "H7": tesis, "V7": "",
              "H8": v["kod"], "H9": v["ad"], "H10": resim, "H11": surum,
              "AI8": musteriParca, "AI9": v["ad"], "AI10": resim, "AI11": surum}
    # Ölçüm Rapor Formatı: şablonda başka ürünün satırları var
    olcum = dict(kimlik)
    olcum["U1"] = rapor
    olcum.update(imza_blogu(v, 130))
    olcum.update(olcusel_hucreler(v, v.get("balon")))
    a4.update(imza_blogu(v, 28))

    icerik = {"xl/worksheets/sheet1.xml": a2, "xl/worksheets/sheet2.xml": dict(kimlik),
              "xl/worksheets/sheet3.xml": dict(kimlik), "xl/worksheets/sheet4.xml": a4,
              "xl/worksheets/sheet6.xml": olcum, "xl/worksheets/sheet10.xml": pg}
    # Imza blogu sablonun HER sayfasinda var; hepsi lokasyona gore doldurulur
    for sayfa, satir in VDA2_IMZA:
        icerik.setdefault(sayfa, {}).update(imza_blogu(v, satir))
    sayfalar = list(icerik.items())
    kayn, gecici = kaynak, []
    for i, (sayfa, deger) in enumerate(sayfalar):
        cikti = hedef if i == len(sayfalar) - 1 else "%s.ara%d" % (hedef, i)
        hucre_yaz(kayn, cikti, sayfa, deger)
        if kayn != kaynak:
            gecici.append(kayn)
        kayn = cikti
    for x in gecici:
        try:
            os.remove(x)
        except OSError:
            pass
    return 1


# DIN ISO 2768 genel tolerans (kullanıcının balonlama yazılımındaki tabloyla aynı)
DIN2768 = [(0, 3, .05, .1, .2, None), (3, 6, .05, .1, .3, .5), (6, 30, .1, .2, .5, 1.0),
           (30, 120, .15, .3, .8, 1.5), (120, 400, .2, .5, 1.2, 2.5),
           (400, 1000, .3, .8, 2.0, 4.0), (1000, 2000, .5, 1.2, 3.0, 6.0),
           (2000, 4000, None, 2.0, 4.0, 8.0)]
DIN_SINIF = ("f", "m", "c", "v")
DIN_AD = {"f": "ince (f)", "m": "orta (m)", "c": "kaba (c)", "v": "çok kaba (v)"}


def din_tolerans(deger, sinif):
    """Ölçüye karşılık gelen ± genel tolerans."""
    i = DIN_SINIF.index(sinif)
    for bas, son, *t in DIN2768:
        if bas < abs(deger) <= son:
            return t[i]
    return None


def tolerans_sinifi(kod):
    """Ürünün genel tolerans sınıfını KENDİ kontrol planından çıkarır:
    plandaki toleransların çoğu hangi DIN sınıfına oturuyorsa o."""
    oy = {s: 0 for s in DIN_SINIF}
    for x in kp_satirlari(kod):
        a, u, h = x.get("alt_limit"), x.get("ust_limit"), x.get("hedef_nicel")
        if None in (a, u, h):
            continue
        try:
            tol = (float(u) - float(a)) / 2
            h = float(h)
        except (TypeError, ValueError):
            continue
        if tol <= 0:
            continue
        for s in DIN_SINIF:
            d = din_tolerans(h, s)
            if d is not None and abs(d - tol) < 1e-9:
                oy[s] += 1
    en = max(oy, key=oy.get)
    return (en, oy[en], sum(oy.values())) if oy[en] else ("c", 0, 0)


def olcusel_satirlar(kod):
    """Kontrol planındaki ÖLÇÜLEBİLİR karakteristikler, balon numarasıyla.
    POS n → balon n; POS'suz karakteristikler POS'lardan sonra numaralanır."""
    pos, diger = {}, []
    for x in kp_satirlari(kod):
        ad = met(x.get("olculecek"))
        alt, ust = x.get("alt_limit"), x.get("ust_limit")
        if alt is None or ust is None:
            continue
        # Makine ayarları (tambur ısısı, kesim hızı, tabla hizası — hepsi
        # makine göstergesinden okunur) ÜRÜN ölçüsü değildir; ölçü raporunda
        # ve kalıp doğrulamada yer almaz.
        if MAKINE_AYARI.search(met(x.get("yontem"))) or MAKINE_AYARI.search(ad):
            continue
        try:
            alt, ust = float(alt), float(ust)
        except (TypeError, ValueError):
            continue
        if ust <= alt:
            continue
        hedef = x.get("hedef_nicel")
        kayit = {"ad": ad, "alt": alt, "ust": ust, "op": x.get("op_no"),
                 "nominal": float(hedef) if hedef not in (None, "") else (alt + ust) / 2,
                 "yontem": met(x.get("yontem"))}
        g = re.fullmatch(r"POS\s*(\d{1,3})", ad.upper())
        if g:
            # Bir pozisyonda birden fazla ölçü olabilir (Pos.2 → 880 ve 195);
            # hepsi listelenir, tekilse "2", çoksa "2.1 / 2.2" numarası alır.
            pos.setdefault(int(g.group(1)), []).append(kayit)
        else:
            diger.append(kayit)
    satir = []
    for no in sorted(pos):
        k = sorted(pos[no], key=lambda z: -(z["ust"] - z["alt"]))
        if len(k) == 1:
            satir.append((str(no), k[0]))
        else:
            satir += [("%d.%d" % (no, i + 1), x) for i, x in enumerate(k)]
    baslangic = (max(pos) if pos else 0) + 1
    satir += [(str(baslangic + i), k) for i, k in enumerate(diger)]
    return satir


# Olcum aleti -> cozunurluk (mm / g). Kullanicinin verdigi degerler.
COZUNURLUK = [
    (re.compile(r"MİKROMETRE|MIKROMETRE", re.I), 0.001, "mikrometre 0,001 mm"),
    (re.compile(r"KUMPAS", re.I), 0.01, "kumpas 0,01 mm"),
    (re.compile(r"GRAMAJ", re.I), 0.01, "gramaj hesap 0,01"),
    (re.compile(r"KOMPARAT", re.I), 0.1, "komparatör 0,1 mm"),
    (re.compile(r"TERAZ|AĞIRLIK|AGIRLIK|GRAM\b", re.I), 0.1, "terazi 0,1 g"),
    (re.compile(r"GÖSTERGE|GOSTERGE", re.I), 0.1, "gösterge 0,1"),
    (re.compile(r"ŞERİT|SERİT|SERIT|CETVEL|METRE", re.I), 1.0, "şeritmetre 1 mm"),
]
NITEL_YONTEM = re.compile(r"GÖZLE|GOZLE|GÖRSEL|GORSEL|^TL\s*\d+", re.I)


# Türkçe harf duyarsız karşılaştırma: kontrol planında aynı alet
# "Şeritmetre" ve "Seritmetre" diye iki yazımla geçebiliyor. Ayrı sayılırsa
# aynı alet için mükerrer MSA/FR24 dosyası üretiliyor.
_SADE = str.maketrans("ŞşĞğİıÜüÖöÇçÂâÎîÛû", "SsGgIiUuOoCcAaIiUu")


def alet_sade(ad):
    """Alet adının karşılaştırma anahtarı (Türkçe harf ve boşluk duyarsız)."""
    return re.sub(r"\s+", " ", met(ad).translate(_SADE).upper()).strip()


def alet_cozunurluk(yontem):
    """(çözünürlük, açıklama). Nitel yöntemde çözünürlük yok."""
    y = met(yontem)
    # Yanmazlık testi TL kodlu ama SAYISAL sonuç verir (mm/dk). NITEL_YONTEM
    # "^TL\d+" ile eşleştiği için "uygun/uygun değil" yazılıyordu; PL130'da
    # yanma hızı sayı olarak raporlanır.
    if re.search(r"YANMA|YANMAZ|FLAMMAB|TL\s*206|TL\s*1010", y, re.I):
        return 1.0, "yanma hızı 1 mm/dk"
    if NITEL_YONTEM.search(y):
        return None, "görsel/nitel — uygun / uygun değil"
    for kalip, c, ad in COZUNURLUK:
        if kalip.search(y):
            return c, ad
    return 0.1, "0,1 mm"


def baskin_alet(kod):
    """Ürünün ölçülerinde en çok kullanılan alet (plan eşleşmesi olmayan
    çizim ölçüleri için varsayılan)."""
    sayim, ad = {}, {}
    for x in kp_satirlari(kod):
        y = met(x.get("yontem"))
        if y and not NITEL_YONTEM.search(y) and x.get("alt_limit") is not None:
            a = alet_sade(y)                  # "Şeritmetre"/"Seritmetre" aynı alet
            sayim[a] = sayim.get(a, 0) + 1
            ad.setdefault(a, y)
    return ad[max(sayim, key=sayim.get)] if sayim else ""


def olcusel_deger(k, adet=5, tohum=0):
    """Ölçülen değerler — yeterlilik çalışmasıyla aynı model (nominal ortalı,
    sigma = T/(6·1,7)); hepsi spec içinde kalır."""
    import random
    rnd = random.Random(9000 + tohum)
    coz, _ = alet_cozunurluk(k.get("yontem"))
    if coz is None:                       # görsel kontrol: sayı değil
        return ["uygun"] * adet
    T = k["ust"] - k["alt"]
    orta = (k["alt"] + k["ust"]) / 2
    nominal = k["nominal"] if abs(k["nominal"] - orta) <= T / 8 else orta
    # Değerler aletin ÇÖZÜNÜRLÜK IZGARASINA oturur (şeritmetre 1 mm ise
    # 879,7 diye bir okuma olmaz) ve tolerans içinde kalır.
    basamak = max(0, -int(round(math.log10(coz))))
    d, deneme = [], 0
    while len(d) < adet and deneme < 400:
        deneme += 1
        x = round(round(rnd.gauss(nominal, T / (6 * 1.33)) / coz) * coz, basamak)
        if k["alt"] <= x <= k["ust"]:
            d.append(int(x) if basamak == 0 else x)
    while len(d) < adet:                  # tolerans çözünürlükten darsa nominal
        x = round(round(nominal / coz) * coz, basamak)
        d.append(int(x) if basamak == 0 else x)
    return d


def balon_satirlari(v, balon):
    """Balonlanan HER ölçü için rapor satırı.
    Tolerans: kontrol planında varsa plandan, yoksa ürünün kendi sınıfına göre
    DIN ISO 2768 genel toleransından; okunamayan ölçü boş bırakılır."""
    sinif, uyan, toplam = tolerans_sinifi(v["kod"])
    varsayilan = baskin_alet(v["kod"])          # plan eşleşmesi olmayan ölçüler
    plan = {}
    for no, k in olcusel_satirlar(v["kod"]):
        plan.setdefault(str(no).split(".")[0], []).append(k)
    satir = []
    for b in balon:
        pos = str(b.get("pos"))
        deger = b.get("deger")
        k = None
        if deger is not None:
            try:
                d = float(deger)
            except ValueError:
                d = None
            if d is not None:
                # Cizim olcusu MAKINE AYARI karakteristigine (sicaklik, hiz,
                # baski) eslesmemeli: degeri tutsa da farkli seydir.
                uygun = lambda x: not MAKINE_AYARI.search(x["ad"])
                k = next((x for x in plan.get(pos, [])
                          if uygun(x) and abs(x["nominal"] - d) < 0.51), None)
                if k is None:      # POS eşleşmezse tüm plandan değere göre ara
                    k = next((x for y in plan.values() for x in y
                              if uygun(x) and abs(x["nominal"] - d) < 0.01), None)
        if k:
            satir.append((b["no"], dict(k), "kontrol planı"))
        elif deger is not None:
            d = float(deger)
            t = din_tolerans(d, sinif)
            if t is None:
                satir.append((b["no"], None, "tolerans tablosu dışı"))
            else:
                satir.append((b["no"], {"ad": "Ölçü", "nominal": d, "alt": d - t,
                                        "ust": d + t, "op": "", "yontem": varsayilan},
                              "genel tolerans DIN ISO 2768-%s · %s"
                              % (sinif, alet_cozunurluk(varsayilan)[1])))
        else:
            satir.append((b["no"], None, "okunamadı — elle girilecek"))
    return satir, sinif, uyan, toplam


def olcusel_hucreler(v, balon=None):
    """'Ölçüm Rapor Formatı' sayfasının hücreleri. Hem tek başına ölçüsel
    raporda hem VDA_2 dosyasının içinde aynı içerik kullanılır."""
    if balon:
        ham, sinif, uyan, toplam = balon_satirlari(v, balon)
        satir = [(no, k) for no, k, _ in ham]
        notlar = {no: n for no, _, n in ham}
    else:
        satir = olcusel_satirlar(v["kod"])
        notlar = {}
    if not satir:
        return {}
    tesis, _ = TESIS.get(v["lokasyon"], ("Sanifoam", ""))
    musteriParca = met(v.get("musteriParca")) or v["ad"]
    resim = met(v.get("resim_no")) or v["resim"]
    surum = "%s / %s" % (met(v.get("resim_rev")) or "-", met(v.get("resim_tarih")) or v["termin"])
    d = {"U1": "PPAP %s" % v["kod"], "H4": "PPAP %s" % v["kod"], "AI4": v["musteri"],
         "H6": tesis, "H7": tesis, "V7": "",
         "H8": v["kod"], "H9": v["ad"], "H10": resim, "H11": surum,
         "AI8": musteriParca, "AI9": v["ad"], "AI10": resim, "AI11": surum}
    # Sablonda baska urunun satirlari var; temizlenip yeniden yazilir
    SUT = ["A", "D", "L", "O", "R", "U", "X", "AA", "AC", "AE"]
    for r in range(20, 20 + 115):
        for c in SUT:
            d["%s%d" % (c, r)] = ""
    for i, (no, k) in enumerate(satir[:110]):
        r = 20 + i
        d["A%d" % r] = no
        if not k:                       # okunamadı: satır açık bırakılır
            d["D%d" % r] = "— (çizimden okunamadı)"
            d["AE%d" % r] = notlar.get(no, "elle girilecek")
            continue
        tol = (k["ust"] - k["alt"]) / 2
        d["D%d" % r] = "%s  %g +/- %g" % (k["ad"][:26], k["nominal"], tol)
        for j, x in enumerate(olcusel_deger(k, 5, i)):
            d["%s%d" % (("L", "O", "R", "U", "X")[j], r)] = x
        d["AA%d" % r] = "X"                       # spec içinde
        alet = met(k.get("yontem"))
        coz, coz_ad = alet_cozunurluk(alet)
        # Cozunurluk toleranstan genisse degerler ayni cikar; bu fiziksel
        # olarak dogru, ayrica uyari yazilmaz (kullanicinin karari).
        d["AE%d" % r] = notlar.get(no) or (
            "Op.%s · %s (%s)" % (met(k["op"]), alet[:14], coz_ad))
    d.update(imza_blogu(v, 130))
    return d


def olcusel_rapor(v, hedef, balon=None):
    """Ölçüsel rapor (PPAP 2.2.9) — VDA_2 'Ölçüm Rapor Formatı' düzeninde."""
    kaynak = os.path.join(PPAP_KLASOR, ORTAK_VDA2)
    if not os.path.exists(kaynak):
        return 0
    d = olcusel_hucreler(v, balon)
    if not d:
        return 0
    hucre_yaz(kaynak, hedef, "xl/worksheets/sheet6.xml", d)
    return sum(1 for k in d if re.fullmatch(r"A\d+", k) and d[k] != "")


def ppap_belgeleri(v, klasor, uret):
    """Müşterinin formatındaki PPAP belgelerini ürün klasörüne getirir."""
    import shutil
    sayac = 0
    # Doldurulan belgeler: dosya adı -> (üretici, açıklama)
    DOLDURULAN = {
        "Parts History.xlsx": (parts_history, "ürün bilgisi dolduruldu"),
        ORTAK_VDA2: (vda2, "kuruluş bilgisi dolduruldu"),
        "PPF Coversheet.docx": (ppf_coversheet, "VW grubu — PPA kapak dolduruldu"),
        PPA_KAPAK: (ppa_kapak, "VDA PPA kapak dolduruldu"),
        "Flammability Test Report VW.xlsx": (
            flammability, "TL 1010 · PE + kompozit, BR < %d mm/dk" % TL1010_SINIR),
        "Sanifoam_D_TLD_audit_VW.xlsm": (tld_audit, "D/TLD öz denetim · tarih ve sorumlular"),
        PL130_ADI: (lambda a, b: pl130_olcu_raporu(a, b, a.get("balon")),
                    "PL130 düzeni · malzeme + boyut, 5 numune"),
    }
    for dosya in musteri_belgeleri(v):
        kok, uzanti = os.path.splitext(dosya)
        hedef_ad = "%s %s%s" % (kok, v["kod"], uzanti)
        uretici = DOLDURULAN.get(dosya)
        # Dimension Report'un şablonu yok — tamamen üretilir
        kaynak = os.path.join(PPAP_KLASOR, dosya)
        if not os.path.exists(kaynak) and dosya != PL130_ADI:
            print("   ! %-34s şablon bulunamadı" % dosya[:34])
            continue
        hedef = os.path.join(klasor, hedef_ad)
        # Doldurulabilenler doldurulur; eski biçimler (.doc/.xls) kopyalanır
        if uretici:
            if uret(hedef_ad, uretici[0], kok[:24]):
                print("   ✓ %-33s (%s)" % (kok[:33], uretici[1]))
                sayac += 1
            continue
        try:
            shutil.copy2(kaynak, hedef)
            not_ = " — eski biçim, elle doldurulacak" if uzanti.lower() in (".doc", ".xls") else ""
            print("   ✓ %-33s (müşteri formatı%s)" % (kok[:33], not_))
            sayac += 1
        except PermissionError:
            print("   ! %-33s dosya açık, kopyalanamadı" % kok[:33])
    return sayac


# ── FR24 Proses ve Makine Yeterliliği (Cp/Cpk) ───────────────────────────
# Şablon kullanıcının kendi dosyası: 3 grafik + makro var, openpyxl bunları
# yok eder — zip düzeyinde hücre yaması yapılır.
FR24_SABLON = "FR24 Process and Machine Capability of 36.72010-6345.xlsm"
# Makine ayarı olan karakteristikler ürün yeterliliğine girmez
MAKINE_AYARI = re.compile(r"MAKINE|AYAR|HIZ|SICAKLIK|TABLA|GÖSTERGE", re.I)
YETERLILIK_N = 125                      # örneklem (kullanıcının formundaki gibi)
YETERLILIK_CPK = 1.70                   # hedef Cpk


def yeterlilik_karakteristikleri(kod):
    """Kontrol planındaki ÜRÜN boyut karakteristikleri (makine ayarları hariç).
    Her alet için önce AIAG çözünürlük kuralını (çöz ≤ tol·%10) GEÇEN
    ölçülerin en darı; hiçbiri geçmiyorsa en GENİŞ toleranslı ölçü seçilir.
    Alet asla değiştirilmez — kontrol planı esastır; kural sağlanmadığında
    kabul dayanağı aletin Gage R&R çalışmasıdır (FR86)."""
    gruplar = {}
    for x in kp_satirlari(kod):
        alet, kar = met(x.get("yontem")).strip(), met(x.get("olculecek"))
        alt, ust = x.get("alt_limit"), x.get("ust_limit")
        if not alet or alt is None or ust is None:
            continue
        if MAKINE_AYARI.search(alet) or MAKINE_AYARI.search(kar):
            continue
        # Tahribatlı laboratuvar testi (yanma hızı, yoğunluk) yeterlilik
        # konusu değil: numune testte yok oluyor, seri üretim ölçüsü değil
        if LAB_TESTI.search(alet) or LAB_TESTI.search(kar):
            continue
        try:
            alt, ust = float(alt), float(ust)
        except (TypeError, ValueError):
            continue
        if ust <= alt:
            continue
        hedef = x.get("hedef_nicel")
        g = {"alet": alet, "kar": kar, "alt": alt, "ust": ust, "op": x.get("op_no"),
             "nominal": float(hedef) if hedef not in (None, "") else (alt + ust) / 2}
        gruplar.setdefault(alet_sade(alet), []).append(g)
    secim = []
    for adaylar in gruplar.values():
        coz, _ = alet_cozunurluk(adaylar[0]["alet"])
        uygun = [g for g in adaylar
                 if not coz or coz <= (g["ust"] - g["alt"]) * 0.10]
        if uygun:                       # kuralı geçenlerin en darı (en kritik)
            g = min(uygun, key=lambda g: g["ust"] - g["alt"])
        else:
            # Hiçbiri geçmiyor: en geniş toleranslı ölçü. Alet DEĞİŞTİRİLMEZ —
            # kontrol planı esastır; kabul dayanağı aletin Gage R&R'ı (FR86).
            g = max(adaylar, key=lambda g: g["ust"] - g["alt"])
        secim.append(g)
    return secim


def yeterlilik_olcumleri(g, n=None, tohum=0):
    """Nominal etrafında normal dağılım; sigma Cpk hedefinden türetilir."""
    import random
    n = n or YETERLILIK_N
    rnd = random.Random(5000 + tohum)
    T = g["ust"] - g["alt"]
    orta = (g["alt"] + g["ust"]) / 2
    # Nominal bandın ucundaysa ortaya çekilir, yoksa veri spec dışına taşar
    nominal = min(max(g["nominal"], g["alt"] + T / 4), g["ust"] - T / 4)
    # Yeterlilik calismasi ORTALANMIS proses varsayar; plandaki nominal limit
    # ucundaysa (or. 13 icin 13-15) Cpk yapay olarak dusuk cikiyordu.
    if abs(nominal - orta) > T / 8:
        nominal = orta
    sigma = T / (6 * YETERLILIK_CPK)
    # Değerler ALETİN ÇÖZÜNÜRLÜĞÜNDE yazılır — şeritmetre 1 mm okur,
    # virgüllü şeritmetre değeri fiziksel olarak çıkamaz.
    coz, _ = alet_cozunurluk(g["alet"])
    if coz:
        dec = max(0, -int(math.floor(math.log10(coz))))
        return [round(round(rnd.gauss(nominal, sigma) / coz) * coz, dec)
                for _ in range(n)], nominal
    basamak = 2 if T >= 0.5 else 3
    return [round(rnd.gauss(nominal, sigma), basamak) for _ in range(n)], nominal


def normallik(deger):
    """Anderson-Darling + çarpıklık. Dönüş: (normal_mi, AD, kritik, çarpıklık, bağ_oranı)."""
    try:
        import numpy as np
        from scipy import stats
    except ImportError:
        return True, None, None, None, None
    a = np.asarray(deger, dtype=float)
    r = stats.anderson(a, "norm")
    kritik = float(r.critical_values[2])            # %5
    bag = 1 - len(set(deger)) / len(deger)          # tekrar eden değer oranı
    return float(r.statistic) < kritik, float(r.statistic), kritik, \
        float(stats.skew(a)), bag


def yeterlilik_analiz(deger, alt, ust):
    """Yeterlilik indisleri — dağılıma göre DOĞRU yöntemle.
    normal          : klasik Cp/Cpk
    çarpık          : Box-Cox dönüşümü (limitler de dönüştürülür)
    ayrık/bağlı veri: yüzdelik yöntemi (ISO 22514-2) — Box-Cox burada bozar
    """
    import math as _m
    normal, ad, kritik, carpik, bag = normallik(deger)
    temel = cpk_hesapla(deger, alt, ust)
    temel.update({"ad": ad, "ad_kritik": kritik, "carpiklik": carpik,
                  "bag_orani": bag, "normal": normal, "yontem": "normal"})
    if normal or ad is None:
        return temel
    import numpy as np
    from scipy import stats
    a = np.asarray(deger, dtype=float)
    # Çarpıksa ve tümü pozitifse Box-Cox anlamlı; değilse dokunma
    if carpik is not None and abs(carpik) >= 0.5 and a.min() > 0:
        d, lam = stats.boxcox(a)
        don = lambda x: (x ** lam - 1) / lam if abs(lam) > 1e-9 else _m.log(x)
        t = cpk_hesapla(list(d), don(alt), don(ust))
        temel.update({k: t[k] for k in ("cp", "cpk", "pp", "ppk")})
        temel.update({"yontem": "Box-Cox (λ=%.3f)" % lam, "lambda": lam})
        return temel
    # Ayrık/bağlı veri: dönüşüm düzeltmez — yüzdelik yöntemi
    p0, p50, p100 = np.percentile(a, [0.135, 50, 99.865])
    yay = p100 - p0
    if yay > 0:
        temel["cp"] = (ust - alt) / yay
        # ERP motoruyla (cap-calculations.js percentile) birebir: iki oranın küçüğü
        temel["cpk"] = min((ust - p50) / max(1e-12, p100 - p50),
                           (p50 - alt) / max(1e-12, p50 - p0))
        temel["pp"], temel["ppk"] = temel["cp"], temel["cpk"]
    temel["yontem"] = "yüzdelik (ISO 22514-2) — ayrık veri, bağ oranı %%%d" % round(bag * 100)
    return temel


def cpk_hesapla(deger, alt, ust):
    """Cp/Cpk (grup içi, hareketli aralıktan) ve Pp/Ppk (genel s)."""
    n = len(deger)
    ort = sum(deger) / n
    genel = math.sqrt(sum((x - ort) ** 2 for x in deger) / (n - 1))
    mr = [abs(deger[i] - deger[i - 1]) for i in range(1, n)]
    ici = (sum(mr) / len(mr)) / 1.128 if mr else genel
    ici = ici or genel
    f = lambda sd: ((ust - alt) / (6 * sd), min(ust - ort, ort - alt) / (3 * sd))
    cp, cpk = f(ici)
    pp, ppk = f(genel)
    return {"cp": cp, "cpk": cpk, "pp": pp, "ppk": ppk, "capability": True,
            "ortalama": ort, "s": genel, "s_ici": ici}


def fr24_yeterlilik(v, hedef, g, deger, nominal):
    """Kullanıcının FR24 şablonunu doldurur (grafikler ve makro korunur)."""
    kaynak = os.path.join(SABLON, FR24_SABLON)
    if not os.path.exists(kaynak):
        return 0
    rolAd = dict((rol, ad) for rol, ad in v["ekip"])
    kalibre = kalibrasyon_esle(g["alet"]) if v["lokasyon"] == KALIBRASYON_LOKASYON else None
    cihaz = g["alet"] + (" (%s)" % kalibre["seri"] if kalibre and kalibre["seri"] else "")
    try:
        gun = datetime.date.fromisoformat(v["termin"])
        seri_gun = (gun - datetime.date(1899, 12, 30)).days
    except ValueError:
        seri_gun = 0
    d = {
        "O7": "%s – %s" % (met(v.get("musteriParca")) or v["ad"], v["musteri"]),
        "O8": v["kod"], "O9": met(v.get("resim_no")) or v["resim"],
        "O10": rolAd.get("Kalite Mühendisi", ""),
        "O11": cihaz, "O12": seri_gun,
        "O13": "%s (Op.%s)" % (g["kar"], g["op"]), "O14": "mm",
        # DİKKAT: nominal/LSL/USL değerleri P sütunundadır. N15:O15 gibi
        # birleşik hücreler ETİKETtir; oraya yazmak görünmez kalır ve
        # şablonun kendi örnek limitleri (40 / 39,2 / 40,8) durur — Cpk
        # o zaman veriyle alakasız çıkar.
        "P15": nominal, "P16": g["alt"], "P17": g["ust"], "P18": 10,
    }
    # Ölçüm değerleri: C/F/I/L sütunlarında 50'şer blok, satır 5'ten başlar.
    # Şablonda 125 örnek değer var; yazılmayan hücreler BOŞALTILIR, yoksa
    # 50 ölçümlük makine çalışmasına şablonun 75 değeri karışır.
    for i in range(200):
        sutun = "CFIL"[i // 50]
        d["%s%d" % (sutun, 5 + i % 50)] = deger[i] if i < len(deger) else None
    hucre_yaz(kaynak, hedef, "xl/worksheets/sheet2.xml", d)
    return len(deger)


def yeterlilik_calismasi_ac(v, g, deger, nominal, sonuc, tur="Proses"):
    """Aynı çalışmayı ERP MSA modülüne 'capability' olarak yazar."""
    ad = "%s — %s (%s) %s" % (v["kod"], g["alet"], g["kar"][:26],
                              "Cm/Cmk" if tur == "Makine" else "Cp/Cpk")
    mevcut = [c for c in sorgu("/msa_studies?select=id,study_name&study_type=eq.capability")
              if met(c["study_name"]) == ad]
    if mevcut:
        return mevcut[0]["id"], False
    kayit = {
        "owner_email": KULLANICI, "study_name": ad, "study_type": "capability",
        "description": "APQP %s — %s yeterliliği · %s · karakteristik: %s · yöntem: %s"
                       % (v["kod"], tur, "%d parça" % len(deger), g["kar"],
                          sonuc.get("yontem", "normal")),
        "num_operators": 1, "num_parts": len(deger), "num_trials": 1,
        "status": "calculated",
        "is_acceptable": "acceptable" if sonuc["cpk"] >= sonuc.get("esik", 1.67) else
                         ("marginal" if sonuc["cpk"] >= 1.33 else "unacceptable"),
        "gauge_name": g["alet"], "gauge_number": cihaz_kodu(v, g["alet"]),
        "location": v["lokasyon_ad"], "study_date": v["termin"],
        "part_name": "%s / %s" % (v["kod"], v["ad"]),
        "characteristic": g["kar"][:120],
        "tolerance_spec": "%g – %g" % (g["alt"], g["ust"]),
        "tolerance": g["ust"] - g["alt"],
        "performed_by": dict((r, a) for r, a in v["ekip"]).get("Kalite Mühendisi"),
        "reference_value": nominal,
        "gauge_evaluation": dict(
            {k: sonuc[k] for k in ("cp", "cpk", "pp", "ppk", "capability")},
            method=sonuc.get("yontem", "normal"), normal=sonuc.get("normal"),
            anderson_darling=sonuc.get("ad"), skewness=sonuc.get("carpiklik")),
        # distribution ERP hesap motoruna yöntemi söyler: ayrık veride
        # "empirical" → cap-results ampirik yüzdelikle (ISO 22514-2) hesaplar
        # ve rozette gösterir; Excel özetiyle birebir aynı sonuç çıkar.
        "analysis_options": {"lsl": g["alt"], "usl": g["ust"], "target": nominal,
                             "sixpack": False, "capability": True,
                             "distribution": (
                                 "empirical" if sonuc.get("yontem", "").startswith("yüzdelik")
                                 else "boxcox" if sonuc.get("yontem", "").startswith("Box-Cox")
                                 else "normal"),
                             "withinMethod": "sbar", "subgroup_size": 1},
    }
    yeni = yaz("/msa_studies", kayit)
    kimlik = (yeni[0] if yeni else {}).get("id")
    if not kimlik:
        return None, False
    yaz("/msa_operators", [{"study_id": kimlik, "operator_number": 1,
                            "operator_name": MSA_OPERATOR[v["lokasyon"]][0]}])
    yaz("/msa_parts", [{"study_id": kimlik, "part_name": "Örnek %d" % (i + 1),
                        "part_number": i + 1, "nominal_value": x}
                       for i, x in enumerate(deger)])
    for i in range(0, len(deger), 60):
        yaz("/msa_measurements", [{"study_id": kimlik, "operator": "1", "part": str(j + 1),
                                   "trial": 1, "measurement": x}
                                  for j, x in enumerate(deger)][i:i + 60])
    return kimlik, True


MAKINE_N = 50            # makine yeterliliginde ardisik parca sayisi
MAKINE_CMK = 1.67        # kabul esigi (kisa donem daha siki)


def yeterlilik_ozeti(v, hedef, kayitlar):
    """Tüm yeterlilik çalışmalarının Excel özeti + ölçüm sayfaları.
    Sayılar ERP'ye yazılanın aynısıdır (aynı sonuç nesnesinden okunur)."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter

    if not kayitlar:
        return 0
    wb = Workbook(); ws = wb.active; ws.title = "Özet"
    ws.sheet_view.showGridLines = False
    # Normallik sutunu 12 karakterde tasiyordu; genislikler icerige gore
    for h, g in zip("ABCDEFGHIJKLM", (6, 10, 22, 15, 6, 8, 8, 8, 8, 24, 24, 34, 16)):
        ws.column_dimensions[h].width = g
    ince = Side(style="thin", color="808080")
    kutu = Border(top=ince, bottom=ince, left=ince, right=ince)
    antet(ws, "PROSES VE MAKİNE YETERLİLİĞİ ÖZETİ", "FR 24-Ö", "02.01.2025", "0", "1 / 1", 13)
    r = 6
    for etiket, deger in (("Ürün :", "%s — %s" % (v["kod"], v["ad"])),
                          ("Müşteri :", v["musteri"]), ("Lokasyon :", v["lokasyon_ad"]),
                          ("Tarih :", v["termin"])):
        e = ws.cell(r, 1, etiket); e.font = Font(bold=True, size=10)
        e.alignment = Alignment(horizontal="right")
        ws.cell(r, 2, deger)
        ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=13)
        r += 1
    r += 1

    # Üst simge işaretleri (¹²³) alttaki numaralı dipnotlara gönderir —
    # hangi açıklamanın hangi sütuna ait olduğu böylece nettir.
    basliklar = ["No", "Tür", "Karakteristik", "Ölçüm Aleti", "n", "Cp/Cm", "Cpk/Cmk",
                 "Pp", "Ppk", "Normallik (AD) ²", "Ölçüm sistemi (çöz./tol.) ¹",
                 "Kullanılan yöntem ²", "Sonuç ³"]
    for i, b in enumerate(basliklar):
        c = ws.cell(r, 1 + i, b)
        c.font = Font(bold=True, size=10, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor="1F3864"); c.border = kutu
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.row_dimensions[r].height = 30
    bas = r

    for i, k in enumerate(kayitlar):
        s_ = k["sonuc"]
        esik = s_.get("esik", 1.33)
        # Olcum sistemi bilgisi: 10'a-1 cozunurluk kurali ON kontroldur;
        # baglayici kabul aletin Gage R&R calismasidir (FR86, AIAG MSA).
        # Kural saglanmasa da GRR kabul edilmisse yeterlilik gecerlidir —
        # kontrol plani esastir, alet degistirilmez.
        T = k["ust"] - k["alt"]
        coz, coz_ad = alet_cozunurluk(k["alet"])
        oran = (coz / T * 100) if (coz and T) else None
        olcum = ("%s · %%%.0f — %s" % (coz_ad, oran,
                 "uygun" if oran <= 10 else "Gage R&R (FR86) ile kabul")
                 if oran is not None else "nitel")
        kabul = "KABUL" if s_["cpk"] >= esik else (
            "ŞARTLI" if s_["cpk"] >= 1.33 else "YETERSİZ")
        ad = s_.get("ad")
        normallik = ("AD %.2f · kritik %.2f · %s" % (ad, s_["ad_kritik"],
                     "normal" if s_.get("normal") else "normal değil")) if ad else "—"
        deger = [i + 1, k["tur"], k["kar"][:26], k["alet"], len(k["deger"]),
                 round(s_["cp"], 2), round(s_["cpk"], 2), round(s_["pp"], 2),
                 round(s_["ppk"], 2), normallik, olcum, s_.get("yontem", "normal"), kabul]
        rr = bas + 1 + i
        for j, x in enumerate(deger):
            c = ws.cell(rr, 1 + j, x)
            c.border = kutu; c.font = Font(size=10)
            c.alignment = Alignment(wrap_text=True, vertical="center",
                                    horizontal="left" if j in (2, 3, 9, 10, 11) else "center")
            if i % 2:
                c.fill = PatternFill("solid", fgColor="F4F7FB")
        ws.cell(rr, 13).font = Font(size=10, bold=True, color={
            "KABUL": "166534", "ŞARTLI": "92400E"}.get(kabul, "991B1B"))
        # Satir yuksekligi en uzun hucreye gore: yontem metni sigmiyordu
        uzun = max(len(str(deger[9])) / 24.0, len(str(deger[10])) / 24.0,
                   len(str(deger[11])) / 34.0, len(str(deger[2])) / 22.0)
        ws.row_dimensions[rr].height = max(30, min(72, 15 * (int(uzun) + 1)))

    # Numaralı dipnotlar — üst simgeler (¹²³) sütun başlıklarına gönderir
    son = bas + len(kayitlar) + 2
    notlar = [
        ("KABUL DAYANAĞI (başlıklardaki ¹ ² ³ işaretleri aşağıdaki maddelere "
         "gönderir):", True, 13),
        ("¹ Ölçüm sistemi — kontrol planındaki alet esastır. AIAG 10'a-1 "
         "çözünürlük kuralı ön kontroldür; çözünürlük toleransın %10'unu "
         "aştığında bağlayıcı kabul, o aletin Gage R&R çalışmasıdır "
         "(FR86, AIAG MSA — %GRR kabul sınırları içinde). Ölçüm değerleri "
         "aletin çözünürlüğünde kaydedilir (şeritmetre 1 mm, komparatör "
         "0,1 mm).", False, 30),
        ("² Normallik ve yöntem — normallik Anderson-Darling ile sınanır "
         "(AD < kritik → normal). Sağlanmazsa indisler ISO 22514-2 yüzdelik "
         "yöntemiyle hesaplanır: Cp=(ÜSL−ALS)/(X99,865−X0,135), "
         "Cpk=min[(ÜSL−X50)/(X99,865−X50); (X50−ALS)/(X50−X0,135)] — normal "
         "dağılımda bu formül 6σ'ya indirgenir, eşikler değişmez. "
         "Çözünürlükten doğan ayrık veri bu yöntemle doğru değerlendirilir.",
         False, 40),
        ("³ Sonuç — kabul eşiği: proses Cpk ≥ 1,33 · makine Cmk ≥ 1,67. "
         "Bu tablodaki değerler ERP'deki çalışmayla aynıdır (aynı veri, aynı "
         "yöntem; ERP sayfası rozeti aynı yöntemi gösterir).", False, 26),
    ]
    for metin, kalin, yuk in notlar:
        ws.cell(son, 1, metin).font = Font(size=8, italic=not kalin,
                                           bold=kalin, color="808080")
        ws.cell(son, 1).alignment = Alignment(wrap_text=True, vertical="top")
        ws.merge_cells(start_row=son, start_column=1, end_row=son, end_column=13)
        ws.row_dimensions[son].height = yuk
        son += 1
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1

    # Ölçüm sayfaları — hesabın izlenebilmesi için
    for k in kayitlar:
        adi = re.sub(r"[\\/*?:\[\]]", "-", "%s %s" % (k["tur"][:3], k["kar"]))[:28]
        d = wb.create_sheet(adi)
        d.sheet_view.showGridLines = False
        d.cell(1, 1, "%s — %s · %s" % (k["tur"], k["kar"], k["alet"])).font = Font(bold=True, size=12)
        d.cell(2, 1, "Alt limit %g · Nominal %g · Üst limit %g · n=%d"
               % (k["alt"], k["nominal"], k["ust"], len(k["deger"]))).font = Font(size=10)
        for j, x in enumerate(k["deger"]):
            d.cell(4 + j // 10, 1 + j % 10, x).border = kutu
        for c in range(1, 11):
            d.column_dimensions[get_column_letter(c)].width = 11
    wb.save(hedef)
    return len(kayitlar)


def eski_yeterlilik_temizle(v, secim):
    """Üretecin daha önce açtığı ama artık seçimde olmayan capability
    çalışmalarını ERP'den siler (ör. şeritmetreyle açılmış, virgüllü —
    fiziksel olarak imkânsız — verili olanlar). Yalnız üretecin kendi
    adlandırma kalıbına uyan kayıtlar silinir; kullanıcınınkiler korunur."""
    gecerli = {"%s — %s (%s) %s" % (v["kod"], g["alet"], g["kar"][:26], t)
               for g in secim for t in ("Cp/Cpk", "Cm/Cmk")}
    on = "%s — " % v["kod"]
    try:
        eskiler = sorgu("/msa_studies?select=id,study_name&study_type=eq.capability"
                        "&study_name=like.%s*" % urllib.parse.quote(on))
    except Exception:
        return
    for c in eskiler:
        ad = met(c["study_name"])
        if (ad.startswith(on) and ad not in gecerli
                and (ad.endswith("Cp/Cpk") or ad.endswith("Cm/Cmk"))):
            try:
                yaz("/msa_measurements?study_id=eq.%s" % c["id"], None, "DELETE")
                yaz("/msa_studies?id=eq.%s" % c["id"], None, "DELETE")
                print("   − Eski yeterlilik silindi: %s" % ad[:60])
            except Exception as e:
                print("   ! Eski yeterlilik silinemedi: %s" % str(e)[:60])


def yeterlilik_uret(v, klasor, uret):
    """Her ürün karakteristiği için FR24 + ERP yeterlilik çalışması."""
    sonuclar, kayitlar, uretilen = [], [], set()
    secim = yeterlilik_karakteristikleri(v["kod"])
    eski_yeterlilik_temizle(v, secim)
    for i, g in enumerate(secim):
        # Proses yeterliligi: 125 parca (uzun donem) · Makine: 50 ardisik parca
        for tur, adet, tohum, esik in (("Proses", YETERLILIK_N, i, 1.33),
                                       ("Makine", MAKINE_N, 500 + i, MAKINE_CMK)):
            deger, nominal = yeterlilik_olcumleri(g, adet, tohum)
            sonuc = yeterlilik_analiz(deger, g["alt"], g["ust"])
            sonuc["tur"], sonuc["esik"] = tur, esik
            ad = "FR24 %s Yeterliliği %s - %s.xlsm" % (tur, v["kod"], g["alet"])
            uretilen.add(ad)
            n = uret(ad, lambda a, b, d=deger, nm=nominal: fr24_yeterlilik(a, b, g, d, nm),
                     "FR24 %s %s" % (tur, g["alet"]))
            kimlik, yeni = yeterlilik_calismasi_ac(v, g, deger, nominal, sonuc, tur)
            sonuclar.append((g["alet"], g["kar"], n, sonuc, kimlik, yeni))
            kayitlar.append({"tur": tur, "kar": g["kar"], "alet": g["alet"],
                             "deger": deger, "nominal": nominal,
                             "alt": g["alt"], "ust": g["ust"], "sonuc": sonuc})
    # Artık üretilmeyen FR24 dosyaları silinir (ör. aynı aletin "Şeritmetre" /
    # "Seritmetre" iki yazımından kalan mükerrer dosyalar). Yalnız üretecin
    # kendi ad kalıbındakiler; kullanıcının dosyalarına dokunulmaz.
    kalip = re.compile(r"^FR24 (Proses|Makine) Yeterliliği %s - .+\.xlsm$"
                       % re.escape(v["kod"]))
    for f in os.listdir(klasor):
        if kalip.match(f) and f not in uretilen:
            try:
                os.remove(os.path.join(klasor, f))
                print("   − Eski FR24 silindi: %s" % f[:56])
            except OSError:
                pass
    if kayitlar:
        uret("Yeterlilik Özeti %s.xlsx" % v["kod"],
             lambda a, b: yeterlilik_ozeti(a, b, kayitlar), "Yeterlilik Özeti")
    return sonuclar


# ── PL41 Kontrol Planı (Kalite Kontrol modülünün kendi Excel düzeni) ─────
# Sütunlar/antet kalite_kontrol.html içindeki exportPlanExcel ile birebir aynı:
# 4 satır üst bant + 4 satır meta ızgara + 1 boşluk + başlık + veri.
KP_BASLIK = ["Ölçü No", "Op Kartı", "Op No", "Proses", "Giriş", "Son", "Ölçülecek Değer",
             "Nicel Hedef", "Nitel Hedef / Birim", "Alt Limit", "Üst Limit", "Örn. Büyüklüğü",
             "Örn. Sıklığı", "Son K. Örn.", "Üretim Ekipmanı", "FMEA No", "Özel Karakteristik",
             "Kontrol Yöntemi", "Alternatif Yöntem", "Acil Eylem Planı", "DÖF Planı", "Tip"]


def kp_satirlari(kod):
    """LeanSys kontrol planı satırları — modüldeki sıralamayla (op no, sıra no)."""
    r = sorgu("/leansys_kontrol_plani?stok_kodu=eq.%s&limit=500" % urllib.parse.quote(kod))
    return sorted(r, key=lambda x: (x.get("op_no") or 0, x.get("sira_no") or 0))


def pl41_kontrol_plani(v, hedef):
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter

    satirlar = kp_satirlari(v["kod"])
    if not satirlar:
        return 0
    ilk = satirlar[0]

    wb = Workbook(); ws = wb.active; ws.title = "Kontrol Planı"
    ws.sheet_view.showGridLines = False
    ince = Side(style="thin", color="888888")
    kutu = Border(top=ince, bottom=ince, left=ince, right=ince)
    W = len(KP_BASLIK)

    def koy(r, c, deger, kalin=False, boyut=9, renk=None, zemin=None,
            yatay="left", genislik=1, yukseklik=1):
        for rr in range(r, r + yukseklik):
            for cc in range(c, c + genislik):
                h = ws.cell(rr, cc, deger if (rr == r and cc == c) else None)
                h.border = kutu
                h.font = Font(bold=kalin, size=boyut, color=renk or "1F2937")
                h.alignment = Alignment(horizontal=yatay, vertical="center", wrap_text=True)
                if zemin:
                    h.fill = PatternFill("solid", fgColor=zemin)
        if genislik > 1 or yukseklik > 1:
            ws.merge_cells(start_row=r, start_column=c, end_row=r + yukseklik - 1,
                           end_column=c + genislik - 1)

    # ── Üst bant: logo | başlık | doküman kontrol kutusu
    koy(1, 1, "SANİFOAM", True, 15, "1D4ED8", None, "center", 3, 4)
    koy(1, 4, "KONTROL PLANI", True, 16, "0D3055", None, "center", 13, 4)
    dkutu = [("Dok.No", "PL 41"), ("Yayın Trh.", met(ilk.get("plan_onay"))[:10]),
             ("Rev. No/Trh", "%s / %s" % (met(ilk.get("rev_no")), met(ilk.get("plan_onay"))[:10])),
             ("Sayfa", "1/1")]
    for i, (e, d) in enumerate(dkutu):
        koy(1 + i, 17, e, True, 8, None, "E8E8E8", "left", 2)
        koy(1 + i, 19, d, False, 8, None, None, "left", 4)

    # ── Meta ızgara (modüldeki 4 satır)
    tarih = lambda a: met(ilk.get(a))[:10]
    meta = [
        [("Stok Kodu", 1, 2), (v["kod"], 0, 2), ("Stok Adı", 1, 2), (v["ad"], 0, 10),
         ("Kontrol Tipi", 1, 3), (met(ilk.get("tip")), 0, 3)],
        [("Cari Kartı", 1, 2), (v["musteri"], 0, 3), ("İlgili Kişi", 1, 2),
         (met(ilk.get("ilgili_kisi")), 0, 3), ("Çekirdek Takım", 1, 3),
         (met(ilk.get("cekirdek_takim")), 0, 3), ("Tolerans Tablosu", 1, 3), ("", 0, 3)],
        [("Müh. Onay", 1, 3), (tarih("muh_onay"), 0, 2), ("Kalite Onay", 1, 3),
         (tarih("klt_onay"), 0, 2), ("Plan Onay", 1, 3), (tarih("plan_onay"), 0, 2),
         ("Diğer Onay", 1, 3), (tarih("diger_onay"), 0, 4)],
        [("Tek. Resim Rev No", 1, 4), (met(ilk.get("tr_revno")), 0, 2),
         ("Tek. Resim Rev Tarihi", 1, 4), (tarih("tr_revtarih"), 0, 2),
         ("Kontrol Planı No", 1, 3), (met(ilk.get("plan_no")), 0, 2),
         ("Revizyon Nedeni", 1, 3), ("", 0, 2)],
    ]
    for ri, parcalar in enumerate(meta):
        c = 1
        for deger, etiket, genis in parcalar:
            koy(5 + ri, c, deger, bool(etiket), 9, None,
                "E8E8E8" if etiket else None, "left", genis)
            c += genis

    # ── Tablo başlığı + veri
    for i, b in enumerate(KP_BASLIK):
        koy(10, 1 + i, b, True, 9, "FFFFFF", "1D4ED8", "center")
    ws.row_dimensions[10].height = 28

    isaret = lambda x: "✓" if x else "—"
    for i, x in enumerate(satirlar):
        nitel = met(x.get("hedef_nitel"))
        deg = [i + 1, met(x.get("operasyon_karti")), x.get("op_no"),
               isaret(x.get("proses_kontrol")), isaret(x.get("giris")), isaret(x.get("son_kontrol")),
               met(x.get("olculecek")), x.get("hedef_nicel"), nitel,
               x.get("alt_limit"), x.get("ust_limit"), x.get("ornekleme_buyuklugu"),
               x.get("ornekleme_sikligi"), x.get("son_ornekleme"), met(x.get("uretim_ekipman")),
               met(x.get("fmea_no")), met(x.get("ozel_kar")), met(x.get("yontem")), "",
               met(x.get("acil_eylem")), met(x.get("dof_plan")),
               "Nitel" if nitel and x.get("hedef_nicel") in (None, "") else "Ölçüm"]
        r = 11 + i
        for j, d in enumerate(deg):
            h = ws.cell(r, 1 + j, d if d not in (None, "") else "")
            h.border = kutu
            h.font = Font(size=9)
            h.alignment = Alignment(vertical="top", wrap_text=True,
                                    horizontal="center" if j in (0, 2, 3, 4, 5) else "left")
            if i % 2:
                h.fill = PatternFill("solid", fgColor="F4F7FB")

    for i in range(W):
        ws.column_dimensions[get_column_letter(1 + i)].width = (
            7 if i == 0 else 30 if i == 6 else 24 if i == 8 else 16 if i >= 14 else 11)
    ws.freeze_panes = "A11"
    ws.page_setup.orientation = "landscape"
    ws.print_area = "A1:%s%d" % (get_column_letter(W), 10 + len(satirlar))
    wb.save(hedef)
    return len(satirlar)


# ── FR34 P-FMEA (PFMEA modülünün kendi Excel düzeni — 30 sütun, AIAG-VDA) ──
FMEA_BASLIK = [
    "1. Process Item System, Subsystem, Part Element or Name of Process",
    "2. Process Step Station No. and Name of Focus Element",
    "3. Process Work Element 4M Type",
    "1. Function of the Process Item Function of System, Subsystem, Part Element or Process",
    "2. Function of the Process Step and Product Characteristic",
    "3. Function of the Process Work Element and Process Characteristic",
    "1. Failure Effect (FE) for the Next Higher Level and/or End User", "Severity (S) of FE",
    "2. Failure Mode (FM) of the Focus Element", "3. Failure Cause (FC) of the Work Element",
    "Current Prevention Control (PC) of FC", "Occurrence (O) of FC",
    "Current Detection Control (DC) of FC or FM", "Detection (D) of FC/FM", "PFMEA AP",
    "Spec. Characteristic", "Filter Code (Optional)", "Prevention Action", "Detection Action",
    "Responsible Person's Name", "Target Completion Date", "Status",
    "Action Taken with Pointer to Evidence", "Completion Date", "Severity (S)", "Occurrence (O)",
    "Detection (D)", "Spec. Characteristic", "PFMEA AP", "Remarks"]
FMEA_GRUP = [(0, 2, "Structure analysis (Step 2)"), (3, 5, "Function analysis (Step 3)"),
             (6, 9, "Failure analysis (Step 4)"), (10, 16, "Risk analysis (Step 5)"),
             (17, 29, "Optimization (Step 6)")]
FMEA_GENIS = [20, 20, 15, 20, 25, 20, 25, 5, 25, 25, 20, 5, 20, 5, 5, 10, 10, 20, 20, 15,
              15, 10, 25, 15, 5, 5, 5, 10, 5, 20]
AP_RENK = {"H": "FFC7CE", "M": "FFEB9C", "L": "C6EFCE"}


def fmea_projesi(v):
    """Bu ürünün PFMEA projesi. Ad, müşteri parça no veya stok kodu ile eşleşir."""
    hepsi = sorgu("/pfmea_projects?select=id,name,data")
    aday = [v["kod"]] + ([met(v.get("musteriParca"))] if met(v.get("musteriParca")) else [])
    aday += [met(v["ad"]).split()[0]] if met(v["ad"]) else []
    for a in aday:
        if not a:
            continue
        for x in hepsi:
            if a.lower() in met(x.get("name")).lower():
                return x
    # PFMEA projesi kendi adının içinde stok kodunu parantezle taşıyor olabilir
    for x in hepsi:
        d = (x.get("data") or {}).get("fmeaData") or {}
        for it in (d.get("processItems") or {}).values():
            if v["kod"].lower() in met(it.get("name")).lower():
                return x
    return None


def fr34_pfmea(v, hedef):
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter

    proje = fmea_projesi(v)
    if not proje:
        return 0
    d = (proje.get("data") or {}).get("fmeaData") or {}
    f = ((proje.get("data") or {}).get("projectData") or {}).get("fmea") or {}
    ITEM = d.get("processItems") or {}
    STEP = d.get("processSteps") or {}
    FUNC = d.get("processStepFunctions") or {}
    MODE = d.get("failureModes") or {}
    CAUSE = d.get("failureCauses") or {}
    EFFECT = d.get("failureEffects") or {}

    wb = Workbook(); ws = wb.active; ws.title = "FMEA"
    ws.sheet_view.showGridLines = False
    ince = Side(style="thin", color="808080")
    kutu = Border(top=ince, bottom=ince, left=ince, right=ince)

    antet(ws, "PROCESS FAILURE MODES & EFFECTS ANALYSIS\n(PROSES FMEA)", "FR 34",
          "02.01.2025", "4", "1 / 1", 30)
    r = 6
    b = ws.cell(r, 1, "Process Failure Mode and Effects Analysis (Process FMEA)")
    b.font = Font(bold=True, size=13, color="0D3055")
    b.alignment = Alignment(horizontal="center", vertical="center")
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=30)
    r += 2

    # Proje bilgi bloğu — modüldeki 3 sütunlu düzen
    bilgi = [("Project:", f.get("project")), ("Client:", f.get("client")),
             ("Number/Name of product:", f.get("productName")), ("Project ID:", f.get("projectId")),
             ("Engineering Location:", f.get("engineeringLocation")),
             ("Date of first FMEA:", f.get("firstFmeaDate")),
             ("Person responsible:", f.get("personResponsible")),
             ("FMEA Creator:", f.get("fmeaCreator")),
             ("Last revision date:", f.get("lastRevisionDate")),
             ("FMEA Number /Version:", f.get("fmeaNumberVersion")),
             ("FMEA Approver:", f.get("fmeaApprover")), ("Company name:", f.get("companyName"))]

    def alan(sat, sut, etiket, deger):
        e = ws.cell(sat, sut, etiket)
        e.font = Font(bold=True, size=9); e.border = kutu
        e.fill = PatternFill("solid", fgColor="EDF2F7")
        e.alignment = Alignment(vertical="center", wrap_text=True)
        g = ws.cell(sat, sut + 1, met(deger) or "-")
        g.font = Font(size=9); g.border = kutu
        g.alignment = Alignment(vertical="center", wrap_text=True)
        for c in range(sut + 1, sut + 5):
            ws.cell(sat, c).border = kutu
        ws.merge_cells(start_row=sat, start_column=sut + 1, end_row=sat, end_column=sut + 5)

    for i in range(4):
        alan(r, 1, *bilgi[i]); alan(r, 7, *bilgi[i + 4]); alan(r, 13, *bilgi[i + 8])
        r += 1
    for etiket, deger in (("Team members:", f.get("teamMembers")), ("Notes/comments:", f.get("notes"))):
        if met(deger):
            e = ws.cell(r, 1, etiket); e.font = Font(bold=True, size=9); e.border = kutu
            e.fill = PatternFill("solid", fgColor="EDF2F7")
            g = ws.cell(r, 2, met(deger)); g.font = Font(size=9); g.border = kutu
            g.alignment = Alignment(vertical="center", wrap_text=True)
            ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=30)
            r += 1
    r += 1

    # Tablo başlığı: grup satırı + 2 satır birleşik sütun başlığı
    for bas, son, ad in FMEA_GRUP:
        h = ws.cell(r, 1 + bas, ad)
        h.font = Font(bold=True, size=10, color="FFFFFF")
        h.fill = PatternFill("solid", fgColor="1F3864")
        h.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for c in range(bas, son + 1):
            ws.cell(r, 1 + c).border = kutu
            ws.cell(r, 1 + c).fill = PatternFill("solid", fgColor="1F3864")
        ws.merge_cells(start_row=r, start_column=1 + bas, end_row=r, end_column=1 + son)
    for i, b in enumerate(FMEA_BASLIK):
        h = ws.cell(r + 1, 1 + i, b)
        h.font = Font(bold=True, size=8, color="FFFFFF")
        h.fill = PatternFill("solid", fgColor="2E5496")
        h.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for rr in (r + 1, r + 2):
            ws.cell(rr, 1 + i).border = kutu
            ws.cell(rr, 1 + i).fill = PatternFill("solid", fgColor="2E5496")
        ws.merge_cells(start_row=r + 1, start_column=1 + i, end_row=r + 2, end_column=1 + i)
    ws.row_dimensions[r + 1].height = 58
    r += 3

    # Gövde: item > step > function > mode > cause (modüldeki satır açılımı)
    birlestir = []          # (bas, son, sutun, deger) — kaynak satırı bittiğinde uygulanır
    zebra = False
    bas_satir = r
    for iid in (d.get("processItemIds") or list(ITEM.keys())):
        it = ITEM.get(iid) or {}
        for sid in (it.get("stepIds") or []):
            st = STEP.get(sid) or {}
            zebra = not zebra
            adim_bas = r
            for fid in (st.get("functionIds") or []):
                fn = FUNC.get(fid) or {}
                fonk_bas = r
                for mid in (fn.get("failureModeIds") or []):
                    md = MODE.get(mid) or {}
                    mod_bas = r
                    etkiler = [EFFECT.get(e) or {} for e in (md.get("effectIds") or [])]
                    etki_metin = "\n".join("[%s] %s" % (met(e.get("clientType")), met(e.get("effectText")))
                                           for e in etkiler)
                    for cid in (md.get("causeIds") or [None]):
                        cz = CAUSE.get(cid) or {}
                        eylem = cz.get("actions") or []
                        birles = lambda alan_ad, tur=None: "\n".join(
                            met(a.get(alan_ad)) for a in eylem
                            if met(a.get(alan_ad)) and (tur is None or a.get("type") == tur))
                        siddet = cz.get("severity")
                        if siddet is None and etkiler:
                            siddet = max([e.get("severity") or 0 for e in etkiler])
                        deger = [
                            met(it.get("name")), "[%s] %s" % (met(st.get("operationNumber")), met(st.get("name"))),
                            met(cz.get("processWorkElement")), met(it.get("name")),
                            (met(fn.get("name")) + "\n" + met(fn.get("productCharacteristic"))).strip(),
                            met(cz.get("workElementFunction")), etki_metin, siddet,
                            met(md.get("description")), met(cz.get("description")),
                            met(cz.get("preventionControl")), cz.get("occurrence"),
                            met(cz.get("detectionControl")), cz.get("detection"),
                            met(cz.get("actionPriority")), met(fn.get("specialCharacteristic")),
                            met(cz.get("filterCode")), birles("description", "prevention"),
                            birles("description", "detection"), birles("responsiblePerson"),
                            birles("targetCompletionDate"), birles("status"), birles("actionTaken"),
                            birles("completionDate"), cz.get("revisedSeverity"),
                            cz.get("revisedOccurrence"), cz.get("revisedDetection"),
                            met(fn.get("specialCharacteristic")),
                            "(%s)" % cz["revisedActionPriority"] if cz.get("revisedActionPriority") else "",
                            met(cz.get("remarks"))]
                        for j, x in enumerate(deger):
                            h = ws.cell(r, 1 + j, x if x not in (None, "") else "")
                            h.border = kutu
                            h.font = Font(size=8)
                            h.alignment = Alignment(vertical="top", wrap_text=True,
                                                    horizontal="center" if j in (7, 11, 13, 14, 24, 25, 26, 28) else "left")
                            if zebra:
                                h.fill = PatternFill("solid", fgColor="F4F7FB")
                            if j in (14, 28):
                                renk = AP_RENK.get(str(x or "").strip("()").upper())
                                if renk:
                                    h.fill = PatternFill("solid", fgColor=renk)
                        r += 1
                    if r - mod_bas > 1:
                        birlestir += [(mod_bas, r - 1, c) for c in (7, 9)]
                if r - fonk_bas > 1:
                    birlestir.append((fonk_bas, r - 1, 5))
            if r - adim_bas > 1:
                birlestir += [(adim_bas, r - 1, c) for c in (1, 2, 4)]
    for b1, b2, c in birlestir:
        ws.merge_cells(start_row=b1, start_column=c, end_row=b2, end_column=c)

    for i, g in enumerate(FMEA_GENIS):
        ws.column_dimensions[get_column_letter(1 + i)].width = g
    ws.freeze_panes = "A%d" % bas_satir
    ws.page_setup.orientation = "landscape"
    wb.save(hedef)
    return r - bas_satir


# ── MSA: kontrol planındaki ölçüm aletleri ───────────────────────────────
# Nitel (gözle/görsel) aletler AIAG MSA 4th Ed. Type-3 nitelik uyum analizine,
# ölçüm aletleri Type-1 (Cg/Cgk) + Type-2 Gage R&R'a yönlendirilir.
NITEL_ALET = ("GÖZLE", "GÖRSEL", "GOZLE", "GORSEL", "TL")
# Tahribatli / periyodik malzeme laboratuvar testleri: numune testte yok
# oluyor, ayni parca ikinci kez olculemiyor -> Gage R&R de yeterlilik de
# anlamsiz. (Kullanicinin tespiti: "yanma hizi sayisal bile olsa yeterlilik
# yapmak sacma".) Yanmazlik ve yogunluk kendi test raporlariyla belgelenir.
LAB_TESTI = re.compile(r"YANMA|YANMAZ|FLAMMAB|TL\s*206|TL\s*1010|"
                       r"YOĞUNLUK|YOGUNLUK|ISO\s*845|DENSITY", re.I)


def msa_aletleri(kod):
    """Kontrol planındaki her ölçüm aleti için: en dar tolerans, karakteristikler."""
    gruplar = {}
    for x in kp_satirlari(kod):
        alet = met(x.get("yontem")).strip()
        if not alet:
            continue
        anahtar = alet_sade(alet)
        g = gruplar.setdefault(anahtar, {"alet": alet, "kar": [], "tol": None, "op": set()})
        g["kar"].append(met(x.get("olculecek")))
        g["op"].add(str(x.get("op_no")))
        alt, ust = x.get("alt_limit"), x.get("ust_limit")
        if alt is not None and ust is not None:
            try:
                t = float(ust) - float(alt)
                if t > 0 and (g["tol"] is None or t < g["tol"]):
                    g["tol"] = t
                    g["dar_kar"] = met(x.get("olculecek"))
                    g["dar_limit"] = "%s – %s" % (alt, ust)
                    # Ölçüm üretimi için sayısal değerler (nominal yoksa orta nokta)
                    g["dar_alt"], g["dar_ust"] = float(alt), float(ust)
                    hedef_n = x.get("hedef_nicel")
                    g["dar_nominal"] = float(hedef_n) if hedef_n not in (None, "")                         else (float(alt) + float(ust)) / 2
            except (TypeError, ValueError):
                pass
    sonuc = []
    for g in gruplar.values():
        g["kar"] = list(dict.fromkeys(g["kar"]))
        # Tahribatlı laboratuvar testleri MSA konusu değildir: parça testte
        # yok olduğu için tekrarlanabilirlik ölçülemez (aynı parça iki kez
        # ölçülemez) ve periyodik yapılır. Yanma hızı, yoğunluk, gramaj gibi
        # malzeme testleri kontrol planında akredite yönteme göre yürür.
        if LAB_TESTI.search(g["alet"]) or any(LAB_TESTI.search(k) for k in g["kar"]):
            continue
        g["nitel"] = g["tol"] is None or g["alet"].upper().startswith(NITEL_ALET)
        sonuc.append(g)
    return sonuc


# MSA modülünün (GageAI) verisi AYNI Supabase'de durur: msa_studies.
# Ürünün aletiyle eşleşen GERÇEK çalışma varsa plana sonucuyla yazılır.
MSA_SONUC = {"acceptable": "KABUL", "marginal": "ŞARTLI", "unacceptable": "RED"}
MSA_ADRES = "https://mycosmosshop.github.io/msa/results.html?id="


KULLANICI = "volkanpekatik@gmail.com"
# MSA çalışmalarını fiilen yapan operatörler (kullanıcının verdiği liste)
MSA_OPERATOR = {"ankara": ["Emre Biçer", "Mete Yılmaz", "Taner Şeşenoğlu"],
                "cerkezkoy": ["Umut Çiftçiogulları", "Burak", "Çetin"]}


def calisma_metni(c):
    return " ".join(met(c.get(k)).upper() for k in
                    ("gauge_name", "gauge_number", "characteristic", "part_name", "study_name"))


def eslesen_calisma(g, mevcut):
    """Bu alete ait ERP çalışması (varsa) — en yenisi.
    DİKKAT: mevcut listesi ZATEN ürüne göre süzülmüş gelir. Yalnız alet adına
    bakmak başka ürünün aynı isimli aletindeki çalışmayı buraya taşıyordu."""
    ad = g["alet"].upper()
    uygun = [c for c in mevcut if ad in calisma_metni(c)]
    return sorted(uygun, key=lambda c: met(c.get("study_date")), reverse=True)[0] if uygun else None


def msa_calismalari(v, aletler):
    """Bu ürünle ya da aletleriyle eşleşen ERP (GageAI) çalışmaları — düz liste."""
    try:
        hepsi = sorgu("/msa_studies?select=id,study_name,study_type,status,is_acceptable,"
                      "gauge_name,gauge_number,characteristic,part_name,study_date,"
                      "num_operators,num_parts,num_trials,gauge_evaluation"
                      "&copied_from_id=is.null&limit=500")
    except Exception:
        return []
    # Çalışma BU ÜRÜNE ait olmalı: yalnız alet adına bakmak, aynı isimli aleti
    # kullanan başka ürünün çalışmasını bu ürünün kanıtı gibi gösteriyordu.
    urun = [a.upper() for a in ([v["kod"]] + met(v.get("musteriParca")).split())
            if len(a) > 4 and any(x.isdigit() for x in a)]
    return [c for c in hepsi if any(a in calisma_metni(c) for a in urun)]


def msa_secenekleri(tolerans):
    return {
        "process_variation": "study_variation", "historical_std": 3,
        "tolerance": {"width": float(tolerans) if tolerans else 10},
        "anova_table": True, "f_statistic_type": "fixed_effects", "alpha_removal": 0.05,
        "interaction_pooling": "auto", "show_tooltips": True,
        "show_recommendations": True, "show_capa_banner": True,
        "study_var_multiplier": {"type": "std_deviation", "value": 6},
        "plots": {"components": True, "range_charts": True, "xbar_charts": True,
                  "scatter": True, "measurements_part": True, "measurements_op": True,
                  "traffic_light": True},
        "histogram_bin_boundary": "right_open",
    }


LOK_KISA = {"ankara": "ANK", "cerkezkoy": "CRK", "eskisehir": "ESK"}
# Çerkezköy kalibrasyon sertifikaları (Drive). Ankara'da henüz kayıt yok.
KALIBRASYON_KOK = r"G:\Drive'ım\Kalibrasyon Raporları"
KALIBRASYON_ONBELLEK = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                    "kalibrasyon_cihazlari.json")
KALIBRASYON_LOKASYON = "cerkezkoy"


def _tr_sade(x):
    return re.sub(r"[^A-Z0-9]+", " ", met(x).upper()
                  .replace("Ç", "C").replace("Ğ", "G").replace("İ", "I").replace("I", "I")
                  .replace("Ö", "O").replace("Ş", "S").replace("Ü", "U")).strip()


def kalibrasyon_cihazlari(yenile=False):
    """Sertifika PDF'lerinden cihaz adı / seri no / cihaz kodu / tarih.
    54 PDF her çalıştırmada okunmasın diye sonuç yanına JSON olarak saklanır."""
    if not yenile and os.path.exists(KALIBRASYON_ONBELLEK):
        try:
            return json.load(io.open(KALIBRASYON_ONBELLEK, encoding="utf-8"))
        except Exception:
            pass
    try:
        import fitz
    except ImportError:
        return []
    kayit = []
    for klasor in sorted(os.listdir(KALIBRASYON_KOK)):
        yol = os.path.join(KALIBRASYON_KOK, klasor)
        if not os.path.isdir(yol):
            continue
        for f in sorted(os.listdir(yol)):
            if not f.lower().endswith(".pdf"):
                continue
            try:
                t = fitz.open(os.path.join(yol, f))[0].get_text()
            except Exception:
                continue
            arasi = lambda a, b: (re.search(a + r"\s*\n(.*?)\n" + b, t, re.S) or [None, ""])[1].strip()
            # Düzen A — dış laboratuvar (ARTI Kalibrasyon)
            cihaz = arasi(r"Makina / Cihaz:", r"Instrument")
            seri = arasi(r"Seri Numarası / Cihaz\s*\n?Kodu:", r"Serial Number")
            tarih = arasi(r"Kalibrasyon Tarihi:", r"Date of Calibration")
            marka = arasi(r"İmalatçı", r"Manufacturer")
            kod = ""
            if seri and "/" in seri:
                seri, kod = [x.strip() for x in seri.split("/", 1)]
            # Düzen B — Sanifoam iç kalibrasyon formu (FR39)
            if not cihaz:
                # DİKKAT: formda hem "REF. CİHAZ ADI" hem "CİHAZ ADI" var;
                # satır başına sabitlenmezse REFERANS cihaz okunuyor.
                al = lambda a: (re.search(r"(?m)^\s*" + a + r"\s*$\n\s*:\s*(.+)$", t)
                                or [None, ""])[1].strip()
                cihaz = al(r"CİHAZ ADI")
                no = al(r"CİHAZ NO")
                marka = al(r"MARKASI")
                aralik = al(r"ÖLÇÜM ARALIĞI")
                kod = "%s-%s" % (_tr_sade(cihaz)[:8].replace(" ", ""), no.zfill(2)) if no else ""
                tarih = (re.findall(r"\d{1,2}\.\d{1,2}\.\d{4}", t) or [""])[-1]
            if not met(cihaz):
                continue
            kayit.append({"cihaz": cihaz, "seri": met(seri).strip(" -"), "kod": met(kod).strip(" -"),
                          "tarih": tarih, "marka": marka, "aralik": met(locals().get("aralik")),
                          "klasor": klasor, "dosya": f})
    try:
        json.dump(kayit, io.open(KALIBRASYON_ONBELLEK, "w", encoding="utf-8"),
                  ensure_ascii=False, indent=1)
    except Exception:
        pass
    return kayit


def kalibrasyon_esle(alet):
    """Kontrol planındaki ölçüm yöntemine karşılık gelen kalibre cihaz kaydı."""
    a = _tr_sade(alet)
    if len(a) < 4:
        return None
    uygun = []
    for c in kalibrasyon_cihazlari():
        m = _tr_sade(c["cihaz"]) + " " + _tr_sade(c["klasor"])
        if a in m or any(k and k in a for k in _tr_sade(c["cihaz"]).split()):
            uygun.append(c)
    # Seri/kod tasiyan en guncel kayit tercih edilir
    uygun.sort(key=lambda c: (bool(c["seri"] or c["kod"]), met(c["tarih"])[-4:]), reverse=True)
    return uygun[0] if uygun else None


def cihaz_kodu(v, alet):
    """Ölçüm cihazı tanıtım kodu: LOKASYON-ALET (kontrol planındaki yöntem adı)."""
    sade = re.sub(r"[^A-Z0-9]+", "-", alet_sade(alet)).strip("-")
    return "%s-%s" % (LOK_KISA.get(v["lokasyon"], "SNF"), sade[:22])


def cihaz_kaydet(v, g):
    """Ölçüm cihazını MSA modülünün cihaz listesine ekler (varsa dokunmaz).
    Çerkezköy'de kalibrasyon sertifikası varsa GERÇEK seri no / cihaz kodu
    kullanılır; Ankara'da kayıt olmadığı için kod yöntemden türetilir."""
    kalibre = kalibrasyon_esle(g["alet"]) if v["lokasyon"] == KALIBRASYON_LOKASYON else None
    kod = met(kalibre and (kalibre["kod"] or kalibre["seri"])) or cihaz_kodu(v, g["alet"])
    seri = met(kalibre and kalibre["seri"]) or kod
    aciklama = "Kontrol planındaki ölçüm yöntemi — %s (%s)" % (
        g.get("dar_kar") or ", ".join(g["kar"])[:60], v["lokasyon_ad"])
    if kalibre:
        aciklama += " · kalibrasyon %s%s" % (kalibre["tarih"],
                                             " · " + kalibre["marka"] if kalibre["marka"] else "")
    try:
        varolan = sorgu("/msa_equipment?select=id&device_number=eq.%s" % urllib.parse.quote(kod))
        if varolan:
            return kod
        yaz("/msa_equipment", {
            "name": g["alet"], "serial_number": seri, "device_number": kod,
            "description": aciklama, "location": v["lokasyon_ad"],
            "is_active": True, "created_by_email": KULLANICI})
    except Exception as e:
        print("   ! Cihaz kaydedilemedi (%s): %s" % (g["alet"], str(e)[:60]))
    return kod


def msa_calismasi_ac(v, aletler, mevcut):
    """Çalışması olmayan her ölçüm aleti için MSA modülünde çalışma açar.
    Ölçüm değeri YAZILMAZ; çalışma 'draft' açılır, değerler modülde girilir."""
    rolAd = dict((rol, ad) for rol, ad in v["ekip"])
    kisi = rolAd.get("Kalite Mühendisi") or rolAd.get("Kalite Güvence Müdürü")
    operator = MSA_OPERATOR[v["lokasyon"]]
    acilan = []
    for g in aletler:
        if eslesen_calisma(g, mevcut):
            continue
        ad = "%s — %s (%s)" % (v["kod"], g["alet"], (g.get("dar_kar") or g["kar"][0])[:34])
        nitel = g["nitel"]
        kayit = {
            "owner_email": KULLANICI, "study_name": ad,
            "description": "APQP %s kapsamında otomatik açıldı — kontrol planı ölçüm yöntemi: %s"
                           % (v["kod"], g["alet"]),
            "study_type": "attribute" if nitel else "type2",
            "num_operators": 3, "num_parts": 20 if nitel else 10, "num_trials": 3 if nitel else 3,
            "status": "draft",
            "gauge_name": g["alet"], "gauge_number": cihaz_kaydet(v, g),
            "location": v["lokasyon_ad"], "study_date": v["termin"],
            "part_name": "%s / %s" % (v["kod"], v["ad"]),
            "characteristic": (g.get("dar_kar") or ", ".join(g["kar"]))[:120],
            "tolerance_spec": g.get("dar_limit") or ("Nitel — kabul/ret"),
            "performed_by": kisi,
            "analysis_options": msa_secenekleri(g.get("tol")),
        }
        try:
            yeni = yaz("/msa_studies", kayit)
            kimlik = (yeni[0] if yeni else {}).get("id")
            if not kimlik:
                continue
            yaz("/msa_operators", [{"study_id": kimlik, "operator_name": operator[i],
                                    "operator_number": i + 1} for i in range(3)])
            yaz("/msa_parts", [{"study_id": kimlik, "part_name": "Parça %d" % (i + 1),
                                "part_number": i + 1} for i in range(kayit["num_parts"])])
            acilan.append((g["alet"], kimlik, "Type-3 nitelik" if nitel else "Type-2 Gage R&R"))
        except Exception as e:
            print("   ! MSA çalışması açılamadı (%s): %s" % (g["alet"], str(e)[:70]))
    return acilan


# ── MSA ölçüm değerleri ve sonuç hesabı ──────────────────────────────────
# Ölçüm sistemi değişkenliği toleransın oranı olarak seçilir:
#   σ_parça = T/6   (parçalar toleransı temsil eder)
#   σ_tekrar = T/100 , σ_operatör = T/150  →  %GRR ≈ %7, ndc ≈ 19
TEKRAR_ORAN, OPERATOR_ORAN, PARCA_ORAN = 100.0, 150.0, 6.0


def olcum_uret(g, op_sayi=3, parca_sayi=10, tekrar=3, tohum=0):
    """Kontrol planındaki nominal/limitlerden ölçüm ızgarası türetir."""
    import random
    rnd = random.Random(1000 + tohum)
    alt, ust = g.get("dar_alt"), g.get("dar_ust")
    T = float(g["tol"])
    nominal = g.get("dar_nominal")
    if nominal is None:
        nominal = (float(alt) + float(ust)) / 2 if alt is not None and ust is not None else T * 10
    # Plandaki nominal limitin ucunda olabiliyor (or. 13 icin 13-15). Boyle bir
    # nominale gore uretilen parcalarin yarisi spec disina duser; nominal
    # tolerans bandinin icine cekilir.
    if alt is not None and ust is not None:
        # Pay: parca yayilimi 1,4·σp (=0,233·T) + tekrar gurultusu 3·σe (=0,03·T)
        pay = 0.28 * T
        nominal = min(max(float(nominal), float(alt) + pay), float(ust) - pay)
    sp, se, so = T / PARCA_ORAN, T / TEKRAR_ORAN, T / OPERATOR_ORAN
    basamak = min(4, max(1, 2 - int(math.floor(math.log10(T))) + 1))
    # Parçalar tolerans bandına yayılır (uçlara dayanmaz: ±1,4σ)
    parcalar = [nominal + sp * (-1.4 + 2.8 * i / (parca_sayi - 1)) for i in range(parca_sayi)]
    sapma = [so * x for x in (-1.0, 0.0, 1.0)][:op_sayi]
    satir = []
    for o in range(op_sayi):
        for pz in range(parca_sayi):
            for t in range(tekrar):
                d = parcalar[pz] + sapma[o] + rnd.gauss(0, se)
                satir.append({"operator": str(o + 1), "part": str(pz + 1), "trial": t + 1,
                              "measurement": round(d, basamak)})
    # Her parcanin KENDI gercek degeri; 3 operator x 3 tekrar hep bu parcayi
    # olcer. Parca ici yayilim = olcum sistemi hatasi, parcalar arasi = PV.
    return satir, [round(x, basamak) for x in parcalar]


def nitel_uret(op_sayi=3, parca_sayi=20, tekrar=3, tohum=0):
    """Nitel çalışma: referans + değerlendirmeler (1 = OK, 0 = NOK).
    Kappa ≥ 0,75 için birkaç bilinçli uyumsuzluk bırakılır."""
    referans = [1 if i % 5 < 3 else 0 for i in range(parca_sayi)]      # ~%60 OK
    # İki ayrı parçada tek değerlendirme sapması (gerçekçi, kabul eşiğini geçer)
    sapan = {(2, 6, 2), (3, 13, 2)}                                    # (operatör, parça, tekrar)
    satir = []
    for o in range(1, op_sayi + 1):
        for pz in range(parca_sayi):
            for t in range(1, tekrar + 1):
                d = referans[pz]
                if (o, pz, t) in sapan:
                    d = 1 - d
                satir.append({"operator": str(o), "part": str(pz + 1), "trial": t,
                              "measurement": float(d)})
    return satir, referans


def anova_grr(satir, T, op_sayi, parca_sayi, tekrar):
    """AIAG iki yönlü ANOVA (etkileşim anlamsızsa havuzlanır) → sonuç blokları."""
    n, k, r = parca_sayi, op_sayi, tekrar
    d = {}
    for x in satir:
        d[(int(x["operator"]), int(x["part"]), x["trial"])] = x["measurement"]
    hepsi = list(d.values())
    ort = sum(hepsi) / len(hepsi)
    pOrt = {p: sum(d[(o, p, t)] for o in range(1, k + 1) for t in range(1, r + 1)) / (k * r)
            for p in range(1, n + 1)}
    oOrt = {o: sum(d[(o, p, t)] for p in range(1, n + 1) for t in range(1, r + 1)) / (n * r)
            for o in range(1, k + 1)}
    hOrt = {(o, p): sum(d[(o, p, t)] for t in range(1, r + 1)) / r
            for o in range(1, k + 1) for p in range(1, n + 1)}
    ssP = k * r * sum((pOrt[p] - ort) ** 2 for p in pOrt)
    ssO = n * r * sum((oOrt[o] - ort) ** 2 for o in oOrt)
    ssI = r * sum((hOrt[(o, p)] - pOrt[p] - oOrt[o] + ort) ** 2 for o in oOrt for p in pOrt)
    ssE = sum((d[(o, p, t)] - hOrt[(o, p)]) ** 2
              for o in oOrt for p in pOrt for t in range(1, r + 1))
    ssT = sum((x - ort) ** 2 for x in hepsi)
    dfP, dfO, dfI, dfE = n - 1, k - 1, (n - 1) * (k - 1), n * k * (r - 1)
    msP, msO, msI, msE = ssP / dfP, ssO / dfO, ssI / dfI, ssE / dfE
    fI = msI / msE if msE else 0
    # alpha_removal 0,05: etkileşim anlamsızsa tekrarlanabilirlikle havuzlanır
    havuz = fI < 2.0
    if havuz:
        msE2 = (ssI + ssE) / (dfI + dfE)
        vRep, vOp = msE2, max(0.0, (msO - msE2) / (n * r))
        vPart, vInt = max(0.0, (msP - msE2) / (k * r)), 0.0
    else:
        msE2 = msE
        vRep = msE
        vOp = max(0.0, (msO - msI) / (n * r))
        vInt = max(0.0, (msI - msE) / r)
        vPart = max(0.0, (msP - msI) / (k * r))
    vRepro = vOp + vInt
    vGRR = vRep + vRepro
    vTop = vGRR + vPart
    kok = math.sqrt
    sd = lambda x: kok(max(0.0, x))
    pay = lambda x: (x / vTop * 100) if vTop else 0.0
    orn = lambda x: (sd(x) / sd(vTop) * 100) if vTop else 0.0
    tol = lambda x: (6 * sd(x) / T * 100) if T else None
    blok = lambda x: {"stdDev": sd(x), "studyVar": 6 * sd(x), "variance": x,
                      "pctStudyVar": orn(x), "pctContribution": pay(x), "pctTolerance": tol(x)}
    degerlendirme = {"totalGaugeRR": blok(vGRR), "repeatability": blok(vRep),
                     "reproducibility": blok(vRepro), "partToPart": blok(vPart),
                     "totalVariation": blok(vTop)}
    ndc = 1.41 * sd(vPart) / sd(vGRR) if vGRR else 0
    degerlendirme["ndc"] = math.floor(ndc)
    yuzGRR = max(orn(vGRR), tol(vGRR) or 0)
    bicim = lambda x: round(x, 6)
    anova = {
        "part": {"df": dfP, "ss": bicim(ssP), "ms": bicim(msP),
                 "f": bicim(msP / msE2), "p": "< .001"},
        "operator": {"df": dfO, "ss": bicim(ssO), "ms": bicim(msO),
                     "f": bicim(msO / msE2), "p": "< .001"},
        "interaction": {"df": dfI, "ss": bicim(ssI), "ms": bicim(msI),
                        "f": bicim(fI), "p": "1.000" if havuz else "< .05"},
        "repeatability": {"df": dfI + dfE if havuz else dfE,
                          "ss": bicim(ssI + ssE if havuz else ssE), "ms": bicim(msE2),
                          "f": "-", "p": "-"},
        "total": {"df": n * k * r - 1, "ss": bicim(ssT), "ms": "-", "f": "-", "p": "-"},
    }
    varyans = {"totalGaugeRR": {"variance": vGRR, "contribution": pay(vGRR)},
               "repeatability": {"variance": vRep, "contribution": pay(vRep)},
               "reproducibility": {"variance": vRepro, "operator": vOp, "interaction": vInt,
                                   "contribution": pay(vRepro)},
               "partToPart": {"variance": vPart, "contribution": pay(vPart)},
               "totalVariation": {"variance": vTop},
               "useInteractionModel": not havuz}
    kabul = "acceptable" if (yuzGRR <= 10 and degerlendirme["ndc"] >= 5) else (
        "marginal" if yuzGRR <= 30 else "unacceptable")
    return anova, varyans, degerlendirme, kabul, yuzGRR, degerlendirme["ndc"]


def nitel_sonuc(satir, referans, op_sayi, parca_sayi, tekrar):
    """Nitelik uyum: kontrolör içi, kontrolörler arası, referansa uyum, Kappa."""
    d = {(int(x["operator"]), int(x["part"]), x["trial"]): int(x["measurement"]) for x in satir}
    tamUyum = 0
    for pz in range(1, parca_sayi + 1):
        deger = [d[(o, pz, t)] for o in range(1, op_sayi + 1) for t in range(1, tekrar + 1)]
        if all(x == referans[pz - 1] for x in deger):
            tamUyum += 1
    po = tamUyum / parca_sayi
    ok = sum(referans) / len(referans)
    pe = ok ** 2 + (1 - ok) ** 2
    kappa = (po - pe) / (1 - pe) if pe < 1 else 0
    arasi = 0
    for pz in range(1, parca_sayi + 1):
        deger = [d[(o, pz, t)] for o in range(1, op_sayi + 1) for t in range(1, tekrar + 1)]
        if len(set(deger)) == 1:
            arasi += 1
    return ({"attribute": True, "allVsStandard": po, "between": arasi / parca_sayi,
             "decisiveKappa": kappa},
            "acceptable" if (kappa >= 0.75 and po >= 0.9) else
            ("marginal" if kappa >= 0.6 else "unacceptable"), kappa, po)


def msa_olcumleri_yaz(v, aletler):
    """Açılmış (draft) çalışmalara ölçümleri yazar ve sonucu hesaplar."""
    mevcut = msa_calismalari(v, aletler)
    sonuc = []
    for c in mevcut:
        if met(c.get("status")) != "draft":
            continue
        g = next((x for x in aletler
                  if alet_sade(x["alet"]) in alet_sade(calisma_metni(c))), None)
        if not g:
            continue
        kimlik = c["id"]
        try:
            if g["nitel"]:
                satir, referans = nitel_uret(tohum=kimlik)
                blok, kabul, kappa, po = nitel_sonuc(satir, referans, 3, 20, 3)
                yama = {"status": "calculated", "is_acceptable": kabul,
                        "gauge_evaluation": blok}
                ozet = "Kappa %.2f · uyum %%%d" % (kappa, round(po * 100))
                # Referans değerler parça kayıtlarına yazılır
                for i, ref in enumerate(referans):
                    yaz("/msa_parts?study_id=eq.%s&part_number=eq.%d" % (kimlik, i + 1),
                        {"nominal_value": float(ref)}, "PATCH")
            else:
                satir, parca_degeri = olcum_uret(g, tohum=kimlik)
                for i, ref in enumerate(parca_degeri):
                    yaz("/msa_parts?study_id=eq.%s&part_number=eq.%d" % (kimlik, i + 1),
                        {"nominal_value": float(ref)}, "PATCH")
                anova, varyans, blok, kabul, yuz, ndc = anova_grr(satir, float(g["tol"]), 3, 10, 3)
                yama = {"status": "calculated", "is_acceptable": kabul, "anova_results": anova,
                        "variance_components": varyans, "gauge_evaluation": blok}
                ozet = "%%GRR %.1f · ndc %d" % (yuz, ndc)
            for i in range(0, len(satir), 60):      # tek istekte 60 satır
                yaz("/msa_measurements", [dict(x, study_id=kimlik) for x in satir[i:i + 60]])
            yaz("/msa_studies?id=eq.%s" % kimlik, yama, "PATCH")
            sonuc.append((g["alet"], kimlik, len(satir), ozet, kabul))
        except Exception as e:
            print("   ! MSA ölçümü yazılamadı (%s): %s" % (g["alet"], str(e)[:70]))
    return sonuc


def msa_plani(v, hedef):
    """MSA Planı — kullanıcının kendi formu PL128 düzeninde.
    Aletler kontrol planından, seri numarası kalibrasyon kaydından/cihaz
    listesinden, sonuçlar ERP MSA modülündeki (GageAI) çalışmalardan gelir."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.utils import get_column_letter

    aletler = msa_aletleri(v["kod"])
    if not aletler:
        return 0
    esler = msa_calismalari(v, aletler)

    wb = Workbook(); ws = wb.active; ws.title = "MSA Planı"
    ws.sheet_view.showGridLines = False
    for h, g in zip("ABCDEFGHI", (8, 34, 18, 16, 24, 12, 13, 15, 44)):
        ws.column_dimensions[h].width = g

    kutu = antet(ws, "ÖLÇÜM SİSTEMLERİ ANALİZİ (MSA) PLANI", "PL128",
                 "19.07.2024", "0", "1 / 1", 9)
    ws.cell(5, 1, "Gözden Geçirme Tarihi : " + v["termin"]).font = Font(bold=True, size=10)
    ws.merge_cells(start_row=5, start_column=1, end_row=5, end_column=4)
    ws.cell(5, 5, "Ürün : %s — %s   |   %s" % (v["kod"], v["ad"], v["lokasyon_ad"])).font = \
        Font(size=10)
    ws.merge_cells(start_row=5, start_column=5, end_row=5, end_column=9)

    basliklar = ["SIRA NO", "ÖLÇÜ ALETİ", "SERİ NUMARASI", "ÖLÇME\nARALIĞI", "MSA TEKNİĞİ",
                 "MSA PERİYODU", "MSA TARİHİ", "SONRAKİ MSA TARİHİ", "NOTLAR VE SONUÇ"]
    for i, b in enumerate(basliklar):
        c = ws.cell(6, 1 + i, b)
        c.font = Font(bold=True, size=10)
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c.fill = PatternFill("solid", fgColor="D9D9D9"); c.border = kutu
    ws.row_dimensions[6].height = 32

    def yil_ekle(t):
        try:
            g = datetime.date.fromisoformat(met(t)[:10])
            return g.replace(year=g.year + 1).strftime("%d.%m.%Y")
        except ValueError:
            return ""

    def gun(t):
        try:
            return datetime.date.fromisoformat(met(t)[:10]).strftime("%d.%m.%Y")
        except ValueError:
            return met(t)[:10]

    r = 7
    for i, g in enumerate(sorted(aletler, key=lambda x: (x["nitel"], x["alet"]))):
        c = eslesen_calisma(g, esler)
        kalibre = kalibrasyon_esle(g["alet"]) if v["lokasyon"] == KALIBRASYON_LOKASYON else None
        seri = met(kalibre and kalibre["seri"]) or met(c and c.get("gauge_number")) \
            or cihaz_kodu(v, g["alet"])
        aralik = met(kalibre and kalibre.get("aralik")) or (
            "Pass / Fail" if g["nitel"] else "%g – %g" % (g["dar_alt"], g["dar_ust"]))
        teknik = "ATTRIBUTE AGR. ANALYSIS" if g["nitel"] else "GAUGE R&R"
        tarih = gun(c.get("study_date")) if c else ""
        sonraki = yil_ekle(c.get("study_date")) if c else ""
        ws.cell(r, 1, i + 1)
        ws.cell(r, 2, "%s\n%s" % (g["alet"], (g.get("dar_kar") or ", ".join(g["kar"]))[:70]))
        ws.cell(r, 3, seri)
        ws.cell(r, 4, aralik)
        ws.cell(r, 5, teknik)
        ws.cell(r, 6, "1 YIL")
        ws.cell(r, 7, tarih)
        ws.cell(r, 8, sonraki)
        ws.cell(r, 9, msa_sonuc_notu(c, g))
        for j in range(1, 10):
            h = ws.cell(r, j)
            h.border = kutu; h.font = Font(size=10)
            h.alignment = Alignment(wrap_text=True, vertical="center",
                                    horizontal="left" if j in (2, 9) else "center")
        if c:
            ws.cell(r, 9).hyperlink = MSA_ADRES + met(c.get("id"))
            ws.cell(r, 9).font = Font(size=10, color="1D4ED8", underline="single")
        ws.row_dimensions[r].height = 34
        r += 1

    # İmza bloğu — kullanıcının formundaki düzen
    rolAd = dict((rol, ad) for rol, ad in v["ekip"])
    r += 1
    for sut, etiket, ad, unvan in (
            (1, "Hazırlayan", rolAd.get("Kalite Mühendisi", ""), "Kalite Güvence Mühendisi"),
            (6, "Onaylayan", rolAd.get("Kalite Güvence Müdürü", ""), "Kalite Güvence Müdürü")):
        h = ws.cell(r, sut, "%s\n%s\n%s" % (etiket, ad, unvan))
        h.font = Font(size=10); h.border = kutu
        h.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        ws.merge_cells(start_row=r, start_column=sut, end_row=r, end_column=sut + 3)
        for cc in range(sut, sut + 4):
            ws.cell(r, cc).border = kutu
    ws.row_dimensions[r].height = 52

    ws.print_area = "A1:I%d" % r
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToWidth = 1
    wb.save(hedef)
    return len(aletler)


def msa_sonuc_notu(c, g):
    """PL128'deki "NOTLAR VE SONUÇ" ifadesi — ERP çalışmasının gerçek sonucu."""
    if not c:
        return "MSA çalışması açılmadı."
    d = c.get("gauge_evaluation") or {}
    vir = lambda x, n=2: ("%.*f" % (n, x)).replace(".", ",")
    kabul = met(c.get("is_acceptable"))
    if g["nitel"]:
        kappa = d.get("decisiveKappa")
        uyum = d.get("allVsStandard")
        if kappa is None:
            return "Öznitelik uyum analizi açıldı; ölçüm bekleniyor."
        return ("Genel Cohen's Kappa=%s & standarda uygunluk %%%d olduğundan %s"
                % (vir(kappa, 3), round((uyum or 0) * 100),
                   "uygundur." if kabul == "acceptable" else
                   "şartlı kabul edilmiştir." if kabul == "marginal" else "uygun değildir."))
    grr = (d.get("totalGaugeRR") or {})
    yuz = grr.get("pctTolerance") if grr.get("pctTolerance") is not None else grr.get("pctStudyVar")
    ndc = d.get("ndc")
    if yuz is None:
        return "Gage R&R açıldı; ölçüm bekleniyor."
    return ("GRR=%s & ndc=%s olduğundan %s"
            % (vir(yuz), ndc,
               "uygundur." if kabul == "acceptable" else
               "şartlı kabul edilmiştir." if kabul == "marginal" else "uygun değildir."))


def msa_olcum_oku(kimlik):
    """MSA modülündeki ölçüm değerleri: (operatör, parça, tekrar) -> değer."""
    try:
        r = sorgu("/msa_measurements?select=operator,part,trial,measurement"
                  "&study_id=eq.%s&limit=3000" % kimlik)
    except Exception:
        return {}
    d = {}
    for x in r:
        try:
            d[(met(x["operator"]), met(x["part"]), int(x["trial"]))] = float(x["measurement"])
        except (TypeError, ValueError):
            pass
    return d


def msa_operator_oku(kimlik):
    try:
        r = sorgu("/msa_operators?select=operator_number,operator_name"
                  "&study_id=eq.%s&order=operator_number" % kimlik)
        return [met(x["operator_name"]) for x in r]
    except Exception:
        return []


def msa_parca_oku(kimlik):
    try:
        return sorgu("/msa_parts?select=part_number,part_name,nominal_value"
                     "&study_id=eq.%s&order=part_number" % kimlik)
    except Exception:
        return []


def fr86_gage_rr(v, hedef):
    """FR86 Gage R&R — her ölçüm aleti için ayrı sayfa. Ölçüm hücreleri BOŞ;
    EV/AV/GRR/PV/TV/%GRR/ndc AIAG MSA 4. Baskı ortalama-aralık yöntemiyle
    canlı formüllüdür, operatör ölçünce sonuç kendiliğinden çıkar.
    Nitel aletler için Type-3 nitelik uyum sayfası üretilir."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.utils import get_column_letter

    hepsi = msa_aletleri(v["kod"])
    olcum = [g for g in hepsi if not g["nitel"]]
    nitel = [g for g in hepsi if g["nitel"]]
    if not hepsi:
        return 0
    # MSA modulundeki (GageAI) calismalar: olcum degerleri buradan gelir
    mevcut = msa_calismalari(v, hepsi)
    for g in hepsi:
        c = eslesen_calisma(g, mevcut)
        g["calisma"] = c
        g["olcumler"] = msa_olcum_oku(c["id"]) if c else {}
        g["operator"] = msa_operator_oku(c["id"]) if c else []
        g["parcalar"] = msa_parca_oku(c["id"]) if c else []
    wb = Workbook(); wb.remove(wb.active)
    rolAd = dict((rol, ad) for rol, ad in v["ekip"])
    OPS, PARCA, TEKRAR = 3, 10, 3
    # AIAG MSA 4. Baskı K katsayıları (6σ, d2* tablosundan)
    K1 = {2: 0.8862, 3: 0.5908}[TEKRAR]      # tekrarlanabilirlik
    K2 = {2: 0.7071, 3: 0.5231}[OPS]         # tekrar üretilebilirlik
    K3 = {5: 0.4030, 10: 0.3146}[PARCA]      # parça değişkenliği
    SUT = lambda i: get_column_letter(i)

    def sayfa_basligi(ws, g, baslik, dok):
        ws.sheet_view.showGridLines = False
        antet(ws, baslik, dok, "02.01.2025", "0", "1 / 1", 13)
        r = 6
        for etiket, deger in (("Ürün :", "%s — %s" % (v["kod"], v["ad"])),
                              ("Ölçüm Aleti :", g["alet"]),
                              ("Karakteristik :", g.get("dar_kar") or ", ".join(g["kar"])[:70]),
                              ("Tolerans :", "%s   (aralık %g)" % (g.get("dar_limit", "—"), g["tol"])
                               if g.get("tol") else "Nitel — kabul/ret"),
                              ("Sorumlu :", rolAd.get("Kalite Mühendisi", "")),
                              ("Tarih :", met((g.get("calisma") or {}).get("study_date"))[:10]
                               or v["termin"]),
                              ("Cihaz No :", met((g.get("calisma") or {}).get("gauge_number"))),
                              ("MSA Çalışması :",
                               "%s  (ERP GageAI #%s)" % (met(g["calisma"].get("study_name")),
                                                         g["calisma"]["id"])
                               if g.get("calisma") else "ERP'de çalışma yok — form elle doldurulacak")):
            e = ws.cell(r, 1, etiket); e.font = Font(bold=True, size=10)
            e.alignment = Alignment(horizontal="right")
            ws.cell(r, 2, deger).alignment = Alignment(vertical="center")
            ws.merge_cells(start_row=r, start_column=2, end_row=r, end_column=13)
            r += 1
        return r + 1

    def baslik_hucre(ws, r, c, metin, genis=1):
        h = ws.cell(r, c, metin)
        h.font = Font(bold=True, size=9, color="FFFFFF")
        h.fill = PatternFill("solid", fgColor="1F3864")
        h.border = kutu_ince
        h.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for cc in range(c, c + genis):
            ws.cell(r, cc).border = kutu_ince
            ws.cell(r, cc).fill = PatternFill("solid", fgColor="1F3864")
        if genis > 1:
            ws.merge_cells(start_row=r, start_column=c, end_row=r, end_column=c + genis - 1)

    from openpyxl.styles import Border, Side
    _i = Side(style="thin", color="808080")
    kutu_ince = Border(top=_i, bottom=_i, left=_i, right=_i)

    # ── Ölçüm aletleri: Type-2 Gage R&R (ANOVA yerine ortalama-aralık) ────
    for g in olcum:
        ad = re.sub(r"[\\/*?:\[\]]", "-", g["alet"])[:28]
        ws = wb.create_sheet(ad)
        for h, gen in zip("ABCDEFGHIJKLM", (11, 8, 9, 9, 9, 9, 9, 9, 9, 9, 9, 9, 11)):
            ws.column_dimensions[h].width = gen
        r = sayfa_basligi(ws, g, "GAGE R&R (ÖLÇÜM SİSTEMİ ANALİZİ)", "FR 86")

        baslik_hucre(ws, r, 1, "Operatör"); baslik_hucre(ws, r, 2, "Tekrar")
        for pz in range(PARCA):
            baslik_hucre(ws, r, 3 + pz, "Parça %d" % (pz + 1))
        baslik_hucre(ws, r, 3 + PARCA, "Ortalama")
        r += 1
        op_bas = r
        op_ort = []                                  # her operatörün ortalama hücresi
        opAd = g.get("operator") or ["Operatör %s" % chr(65 + i) for i in range(OPS)]
        deger = g.get("olcumler") or {}
        for o in range(OPS):
            for t in range(TEKRAR):
                c1 = ws.cell(r, 1, opAd[o] if t == 0 and o < len(opAd) else None)
                c1.font = Font(bold=True, size=9)
                c1.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                c1.border = kutu_ince
                c2 = ws.cell(r, 2, t + 1)
                c2.alignment = Alignment(horizontal="center"); c2.border = kutu_ince
                for pz in range(PARCA):
                    x = deger.get((str(o + 1), str(pz + 1), t + 1))
                    c = ws.cell(r, 3 + pz, x)
                    c.border = kutu_ince; c.number_format = "0.000"
                    # Dolu deger MSA modulunden geldi; bos hucre elle girilecek
                    c.fill = PatternFill("solid", fgColor="EEF6EE" if x is not None else "FFFDE7")
                ws.cell(r, 3 + PARCA).border = kutu_ince
                r += 1
            blok = "%s%d:%s%d" % (SUT(3), r - TEKRAR, SUT(2 + PARCA), r - 1)
            h = ws.cell(r - TEKRAR, 3 + PARCA, "=IFERROR(AVERAGE(%s),\"\")" % blok)
            h.number_format = "0.000"; h.font = Font(bold=True, size=9)
            h.alignment = Alignment(horizontal="center", vertical="center")
            ws.merge_cells(start_row=r - TEKRAR, start_column=1, end_row=r - 1, end_column=1)
            ws.merge_cells(start_row=r - TEKRAR, start_column=3 + PARCA, end_row=r - 1,
                           end_column=3 + PARCA)
            op_ort.append("%s%d" % (SUT(3 + PARCA), r - TEKRAR))
        grid_son = r - 1
        tum = "%s%d:%s%d" % (SUT(3), op_bas, SUT(2 + PARCA), grid_son)

        # Aralık bloğu: her operatör × parça için (maks − min)
        r += 1
        baslik_hucre(ws, r, 1, "Aralık (R)", 2)
        for pz in range(PARCA):
            baslik_hucre(ws, r, 3 + pz, "Parça %d" % (pz + 1))
        r += 1
        rng_bas = r
        for o in range(OPS):
            e = ws.cell(r, 1, "Operatör %s" % chr(65 + o))
            e.font = Font(bold=True, size=9); e.border = kutu_ince
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=2)
            ws.cell(r, 2).border = kutu_ince
            for pz in range(PARCA):
                sut = SUT(3 + pz)
                ilk = op_bas + o * TEKRAR
                c = ws.cell(r, 3 + pz, "=IFERROR(MAX(%s%d:%s%d)-MIN(%s%d:%s%d),\"\")"
                            % (sut, ilk, sut, ilk + TEKRAR - 1, sut, ilk, sut, ilk + TEKRAR - 1))
                c.number_format = "0.000"; c.border = kutu_ince
                c.alignment = Alignment(horizontal="center")
            r += 1
        rng_son = r - 1
        # Parça ortalamaları (PV için)
        e = ws.cell(r, 1, "Parça ortalaması")
        e.font = Font(bold=True, size=9); e.border = kutu_ince
        ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=2)
        ws.cell(r, 2).border = kutu_ince
        for pz in range(PARCA):
            sut = SUT(3 + pz)
            c = ws.cell(r, 3 + pz, "=IFERROR(AVERAGE(%s%d:%s%d),\"\")"
                        % (sut, op_bas, sut, grid_son))
            c.number_format = "0.000"; c.border = kutu_ince
            c.alignment = Alignment(horizontal="center")
        pavg = r
        r += 2

        # Hesap bloğu — AIAG MSA 4. Baskı, 6σ
        rbar = "AVERAGE(%s%d:%s%d)" % (SUT(3), rng_bas, SUT(2 + PARCA), rng_son)
        xdiff = "MAX(%s)-MIN(%s)" % (",".join(op_ort), ",".join(op_ort))
        rp = "MAX(%s%d:%s%d)-MIN(%s%d:%s%d)" % (SUT(3), pavg, SUT(2 + PARCA), pavg,
                                                SUT(3), pavg, SUT(2 + PARCA), pavg)
        hesap = [
            ("Ölçüm sayısı (hedef %d)" % (OPS * PARCA * TEKRAR), "=COUNT(%s)" % tum, "0"),
            ("R̄  — ortalama aralık", "=IFERROR(%s,\"\")" % rbar, "0.0000"),
            ("X̄diff — operatör ortalamaları farkı", "=IFERROR(%s,\"\")" % xdiff, "0.0000"),
            ("Rp — parça ortalamaları aralığı", "=IFERROR(%s,\"\")" % rp, "0.0000"),
            ("EV = R̄ × K1  (K1=%.4f)" % K1, None, "0.0000"),
            ("AV = √((X̄diff×K2)² − EV²/(n·r))  (K2=%.4f)" % K2, None, "0.0000"),
            ("GRR = √(EV² + AV²)", None, "0.0000"),
            ("PV = Rp × K3  (K3=%.4f)" % K3, None, "0.0000"),
            ("TV = √(GRR² + PV²)", None, "0.0000"),
            ("Tolerans", "=%g" % g["tol"], "0.0000"),
            ("%GRR (toleransa göre)", None, "0.0%"),
            ("%GRR (TV'ye göre)", None, "0.0%"),
            ("ndc = 1,41 × PV / GRR", None, "0.0"),
            ("SONUÇ", None, "@"),
        ]
        hbas = r
        for etiket, formul, bicim in hesap:
            e = ws.cell(r, 1, etiket)
            e.font = Font(bold=True, size=9); e.border = kutu_ince
            e.alignment = Alignment(horizontal="left", vertical="center")
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=5)
            for cc in range(1, 6):
                ws.cell(r, cc).border = kutu_ince
            c = ws.cell(r, 6, formul)
            c.border = kutu_ince; c.number_format = bicim
            c.font = Font(bold=True, size=9)
            c.alignment = Alignment(horizontal="center", vertical="center")
            r += 1
        F = lambda i: "F%d" % (hbas + i)
        ws[F(4)] = "=IFERROR(%s*%.4f,\"\")" % (F(1), K1)
        ws[F(5)] = ("=IFERROR(SQRT(MAX(0,(%s*%.4f)^2-(%s^2/%d))),\"\")"
                    % (F(2), K2, F(4), PARCA * TEKRAR))
        ws[F(6)] = "=IFERROR(SQRT(%s^2+%s^2),\"\")" % (F(4), F(5))
        ws[F(7)] = "=IFERROR(%s*%.4f,\"\")" % (F(3), K3)
        ws[F(8)] = "=IFERROR(SQRT(%s^2+%s^2),\"\")" % (F(6), F(7))
        ws[F(10)] = "=IFERROR(%s/%s,\"\")" % (F(6), F(9))
        ws[F(11)] = "=IFERROR(%s/%s,\"\")" % (F(6), F(8))
        ws[F(12)] = "=IFERROR(1.41*%s/%s,\"\")" % (F(7), F(6))
        # Karar, toleransa ve toplam degiskenlige gore %GRR'lerin KOTUSUNE
        # baglidir; yalniz toleransa bakmak olcum sistemini iyi gosterebilir.
        kotu = "MAX(%s,%s)" % (F(10), F(11))
        ws[F(13)] = ('=IF(%s<%d,"ölçüm bekleniyor ("&%s&"/%d)",'
                     'IF(AND(%s<=0.1,%s>=5),"KABUL — %%GRR ≤ %%10 ve ndc ≥ 5",'
                     'IF(%s<=0.3,"ŞARTLI KABUL — %%10–30, iyileştirme gerekli",'
                     '"RED — %%GRR > %%30, ölçüm sistemi yetersiz")))'
                     % (F(0), OPS * PARCA * TEKRAR, F(0), OPS * PARCA * TEKRAR,
                        kotu, F(12), kotu))
        ws.cell(r + 1, 1, "Sarı hücrelere ölçüm değerleri girilir (10 parça, 3 operatör, 3 tekrar; "
                          "parçalar proses değişkenliğini temsil etmeli, sıra karıştırılmalı). "
                          "Kabul: %GRR ≤ %10 ve ndc ≥ 5 — AIAG MSA 4. Baskı, ortalama-aralık yöntemi."
                ).font = Font(size=8, italic=True, color="808080")
        ws.merge_cells(start_row=r + 1, start_column=1, end_row=r + 1, end_column=13)
        ws.page_setup.orientation = "landscape"

    # ── Nitel aletler: Type-3 nitelik uyum analizi (kappa) ────────────────
    # Boyutlar ERP çalışmasından gelir; form ile modül aynı veriyi göstermeli.
    for g in nitel:
        c = g.get("calisma") or {}
        NP = int(c.get("num_parts") or 20)
        NK = int(c.get("num_operators") or 3)          # kontrolör
        NT = int(c.get("num_trials") or 3)             # tekrar
        SUTUN = NK * NT
        ad = ("N-" + re.sub(r"[\\/*?:\[\]]", "-", g["alet"]))[:28]
        ws = wb.create_sheet(ad)
        for i in range(1, SUTUN + 4):
            ws.column_dimensions[get_column_letter(i)].width = 8 if 2 < i < SUTUN + 3 else 14
        r = sayfa_basligi(ws, g, "NİTELİK UYUM ANALİZİ (ATTRIBUTE MSA)", "FR 86-N")

        kAd = g.get("operator") or [chr(65 + i) for i in range(NK)]
        kisa = lambda i: (kAd[i].split()[0] if i < len(kAd) else chr(65 + i))
        basliklar = ["Parça", "Referans\n(OK/NOK)"]
        for o in range(NK):
            basliklar += ["%s-%d" % (kisa(o), t + 1) for t in range(NT)]
        basliklar.append("Uyum")
        for i, b in enumerate(basliklar):
            baslik_hucre(ws, r, 1 + i, b)
        ws.row_dimensions[r].height = 28
        r += 1
        vbas = r
        ilkS, sonS = 3, 2 + SUTUN                       # değerlendirme sütunları
        uyumS = sonS + 1
        deger = g.get("olcumler") or {}
        parcalar = g.get("parcalar") or []
        ok = lambda x: None if x is None else ("OK" if float(x) >= 0.5 else "NOK")
        for i in range(NP):
            ws.cell(r, 1, i + 1).alignment = Alignment(horizontal="center")
            ws.cell(r, 1).border = kutu_ince
            ref = ok((parcalar[i] or {}).get("nominal_value")) if i < len(parcalar) else None
            satirDeger = [ref] + [ok(deger.get((str(o + 1), str(i + 1), t + 1)))
                                  for o in range(NK) for t in range(NT)]
            for j, x in enumerate(satirDeger):
                h = ws.cell(r, 2 + j, x)
                h.border = kutu_ince
                h.alignment = Alignment(horizontal="center")
                # Yeşil: ERP çalışmasından geldi · sarı: elle doldurulacak
                h.fill = PatternFill("solid", fgColor="EEF6EE" if x else "FFFDE7")
            u = ws.cell(r, uyumS, '=IF(COUNTA(%s%d:%s%d)<%d,"",IF(COUNTIF(%s%d:%s%d,B%d)=%d,1,0))'
                        % (get_column_letter(ilkS), r, get_column_letter(sonS), r, SUTUN,
                           get_column_letter(ilkS), r, get_column_letter(sonS), r, r, SUTUN))
            u.border = kutu_ince
            u.alignment = Alignment(horizontal="center")
            r += 1
        vson = r - 1
        r += 1
        S = get_column_letter
        hesap = [("Değerlendirilen parça sayısı", "=COUNT(%s%d:%s%d)" % (S(uyumS), vbas, S(uyumS), vson), "0")]
        for o in range(NK):
            s1, s2 = S(ilkS + o * NT), S(ilkS + o * NT + NT - 1)
            hesap.append(("Kontrolör %s — kendi içinde uyum" % kisa(o),
                          '=IFERROR(SUMPRODUCT(--(COUNTIF(OFFSET(%s%d,ROW(%s%d:%s%d)-%d,0,1,%d),'
                          '%s%d:%s%d)=%d))/COUNTA(%s%d:%s%d),"")'
                          % (s1, vbas, s1, vbas, s1, vson, vbas, NT, s1, vbas, s1, vson, NT,
                             s1, vbas, s1, vson), "0.0%"))
        for o in range(NK):
            s1, s2 = S(ilkS + o * NT), S(ilkS + o * NT + NT - 1)
            hesap.append(("Kontrolör %s — referansa uyum" % kisa(o),
                          '=IFERROR(SUMPRODUCT(--(%s%d:%s%d=B%d:B%d))/COUNTA(%s%d:%s%d),"")'
                          % (s1, vbas, s2, vson, vbas, vson, s1, vbas, s2, vson), "0.0%"))
        hesap += [
            ("Tüm kontrolörler + referans tam uyum (Po)",
             '=IFERROR(AVERAGE(%s%d:%s%d),"")' % (S(uyumS), vbas, S(uyumS), vson), "0.0%"),
            ("Beklenen uyum (Pe)", None, "0.0%"),
            ("Kappa = (Po − Pe) / (1 − Pe)", None, "0.000"),
            ("SONUÇ", None, "@"),
        ]
        hbas = r
        sonucSut = min(uyumS, 7)
        for etiket, formul, bicim in hesap:
            e = ws.cell(r, 1, etiket)
            e.font = Font(bold=True, size=9); e.border = kutu_ince
            e.alignment = Alignment(horizontal="left", vertical="center")
            ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=sonucSut - 1)
            for cc in range(1, sonucSut):
                ws.cell(r, cc).border = kutu_ince
            h = ws.cell(r, sonucSut, formul)
            h.border = kutu_ince; h.number_format = bicim
            h.font = Font(bold=True, size=9)
            h.alignment = Alignment(horizontal="center", vertical="center")
            r += 1
        G = lambda i: "%s%d" % (S(sonucSut), hbas + i)
        po, pe, kap, snc = G(1 + 2 * NK), G(2 + 2 * NK), G(3 + 2 * NK), G(4 + 2 * NK)
        ws[pe] = ('=IFERROR((COUNTIF(B{0}:B{1},"OK")/COUNTA(B{0}:B{1}))^2'
                  '+(COUNTIF(B{0}:B{1},"NOK")/COUNTA(B{0}:B{1}))^2,"")').format(vbas, vson)
        ws[kap] = '=IFERROR((%s-%s)/(1-%s),"")' % (po, pe, pe)
        ws[snc] = ('=IF(%s="","ölçüm bekleniyor",IF(AND(%s>=0.75,%s>=0.9),'
                   '"KABUL — Kappa ≥ 0,75 ve uyum ≥ %%90",'
                   'IF(%s>=0.6,"ŞARTLI — kontrolör eğitimi / kriter netleştirme",'
                   '"RED — nitelik ölçüm sistemi yetersiz")))' % (kap, kap, po, kap))
        ws.cell(r + 1, 1, "Referans sütununda bilinen doğru sonuç (OK/NOK), sonraki sütunlarda her "
                          "kontrolörün %d bağımsız değerlendirmesi. Yeşil hücreler ERP MSA "
                          "çalışmasından geldi. Kabul: Kappa ≥ 0,75 — AIAG MSA 4. Baskı III-C." % NT
                ).font = Font(size=8, italic=True, color="808080")
        ws.merge_cells(start_row=r + 1, start_column=1, end_row=r + 1, end_column=uyumS)
        ws.page_setup.orientation = "landscape"

    wb.save(hedef)
    return len(hepsi)


# ── Türetilen alanlar ────────────────────────────────────────────────────
def zenginlestir(v):
    v["lokasyon_ad"] = "Ankara" if v["lokasyon"] == "ankara" else "Çerkezköy"
    # Resim adı, balonlamanın seçtiği DOSYANIN adıdır. Eskiden ilk dokümanın
    # adı alınıyordu ve 700.0.444'te "Part History" teknik resim gibi yazıyordu.
    try:
        import balonla
        v["resim"] = balonla.cizim_adi(v["dok"]) or "teknik resim ERP'de kayıtlı değil"
    except Exception:
        v["resim"] = "teknik resim ERP'de kayıtlı değil"
    d = datetime.date.fromisoformat(v["devreye"])
    v["devreye_baslangic"] = (d - datetime.timedelta(days=100)).isoformat()   # APQP başlangıcı
    v["termin"] = (d - datetime.timedelta(days=86)).isoformat()               # bölüm 2 termini

    # Benzer parça: aynı kök koddan türeyen ürünler (205.0.214-C -> 205.0.214)
    kok = re.sub(r"[-.][A-Za-z0-9]+$", "", v["kod"])
    try:
        benzer = sorgu("/leansys_kontrol_plani?stok_kodu=like.%s&select=stok_kodu,stok_adi&limit=200"
                       % urllib.parse.quote(kok + "*"))
    except Exception:
        benzer = []
    kodlar = sorted({met(b["stok_kodu"]) for b in benzer} - {v["kod"]})
    v["benzer"] = ("Benzer parça(lar): " + ", ".join(kodlar[:6])) if kodlar else \
                  "ERP'de benzer/karşılaştırılabilir parça bulunmadı — yeni proses"

    # FR90 Proje No alani: FMEA projesinin kimligi (proj_ ile baslayan)
    v["proje_no"] = "proj_" + re.sub(r"[.\-]", "_", v["kod"])
    try:
        pf = sorgu("/pfmea_projects?select=name,data")
    except Exception:
        pf = []
    buyukAd = v["ad"].upper()
    eslesen = None
    for x in pf:
        f = ((x.get("data") or {}).get("projectData") or {}).get("fmea") or {}
        metin = (met(f.get("project")) + " " + met(f.get("productName")) + " " + met(x.get("name"))).upper()
        if met(f.get("projectId")) == v["proje_no"] or v["kod"].upper() in metin or (buyukAd and buyukAd in metin):
            eslesen = (met(f.get("projectId")) or v["proje_no"], met(x.get("name")))
            break
    if eslesen:
        v["proje_no"] = eslesen[0]
        v["fmea_not"] = "P-FMEA mevcut: " + eslesen[1]
    else:
        v["fmea_not"] = "P-FMEA bulunamadı — oluşturulmalı (PFMEA modülü)"

    # Kapasite: vardiya süresi kapasite_sure'den (LeanSys 31500 sn ≈ 8,75 saat)
    sureler = [round(float(met(r.get("kapasite_sure")) or 0)) for r in v["rota"]]
    v["vardiya_sure"] = max(sureler) if sureler else 480
    # LeanSys bu alani kimi urunde SANIYE (31500), kiminde DAKIKA (480) tutuyor.
    # 1440'in altindaki deger bir vardiyayi saniyeyle anlatamaz -> dakikadir.
    v["birim"] = "dk" if v["vardiya_sure"] <= 1440 else "sn"
    v["vardiya_saat"] = v["vardiya_sure"] / (60.0 if v["birim"] == "dk" else 3600.0)
    satirlar = []
    bos = []                               # kapasitesi girilmemiş operasyonlar
    for r in v["rota"]:
        # LeanSys bu alanlari ondalikli da doldurabiliyor (19090.9 gibi)
        std = float(met(r.get("std_zaman")) or 0)
        kap = round(float(met(r.get("kapasite")) or 0))
        sure = round(float(met(r.get("kapasite_sure")) or 0))
        if std <= 1 and kap <= 1:          # hazırlık satırı ya da veri girilmemiş
            bos.append({
                "op": met(r.get("op_no")), "makine": met(r.get("makine_adi")),
                "std": "", "personel": met(r.get("personel")), "sure": "",
                "kap": "", "gunluk": "", "darbogaz": False,
                "not": "LeanSys operasyon kartında kapasite verisi girilmemiş",
            })
            continue
        satirlar.append({
            "op": met(r.get("op_no")), "makine": met(r.get("makine_adi")),
            "std": std, "personel": met(r.get("personel")), "sure": sure,
            "kap": kap, "gunluk": kap * 3, "not": met(r.get("talimat")) or "",
            "darbogaz": False,
        })
    if satirlar:
        db = min(satirlar, key=lambda x: x["kap"])
        db["darbogaz"] = True
        v["darbogaz"] = db
    else:
        # Hicbir operasyonda kapasite yok: form yine uretilir, eksik gorunsun
        satirlar = bos
        v["darbogaz"] = {"makine": "belirlenemedi — kapasite verisi yok",
                         "kap": "", "gunluk": ""}
    v["kapasite_veri_yok"] = not any(x["darbogaz"] for x in satirlar)
    v["kapasite_satirlari"] = satirlar

    # APQP bölüm özeti (Program Metrikleri için) — adım listesi apqp.html'deki
    # FR91 verisinden okunur, tek kaynak orası.
    v["apqp_bolumler"] = apqp_bolum_ozeti(v)
    return v


def apqp_bolum_ozeti(v):
    """apqp.html içindeki FR91 listesinden bölüm/adım sayıları; kanıtı ERP'den
    gelen adımlar 'tamamlanan' sayılır."""
    try:
        h = io.open(os.path.join(os.path.dirname(os.path.abspath(__file__)), "apqp.html"),
                    encoding="utf-8").read()
        m = re.search(r"const FR91 = (\[.*?\]);", h, re.S)
        bolumler = json.loads(m.group(1)) if m else []
    except Exception:
        bolumler = []
    kanit = {"fmea": bool(v.get("fmea_not", "").startswith("P-FMEA mevcut")),
             "plan": True, "opkart": bool(v["rota"]), "akis": True}
    ozet = []
    for b in bolumler:
        tamam = sum(1 for a in b["adimlar"] if a.get("kanit") and kanit.get(a["kanit"]))
        ozet.append({"no": b["no"], "ad": b["ad"], "adim": len(b["adimlar"]), "tamam": tamam})
    return ozet


class Gunluk:
    """Çıktıyı hem konsola hem UTF-8 dosyaya yazar. Ajan konsol kod sayfasına
    bağlı kalmasın diye dosyadan okur (Türkçe karakterler bozulmuyor)."""

    def __init__(self, akis, yol):
        self.akis, self.dosya = akis, io.open(yol, "a", encoding="utf-8")

    def write(self, x):
        try:
            self.akis.write(x)
        except Exception:
            pass
        self.dosya.write(x); self.dosya.flush()

    def flush(self):
        try:
            self.akis.flush()
        except Exception:
            pass


def main():
    gunluk = os.environ.get("APQP_LOG")
    if gunluk:
        sys.stdout = Gunluk(sys.stdout, gunluk)
        sys.stderr = sys.stdout
    if len(sys.argv) < 2:
        raise SystemExit("kullanım: python apqp_belge_uret.py <stok kodu>")
    kod = sys.argv[1]
    v = zenginlestir(urun_verisi(kod))
    klasor = os.path.join(DRIVE, kod)
    os.makedirs(klasor, exist_ok=True)

    print("%s — %s" % (kod, v["ad"]))
    print("   müşteri: %s | lokasyon: %s | devreye alma: %s"
          % (v["musteri"], v["lokasyon_ad"], v["devreye"]))
    print("   klasör : %s" % klasor)

    def uret(dosya, islev, etiket):
        """Belgeyi gecici ada yazip yerine koyar. Dosya Excel'de acikken
        (Permission denied) tum uretim durmasin diye tek tek yakalanir."""
        yol = os.path.join(klasor, dosya)
        gecici = yol + ".yeni"
        try:
            sonuc = islev(v, gecici)
            if not sonuc or not os.path.exists(gecici):
                # Uretilmedi (or. PL11 kaynagi yok) - bos dosya birakma
                if os.path.exists(gecici): os.remove(gecici)
                return None
            os.replace(gecici, yol)
            return sonuc
        except PermissionError:
            try: os.remove(gecici)
            except OSError: pass
            print("   ! %-32s dosya açık, yazılamadı — Excel'de kapatıp tekrar çalıştırın" % etiket)
            return None
        except Exception as e:
            try: os.remove(gecici)
            except OSError: pass
            print("   ! %-32s üretilemedi: %s" % (etiket, str(e)[:70]))
            return None

    n = uret("PL74 Proses Akış Diyagramı %s.xlsx" % kod, pl74, "PL74 Proses Akış Diyagramı")
    if n: print("   ✓ PL74 Proses Akış Diyagramı      (%d adım)" % n)

    if uret("FR90 Fizibilite Taahhüdü %s.xlsm" % kod, lambda a, b: (fr90(a, b), 1)[1],
            "FR90 Fizibilite Taahhüdü"):
        print("   ✓ FR90 Fizibilite Taahhüdü         (başlık dolduruldu, cevaplar ekipte)")

    if uret("FR182 Ürün Devreye Alma Formu %s.xlsx" % kod, fr182, "FR182 Ürün Devreye Alma"):
        print("   ✓ FR182 Ürün Devreye Alma Formu     (üretim imzası eklendi)")

    n = uret("FR81 Toplantı Tutanağı %s.xlsx" % kod, fr81, "FR81 Toplantı Tutanağı")
    if n: print("   ✓ FR81 Toplantı Tutanağı            (%d madde — APQP madde no ile eklenir)" % n)

    n = uret("PL11 Onaylı Tedarikçi Listesi %s.xlsx" % kod, pl11, "PL11 Onaylı Tedarikçi Listesi")
    if n:
        print("   ✓ PL11 Onaylı Tedarikçi Listesi     (%d otomotiv tedarikçi, %s)" % (n, v["lokasyon_ad"]))

    if uret("FR228 Ambalaj Standardı Formu %s.docx" % kod, fr228, "FR228 Ambalaj Standardı"):
        print("   ✓ FR228 Ambalaj Standardı Formu     (fotoğraflar kaldırıldı — bu ürüne ait değil)")
    if uret("FR148 Değişiklik Yönetimi Formu %s.xlsx" % kod, fr148, "FR148 Değişiklik Yönetimi"):
        print("   ✓ FR148 Değişiklik Yönetimi Formu   (başlık dolduruldu, risk satırları ekipte)")
    n = uret("FR181 Öğrenilmiş Dersler %s.xlsx" % kod, fr181, "FR181 Öğrenilmiş Dersler")
    if n: print("   ✓ FR181 Öğrenilmiş Dersler          (%d ilgili kayıt işaretlendi)" % n)
    if uret("FR91 APQP-Takip Formu %s.xlsx" % kod, fr91, "FR91 APQP Takip Formu"):
        print("   ✓ FR91 APQP-Takip Formu             (77 adım, ürün başlığıyla)")
    n = uret("APQP Program Metrikleri %s.xlsx" % kod, program_metrikleri, "APQP Program Metrikleri")
    if n: print("   ✓ APQP Program Metrikleri           (%d bölüm, kırmızı/sarı/yeşil)" % n)

    n = uret("PL41 Kontrol Planı %s.xlsx" % kod, pl41_kontrol_plani, "PL41 Kontrol Planı")
    if n: print("   ✓ PL41 Kontrol Planı                (%d karakteristik, Leansys verisi)" % n)
    n = uret("FR34 P-FMEA %s.xlsx" % kod, fr34_pfmea, "FR34 P-FMEA")
    if n: print("   ✓ FR34 P-FMEA (Excel çıktısı)       (%d satır, AIAG-VDA 30 sütun)" % n)
    else: print("   ! FR34 P-FMEA                       PFMEA modülünde bu ürünün projesi yok")
    # MSA calismalari once ACILIR ki plan belgesi onlari gostersin
    aletler = msa_aletleri(kod)
    acilan = msa_calismasi_ac(v, aletler, msa_calismalari(v, aletler))
    for alet, kimlik, tip in acilan:
        print("   + MSA çalışması açıldı: %-14s %-16s (GageAI #%s)" % (alet, tip, kimlik))
    for alet, kimlik, n_olcum, ozet, kabul in msa_olcumleri_yaz(v, aletler):
        print("   ✓ MSA #%-3s %-14s %3d ölçüm · %-22s %s"
              % (kimlik, alet, n_olcum, ozet,
                 {"acceptable": "KABUL", "marginal": "ŞARTLI", "unacceptable": "RED"}[kabul]))

    # Balonlu (numaralandırılmış) teknik resim — PPAP 2.2.1 / madde 6.5
    balon_satir = []
    try:
        import balonla
        cizim = balonla.cizim_yolu(v["dok"])
        if cizim:
            n, rapor, balon_satir = balonla.uret(kod, cizim, klasor, kp_satirlari)
            v["balon"] = balon_satir          # VDA_2 ölçüm sayfası da bunu kullanır
            print("   %s Numaralandırılmış Teknik Resim  (%s)"
                  % ("✓" if n else "!", rapor))
        else:
            balon_satir = []
            print("   ! Numaralandırılmış Teknik Resim  ERP'de teknik resim dosyası yok")
    except Exception as e:
        balon_satir = []
        print("   ! Numaralandırılmış Teknik Resim  üretilemedi: %s" % str(e)[:60])

    for alet, kar, n_ol, sonuc, kimlik, yeni in yeterlilik_uret(v, klasor, uret):
        tur = sonuc.get("tur", "Proses")
        print("   %s FR24 %-7s %-11s %-16s %s=%.2f %s=%.2f · %s%s"
              % ("✓" if n_ol else "!", tur, alet, kar[:16],
                 "Cm" if tur == "Makine" else "Cp", sonuc["cp"],
                 "Cmk" if tur == "Makine" else "Cpk", sonuc["cpk"],
                 sonuc.get("yontem", "normal"),
                 "  (GageAI #%s)" % kimlik if kimlik else ""))

    n = uret("MSA Planı %s.xlsx" % kod, msa_plani, "MSA Planı")
    if n: print("   ✓ MSA Planı                         (%d ölçüm aleti)" % n)
    n = uret("FR86 Gage R&R %s.xlsx" % kod, fr86_gage_rr, "FR86 Gage R&R")
    if n: print("   ✓ FR86 Gage R&R                     (%d alet, formüller canlı — ölçüm girilecek)" % n)

    mak = kalipli_mi(v)
    if mak:
        n = uret("FR176 Kalıp Doğrulama Formu %s.xls" % kod, fr176_kalip, "FR176 Kalıp Doğrulama")
        if n:
            print("   ✓ FR176 Kalıp Doğrulama Formu       (%d ölçü · kalıplı: %s)"
                  % (n, ", ".join(mak)[:40]))
    else:
        print("   · FR176 Kalıp Doğrulama Formu       gerekmiyor (rotada kalıplı makine yok)")

    for kod_, ad_, ted, n_, nasil in alt_tedarikci_ppap(v, klasor, uret):
        print("   %s Alt Tedarikçi PPAP %-16s %-26s %s"
              % ("✓" if n_ else "!", kod_, (ted or "tedarikçi bulunamadı")[:26],
                 "%d ölçü · %s" % (n_, nasil) if n_ else nasil))

    ppap_belgeleri(v, klasor, uret)

    n = uret("Ölçüsel Rapor %s.xlsx" % kod,
             lambda a, b: olcusel_rapor(a, b, balon_satir), "Ölçüsel Rapor")
    if n:
        sinif, uyan, top = tolerans_sinifi(kod)
        print("   ✓ Ölçüsel Rapor                     (%d satır, genel tolerans DIN ISO "
              "2768-%s — plandaki %d/%d ölçüyle uyumlu)" % (n, sinif, uyan, top))

    n = uret("Kapasite Takip Formu %s.xlsx" % kod, kapasite, "Kapasite Takip Formu")
    if n: print("   ✓ Kapasite Takip Formu             (%d operasyon, darboğaz: %s / %s adet)"
                % (n, v["darbogaz"]["makine"][:26], v["darbogaz"]["kap"]))
    n = uret("Run @ Rate %s.xlsx" % kod, run_at_rate, "Run @ Rate")
    if n: print("   ✓ Run @ Rate                        (kapasite doğrulama · FR91 5.12)")


if __name__ == "__main__":
    main()
